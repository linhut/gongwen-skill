# -*- coding: utf-8 -*-
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# https://github.com/linhut/gongwen-skill
# Licensed under the MIT License. See the LICENSE file for details.
#
# -*- coding: utf-8 -*-
"""
座签默认模板生成器（方案六 P2-2：内置默认模板）。

座签排版严重依赖 .dotx 模板（表格结构、行高、字体等均由模板承载）。
本模块在 skill 仓库中生成并附带一个与 WPS 座签模板参数一致的默认 .dotx
模板，使用户零配置即可使用 table-signs 命令；用户仍可通过 --template 覆盖。

默认模板参数（WPS 模板实测）：
- 纸张：A4 纵向（21×29.7cm）
- 页边距：上 2.69cm 下 2.69cm 左 2.79cm 右 2.79cm
- 页眉距离 1.50cm / 页脚距离 1.75cm
- 表格：1 列，行高 9.65cm（At Least），对齐居中，宽度 541.9pt
- 字体：华文新魏，156pt（两字名），不加粗
- 占位符：Jose AI
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# 默认模板文件路径（相对本模块：engine/templates/table_sign.dotx）
TEMPLATES_DIR = Path(__file__).resolve().parent / "templates"
DEFAULT_TEMPLATE_PATH = TEMPLATES_DIR / "table_sign.dotx"

# 默认模板参数（与 WPS 座签模板实测一致）
PLACEHOLDER = "Jose AI"
FONT_NAME = "华文新魏"
FONT_SIZE_PT = 156  # 两字名
TABLE_COLUMNS = 1
TABLE_ROWS = 2       # 每名占 2 行（正反各 1 行，对齐示例模板结构）
TABLE_ROW_HEIGHT_CM = 9.65
TABLE_WIDTH_PT = 541.9


def _set_cell_shading_and_borders(table) -> None:
    """为表格设置边框（示例模板：single sz=4 color=auto），使座签在打印时轮廓清晰。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 边框定义（对齐示例文档：val=single sz=4 color=auto）
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")  # 0.5pt
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_row_height(row, height_cm: float) -> None:
    """设置行高（规则 At Least），确保单元格有足够的竖排空间。"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # 1cm ≈ 567 twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def build_default_template(output_path: Path | None = None) -> Path:
    """构建与 WPS 座签模板参数一致的默认 .dotx 模板。

    Args:
        output_path: 输出 .dotx 路径（默认 engine/templates/table_sign.dotx）

    Returns:
        生成的模板文件路径
    """
    out = Path(output_path) if output_path else DEFAULT_TEMPLATE_PATH
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 页面设置：A4 纵向 + 页边距 + 页眉/页脚距离（对齐示例模板实测）
    section = doc.sections[0]
    section.page_width = Cm(20.99)
    section.page_height = Cm(29.70)
    section.top_margin = Cm(2.69)
    section.bottom_margin = Cm(2.69)
    section.left_margin = Cm(2.79)
    section.right_margin = Cm(2.79)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)

    # 1 列表格（每名占 2 行：正反各 1 行，对齐示例 42 行整卷模板的行结构）：
    # 行高 9.65cm At Least，对齐居中，宽度 541.9pt
    table = doc.add_table(rows=TABLE_ROWS, cols=TABLE_COLUMNS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 设置表格宽度（541.9pt ≈ 19.12cm，对齐示例 10838 twips）
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 修复：doc.add_table 会自动创建 <w:tblW w:w="0" type="auto"/>，若直接 append 新 tblW
    # 会形成两个节点，python-docx 读取到的是 auto 0 值。先移除旧节点再写入正确宽度。
    old_tblW = tblPr.find(qn("w:tblW"))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(TABLE_WIDTH_PT * 20)))  # 1pt = 20 twips → 10838
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    _set_cell_shading_and_borders(table)

    # 每行：行高 9.65cm At Least + 单元格内容（华文新魏 156pt 不加粗，占位符）
    for row in table.rows:
        _set_row_height(row, TABLE_ROW_HEIGHT_CM)
        cell = row.cells[0]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(PLACEHOLDER)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(FONT_SIZE_PT)
        run.font.bold = False
        # 段落无缩进，垂直方向尽量居中
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

    # 保存为 .dotx 模板
    doc.save(str(out))
    return out


def ensure_default_template() -> Path:
    """确保默认模板存在；不存在则构建。返回模板路径。"""
    if not DEFAULT_TEMPLATE_PATH.exists():
        build_default_template()
    return DEFAULT_TEMPLATE_PATH
