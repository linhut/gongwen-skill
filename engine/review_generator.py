# -*- coding: utf-8 -*-
"""
审稿流转文书生成器 —— 为五角色审稿机制提供模板文档。

用于党政机关公文起草后的审核流转流程，生成结构化审稿单：
  - 完整版（5 角色）：撰稿人→业务审核→文字校对→综合核稿→领导签发
  - 精简版（3 角色）：撰稿人→业务+文字复合审核→综合负责人终审

用法（CLI）：
  python gongwen.py review 通知 -o 审稿单.docx
  python gongwen.py review 请示 --scheme compact -o 审稿-精简版.docx
"""
from __future__ import annotations
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from engine.utils.logger import logger
# 跨模块#2 修复：统一字体设置入口（避免直接 font.name 导致 eastAsia 回退）
from engine.core.document.font_utils import set_run_font

# 审稿角色定义（完整版 5 角色）
FULL_SCHEME = [
    ("①撰稿人", "起草人", "业务事实准确、数据来源可靠、逻辑通顺、覆盖完整"),
    ("②业务审核人", "处室/部门负责人", "业务口径、事实真实、权责分工、工作可行性"),
    ("③文字校对岗", "综合岗", "错别字/标点/语病、序号规范、术语统一、格式合规"),
    ("④综合核稿人", "综合办/专班负责人", "政策口径、全局逻辑、风险排查、行文基调"),
    ("⑤签发领导", "分管领导", "核心观点确认、重大工作安排、是否同意印发/报送"),
]

# 审稿角色定义（精简版 3 角色）
COMPACT_SCHEME = [
    ("①撰稿人", "起草人", "业务事实准确、数据来源可靠、逻辑通顺、覆盖完整"),
    ("②业务+文字复合审核", "业务处室+综合岗", "业务口径+文字格式双审"),
    ("③综合负责人终审", "综合办负责人", "全局把关、风险排查、是否同意报送"),
]


def _set_cell_shading(cell, color_hex: str) -> None:
    """设置单元格底色。"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_cell(row, text: str, bold: bool = False, width_mm: float = None,
                     shading: str = None, alignment: str = 'left') -> None:
    """向表格行添加带样式的单元格。"""
    cell = row.add_cell()
    if width_mm:
        cell.width = Mm(width_mm)
    if shading:
        _set_cell_shading(cell, shading)
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.font.size = Pt(12)
    set_run_font(run, '仿宋_GB2312')  # 跨模块#2: 统一字体入口
    if bold:
        run.bold = True


def _add_styled_cell_by_index(table, row_idx: int, col_idx: int, text: str,
                              bold: bool = False, width_mm: float = None,
                              shading: str = None, alignment: str = 'left') -> None:
    """按行列索引填充已有单元格样式。"""
    cell = table.cell(row_idx, col_idx)
    if width_mm:
        cell.width = Mm(width_mm)
    if shading:
        _set_cell_shading(cell, shading)
    # 清空默认段落
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.font.size = Pt(12)
    set_run_font(run, '仿宋_GB2312')  # 跨模块#2: 统一字体入口
    if bold:
        run.bold = True


def _set_table_borders(table) -> None:
    """设置表格边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)


def generate_review_template(
    doc_type: str,
    output_path: str | Path,
    scheme: str = "full",
    doc_title: str = "",
) -> Path:
    """
    生成审稿流转单。

    Args:
        doc_type: 公文类型（通知/请示/报告/函等）
        output_path: 输出 .docx 路径
        scheme: 'full'（5角色）或 'compact'（3角色）
        doc_title: 待审文稿标题（可选）

    Returns:
        生成的 .docx 文件路径
    """
    doc = Document()

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 标题
    title_text = f"公文审稿流转单"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(18)
    # S10 修复：使用 set_run_font 设置 4 属性字体，避免直接 font.name 导致 eastAsia 回退
    from engine.core.document.font_utils import set_run_font
    set_run_font(title_run, '黑体')
    title_run.bold = True

    # 基本信息区
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    scheme_label = "完整版（5角色）" if scheme == "full" else "精简版（3角色）"
    info_text = f"公文类型：{doc_type}   审核方案：{scheme_label}"
    if doc_title:
        info_text += f"\n文稿标题：{doc_title}"
    info_run = info.add_run(info_text)
    info_run.font.size = Pt(12)
    set_run_font(info_run, '仿宋_GB2312')  # 跨模块#2: 统一字体入口

    doc.add_paragraph()  # 空行

    # 选择角色列表
    roles = FULL_SCHEME if scheme == "full" else COMPACT_SCHEME

    # 审稿表格：角色 | 责任主体 | 审稿侧重点 | 审稿意见 | 是否通过 | 签名
    table = doc.add_table(rows=1 + len(roles), cols=6)
    table.style = 'Table Grid'
    _set_table_borders(table)

    # 表头
    headers = ['审稿角色', '责任主体', '审稿侧重点', '审稿意见', '是否通过', '签名/日期']
    for ci, h in enumerate(headers):
        _add_styled_cell_by_index(table, 0, ci, h, bold=True, width_mm=22,
                                  shading='D9E2F3', alignment='center')

    # 数据行
    for ri, (role, responsible, focus) in enumerate(roles):
        row_idx = ri + 1
        _add_styled_cell_by_index(table, row_idx, 0, role, bold=True, width_mm=22)
        _add_styled_cell_by_index(table, row_idx, 1, responsible, width_mm=26)
        _add_styled_cell_by_index(table, row_idx, 2, focus, width_mm=42)
        _add_styled_cell_by_index(table, row_idx, 3, '', width_mm=38)  # 审稿意见（留空手写）
        _add_styled_cell_by_index(table, row_idx, 4, '□通过 □退回 □修改', width_mm=22, alignment='center')
        _add_styled_cell_by_index(table, row_idx, 5, '', width_mm=22)

    # 设置列宽
    if table.columns:
        widths = [22, 26, 42, 38, 22, 22]
        for ci, w in enumerate(widths):
            for row in table.rows:
                row.cells[ci].width = Mm(w)

    doc.add_paragraph()  # 空行

    # 审核流转记录
    flow_title = doc.add_paragraph()
    flow_title_run = flow_title.add_run("流转记录")
    flow_title_run.font.size = Pt(14)
    set_run_font(flow_title_run, '黑体')  # 跨模块#2: 统一字体入口
    flow_title_run.bold = True

    flow_table = doc.add_table(rows=4, cols=3)
    flow_table.style = 'Table Grid'
    _set_table_borders(flow_table)

    flow_headers = ['流转环节', '完成时间', '备注']
    for ci, h in enumerate(flow_headers):
        _add_styled_cell_by_index(flow_table, 0, ci, h, bold=True, shading='D9E2F3', alignment='center')
    for ri in range(1, 4):
        for ci in range(3):
            _add_styled_cell_by_index(flow_table, ri, ci, '', width_mm=50)

    # 使用说明
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        "使用说明：\n"
        "1. 每轮审稿完成后由对应角色填写审稿意见并签名，流转至下一环节\n"
        "2. 「是否通过」标注通过/退回/修改，退回需附具体修改意见\n"
        "3. 全部角色审核通过后进入签发环节\n"
        "4. 本单随文稿流转，最终归档保存"
    )
    note_run.font.size = Pt(10)
    set_run_font(note_run, '仿宋_GB2312')  # 跨模块#2: 统一字体入口
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    out = Path(output_path)
    doc.save(str(out))
    logger.info(f"审稿流转单已生成: {out}（{scheme_label}）")
    return out
