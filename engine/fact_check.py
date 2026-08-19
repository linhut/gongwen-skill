# -*- coding: utf-8 -*-
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# https://github.com/linhut/gongwen-skill
# Licensed under the MIT License. See the LICENSE file for details.
#
"""
事实核验模块 —— 对用户文档中的关键人事信息做主动交叉核验。

依据「公文技能提质方案 v2.1」问题三设计：
- Step 1 实体提取：从用户文档提取领导姓名+职务 / 机构全称+简称 / 项目名称 / 文号 / 关键数据
- Step 2 基准构建：从背景资料（docx/pdf/md/txt/URL）提取同类实体，构建事实基准表
- Step 3 互联网交叉核验：对"待核验"关键项（领导姓名、机构全称）搜索官方来源比对
- Step 4 标记输出：一致→不标记；存疑→批注提醒；无法核验→批注标注"未经核验，请人工确认"
- 核验报告：生成 .fact_check.json 摘要（命令行 + 文件）

用法：
  from fact_check import FactChecker, run_fact_check
  report = run_fact_check("文档.docx", ["背景资料1.pdf", "背景资料2.md"])
"""
from __future__ import annotations
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from engine.utils.logger import logger


# ---------------------------------------------------------------------------
#  数据结构
# ---------------------------------------------------------------------------

@dataclass
class Entity:
    """从文档中提取的实体。"""
    entity_type: str      # person / org / project / doc_no / data
    entity_name: str      # 实体名（如"XXX"）
    doc_attribute: str = ""   # V2 新增：文档中的属性（如"省民宗委党组成员、副主任"）
    doc_context: str = ""     # V2 新增：完整上下文（如"省民宗委党组成员、副主任XXX..."）
    context: str = ""     # 上下文片段
    paragraph_index: int = 0
    status: str = "待核验"  # 待核验 / 已确认 / 存疑 / 无法核验
    source: str = ""      # 核验来源
    note: str = ""        # 说明


@dataclass
class FactCheckReport:
    """事实核验报告。"""
    document: str = ""
    entities: List[Entity] = field(default_factory=list)
    confirmed: List[Entity] = field(default_factory=list)
    doubtful: List[Entity] = field(default_factory=list)
    unverified: List[Entity] = field(default_factory=list)

    def summary_text(self) -> str:
        """生成命令行摘要。"""
        lines = [
            "═══ 事实核验摘要 ═══",
            f"文档：{self.document}",
            f"核验实体：{len(self.entities)} 项",
            "─────────────────",
        ]
        if self.confirmed:
            lines.append(f"✅ 已确认：{len(self.confirmed)} 项")
            for e in self.confirmed:
                lines.append(f"- {e.entity_name}：{e.note}")
        if self.doubtful:
            lines.append(f"⚠️ 存疑：{len(self.doubtful)} 项")
            for e in self.doubtful:
                lines.append(f"- [P1] {e.entity_name} → {e.note}")
        if self.unverified:
            lines.append(f"❓ 未核实：{len(self.unverified)} 项")
            for e in self.unverified:
                lines.append(f"- [P2] {e.entity_name} → 无官方来源确认")
        lines.append("─────────────────")
        lines.append(f"详情见：{self.document}.fact_check.json")
        lines.append("═══════════════════")
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "document": self.document,
            "entities": [
                {"entity_type": e.entity_type, "entity_name": e.entity_name,
                 "context": e.context, "paragraph_index": e.paragraph_index,
                 "status": e.status, "source": e.source, "note": e.note}
                for e in self.entities
            ],
            "confirmed": [e.entity_name for e in self.confirmed],
            "doubtful": [{"name": e.entity_name, "note": e.note} for e in self.doubtful],
            "unverified": [e.entity_name for e in self.unverified],
        }


# ---------------------------------------------------------------------------
#  实体提取规则
# ---------------------------------------------------------------------------

# 职务后缀（用于人名+职务识别）
_TITLE_SUFFIXES = [
    "主任", "副主任", "书记", "副书记", "部长", "副部长", "厅长", "副厅长",
    "局长", "副局长", "处长", "副处长", "秘书长", "副秘书长", "执行副秘书长",
    "部长助理", "董事长", "总经理", "总工程师", "党组成员", "常务副部长",
    "常务副部长", "一级调研员", "机关党委书记",
]

# 机构常见后缀
_ORG_SUFFIXES = [
    "办公室", "委员会", "联合会", "协会", "研究院", "设计院", "集团",
    "有限公司", "公司", "部", "厅", "局", "处", "办", "中心", "学校", "大学", "医院",
]
# N2 修复：常见动词/助词组合（过滤非实体片段，如"会议听取""以...护航"）
_NON_ENTITY_PATTERNS = [
    r'^会议听取', r'^会议指出', r'^会议认为', r'^会议强调', r'^会议要求',
    r'^以', r'^为', r'^通过', r'^围绕', r'^聚焦', r'^持续', r'^不断',
    r'^推动', r'^加快', r'^实现', r'^强化', r'^深化', r'^健全', r'^完善',
    r'^确保', r'^坚持', r'^充分', r'^切实', r'^全力',
]
# 机构名最小长度（M3 修复：从 4 提升到 5，避免"各部""副部"等过短残片）
_ORG_MIN_LEN = 5
# M3 修复：单字机构后缀右边界扩展映射（"信息技术部"+"门"→"信息技术部门"）
_ORG_SUFFIX_EXTENSIONS = {
    "部": ["门", "长"],
    "副部": ["长"],
    "办": ["公室"],
}
# M3 修复：句子特征词（包含 2 个以上则判定为短句而非实体）
_SENTENCE_INDICATORS = [
    "赋能", "保障", "贴合", "服务", "构建", "坚持", "推动", "深化",
    "既", "又", "实现", "围绕", "聚焦", "促进", "确保", "凝聚",
]
# P1 修复：不完整职务/机构后缀黑名单（"常务副部""副部""各部"等残片）
# 注意："限公司"不放此处——合法"XX有限公司"以"限公司"结尾，应用 ^ 开头检测拦截独立残片
_INCOMPLETE_SUFFIX_BLACKLIST = ["常务副部", "副部", "各部", "部的"]
# P1 修复：动词/系词开头（"是服务全省大局""构建政企学研…"等句子片段）
_VERB_LEADING_PATTERNS = [
    r'^是', r'^构建', r'^推动', r'^服务', r'^打造', r'^建设',
    r'^强化', r'^深化', r'^实现', r'^确保', r'^凝聚', r'^激发',
]


def _looks_like_sentence(text: str) -> bool:
    """M3 修复：长候选实体中包含多个动词/连词 → 判定为短句。"""
    return sum(1 for w in _SENTENCE_INDICATORS if w in text) >= 2


def _is_valid_entity_name(e_type: str, name: str, context: str = "") -> bool:
    """N2 + M3 + P1 修复：实体名有效性过滤（长度/动词片段/截断残片/右边界/长句/不完整后缀）。"""
    if not name:
        return False
    if e_type in ('person', 'org'):
        if not (2 <= len(name) <= 20):
            return False
    if e_type == 'org' and len(name) < _ORG_MIN_LEN:
        return False
    # P1：不完整后缀黑名单（"常务副部""副部"等残片）
    for bad in _INCOMPLETE_SUFFIX_BLACKLIST:
        if name.endswith(bad):
            return False
    # P1：重复尾字检测（"信息技术部部"= 右边界扩展重复）——机构名最后两字相同且同后缀
    if e_type == 'org' and len(name) >= 2:
        if name[-2] == name[-1]:
            return False
    # P1：动词/系词开头 → 句子片段，非实体
    for vpat in _VERB_LEADING_PATTERNS:
        if re.match(vpat, name):
            return False
    # M3：长度 > 10 的候选做句子检测（含 2+ 动词/连词 → 短句非实体）
    if len(name) > 10 and _looks_like_sentence(name):
        return False
    # M3 + P1：右边界检测——机构名以单字后缀结尾且原文紧跟扩展字（"部"+"门/长"）
    if e_type == 'org':
        for suffix, extensions in _ORG_SUFFIX_EXTENSIONS.items():
            if name.endswith(suffix):
                for ext in extensions:
                    if ext in context and context.find(name) >= 0:
                        after = context[context.find(name) + len(name):]
                        if after.startswith(ext):
                            return False  # 右边界截断（原文有"门/长"但未纳入实体）
    # 动词/助词开头 → 非实体（如"限公司""表分别"）
    for pat in _NON_ENTITY_PATTERNS:
        if re.match(pat, name):
            return False
    # 截断残片：以"限公司""表"等结尾但缺少前半部分
    if re.match(r'^(限公司|表|钟扬关|分别|副部|各部)', name):
        return False
    return True


def extract_entities(paragraphs: list[str]) -> List[Entity]:
    """
    Step 1：从文档段落提取关键实体。

    N2 修复：提取后经 _is_valid_entity_name 过滤，剔除动词片段/截断残片等非实体。

    Args:
        paragraphs: 文档段落文本列表

    Returns:
        实体列表
    """
    entities: List[Entity] = []
    seen = set()

    def _add(e_type: str, name: str, ctx: str, para_idx: int,
             doc_attribute: str = "", doc_context: str = "") -> None:
        if not _is_valid_entity_name(e_type, name):
            return
        key = (e_type, name)
        if key in seen:
            return
        seen.add(key)
        entities.append(Entity(entity_type=e_type, entity_name=name,
                               context=ctx[:60], paragraph_index=para_idx,
                               doc_attribute=doc_attribute, doc_context=doc_context))

    for idx, text in enumerate(paragraphs):
        if not text.strip():
            continue
        # 1. 人名+职务：{职务}XXX（职务在名前）——V2：保存完整职务到 doc_attribute/doc_context
        for m in re.finditer(
                r'([\u4e00-\u9fa5、，]{2,12}(?:' + '|'.join(_TITLE_SUFFIXES) + r'))([\u4e00-\u9fa5]{2,3})',
                text):
            _add('person', m.group(2), m.group(0), idx,
                 doc_attribute=m.group(1), doc_context=m.group(0))
        # 2. 人名（常见三字姓名模式：职务后或"听取了……的通报"结构）
        for m in re.finditer(r'听取了([\u4e00-\u9fa5]{2,3})(?:关于|对)', text):
            _add('person', m.group(1), m.group(0), idx)
        for m in re.finditer(r'([\u4e00-\u9fa5]{2,3})(?:同志)?(?:作了|进行了|主持)', text):
            _add('person', m.group(1), m.group(0), idx)
        # 3. 机构全称
        for m in re.finditer(
                r'([\u4e00-\u9fa5]{4,20}(?:' + '|'.join(_ORG_SUFFIXES) + r'))', text):
            _add('org', m.group(1), m.group(0), idx)
        # 4. 发文字号（XX发〔2026〕X号）
        for m in re.finditer(r'[\u4e00-\u9fa5]{2,6}发〔\d{4}〕\d+号', text):
            _add('doc_no', m.group(0), m.group(0), idx)
        # 5. 关键数据（金额/百分比）
        for m in re.finditer(r'(\d+(?:\.\d+)?(?:万元|亿元|%|％|万人|人次))', text):
            _add('data', m.group(1), m.group(0), idx)
        # 6. 项目名称（"关于……课题"）
        for m in re.finditer(r'关于([\u4e00-\u9fa5]{4,20}课题)', text):
            _add('project', m.group(1), m.group(0), idx)

    return entities


# ---------------------------------------------------------------------------
#  R3 修复：LLM 内容理解实体提取（主通道）+ 规则提取（兜底）交叉验证
# ---------------------------------------------------------------------------

_LLM_EXTRACT_PROMPT = (
    "请从以下中文公文中提取所有需事实核验的实体，包括：\n"
    "1. 人名（含其职务描述，如\"省民宗委党组成员、副主任XXX\"）\n"
    "2. 组织机构名（含完整全称）\n"
    "3. 发文字号、关键数据\n"
    "输出 JSON 数组，格式：[{{\"type\": \"person|org|doc_no|data\", \"name\": \"实体名\", "
    "\"title\": \"职务描述(仅person)\", \"paragraph_index\": 段落序号}}]。\n"
    "只输出 JSON，不要其他文字。\n\n"
    "公文内容：\n{content}"
)


def extract_entities_llm(paragraphs: list[str]) -> Optional[List[Entity]]:
    """R3 修复：LLM 内容理解提取实体（主通道，准确率优先）。

    通过环境变量 GONGWEN_LLM_API 指定 OpenAI 兼容 API（如本地 Ollama）启用；
    未配置时返回 None（调用方回退到规则提取，保持自包含、克隆即用）。

    Args:
        paragraphs: 文档段落文本列表

    Returns:
        LLM 提取的实体列表（成功时）；None（未配置/失败时回退规则）
    """
    import json as _json
    import os

    api_url = os.environ.get("GONGWEN_LLM_API", "").strip()
    api_key = os.environ.get("GONGWEN_LLM_API_KEY", "").strip()
    model = os.environ.get("GONGWEN_LLM_MODEL", "gpt-4o-mini").strip()
    if not api_url:
        logger.info("LLM 实体提取未配置（GONGWEN_LLM_API），回退规则提取")
        return None

    content = "\n".join(f"[{i}] {t}" for i, t in enumerate(paragraphs))
    payload = {
        "model": model,
        "messages": [
            {"role": "user", "content": _LLM_EXTRACT_PROMPT.format(content=content[:8000])},
        ],
        "temperature": 0,
    }
    try:
        import urllib.request
        req = urllib.request.Request(
            api_url.rstrip("/") + "/chat/completions",
            data=_json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json",
                     **({"Authorization": f"Bearer {api_key}"} if api_key else {})},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        raw = data["choices"][0]["message"]["content"]
        # 提取 JSON 数组（LLM 可能附带 ```json 围栏）
        m = re.search(r'\[\s*\{.*\}\s*\]', raw, re.S)
        if not m:
            logger.warning("LLM 返回无 JSON 数组，回退规则提取")
            return None
        items = _json.loads(m.group(0))
        entities: List[Entity] = []
        for it in items:
            etype = it.get("type", "")
            name = str(it.get("name", "")).strip()
            if not name or etype not in ("person", "org", "doc_no", "data"):
                continue
            pi = int(it.get("paragraph_index", 0) or 0)
            ctx = paragraphs[pi][:60] if 0 <= pi < len(paragraphs) else ""
            # V2：LLM 的 title 字段 → doc_attribute（职务/属性描述）
            entities.append(Entity(
                entity_type=etype, entity_name=name, context=ctx, paragraph_index=pi,
                doc_attribute=str(it.get("title", "") or ""),
                doc_context=ctx,
            ))
        logger.info(f"LLM 实体提取完成: {len(entities)} 项")
        return entities
    except Exception as e:
        logger.warning(f"LLM 实体提取失败（{e}），回退规则提取")
        return None


def extract_entities_hybrid(paragraphs: list[str]) -> List[Entity]:
    """R3 修复：混合提取——LLM 主通道 + 规则兜底 + 合并去重（LLM 结果优先）。

    Args:
        paragraphs: 文档段落文本列表

    Returns:
        合并去重后的实体列表
    """
    llm_entities = extract_entities_llm(paragraphs) or []
    rule_entities = extract_entities(paragraphs)

    if not llm_entities:
        return rule_entities

    # 合并去重：LLM 结果优先，规则结果补充（按 类型+名称 去重）
    merged: List[Entity] = []
    seen: set = set()
    for e in llm_entities + rule_entities:
        key = (e.entity_type, e.entity_name)
        if key in seen:
            continue
        seen.add(key)
        merged.append(e)
    return merged


# ---------------------------------------------------------------------------
#  基准构建（背景资料解析）
# ---------------------------------------------------------------------------

# SEC-2/SEC-3 修复：安全 URL 抓取辅助——拒绝内网/回环地址（SSRF 防护），
# 并限制重定向跳数（防重定向循环 / 跳转到内网后探测）
_PRIVATE_IP_PREFIXES = (
    "127.", "10.", "192.168.",
    "172.16.", "172.17.", "172.18.", "172.19.", "172.2", "172.30.", "172.31.",
    "0.", "169.254.",
)


def _safe_fetch_url(url: str, timeout: int = 10) -> Optional[str]:
    """安全抓取 URL 文本内容（SSRF 防护 + 重定向限制）。

    Returns:
        响应文本（utf-8 解码），失败/不安全返回 None
    """
    import ipaddress
    import socket
    import urllib.request

    # 1. scheme 校验：仅 http/https
    if not str(url).lower().startswith(("http://", "https://")):
        logger.warning(f"拒绝非 http(s) URL: {url[:60]}")
        return None

    # 2. 主机解析 + 内网/回环地址拒绝
    try:
        from urllib.parse import urlparse
        host = urlparse(url).hostname or ""
        # 处理 localhost 字面量
        if host in ("localhost", "localhost.localdomain", ""):
            logger.warning(f"拒绝内网/回环主机: {url[:60]}")
            return None
        # IP 字面量检查
        try:
            ip = ipaddress.ip_address(host)
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                logger.warning(f"拒绝内网/保留地址: {url[:60]}")
                return None
        except ValueError:
            # 域名：解析后校验（解析失败按不安全处理）
            try:
                for _info in socket.getaddrinfo(host, None):
                    _ip = _info[4][0]
                    if _ip.startswith(_PRIVATE_IP_PREFIXES) or _ip in ("::1", "0.0.0.0"):
                        logger.warning(f"拒绝内网解析地址: {url[:60]}")
                        return None
            except Exception:
                logger.warning(f"域名解析失败，视为不安全: {url[:60]}")
                return None
    except Exception as e:
        logger.warning(f"URL 主机校验失败: {url[:60]}: {e}")
        return None

    # 3. 限制重定向跳数（默认 urllib 最多 10 跳，收紧为 3 跳）
    class _LimitedRedirect(urllib.request.HTTPRedirectHandler):
        max_redirections = 3

    opener = urllib.request.build_opener(_LimitedRedirect)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with opener.open(req, timeout=timeout) as resp:
            return resp.read().decode("utf-8", errors="ignore")
    except Exception as e:
        logger.debug(f"安全 URL 抓取失败: {url[:60]}: {e}")
        return None


def _extract_text_from_background(path: str) -> str:
    """从背景资料提取纯文本（支持 docx/pdf/md/txt/url）。"""
    p = str(path).strip().lower()
    if p.startswith(('http://', 'https://')):
        # URL：尽力抓取正文（网络不可用时返回空；SEC-2 防 SSRF、SEC-3 限重定向）
        try:
            html = _safe_fetch_url(path, timeout=10)
            if html is None:
                return ""
            # 简单去 HTML 标签
            import re as _re
            return _re.sub(r'<[^>]+>', ' ', html)
        except Exception as e:
            logger.debug(f"URL 抓取失败: {path}: {e}")
            return ""
    path_obj = Path(path)
    if not path_obj.exists():
        logger.warning(f"背景资料不存在: {path}")
        return ""
    if path_obj.suffix.lower() == '.docx':
        try:
            from docx import Document
            doc = Document(str(path_obj))
            return '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.warning(f"docx 读取失败 {path}: {e}")
            return ""
    if path_obj.suffix.lower() == '.pdf':
        try:
            import fitz  # pymupdf
            doc = fitz.open(str(path_obj))
            return '\n'.join(page.get_text() for page in doc)
        except Exception as e:
            logger.warning(f"pdf 读取失败 {path}: {e}")
            return ""
    # md / txt / 其他文本
    try:
        for enc in ('utf-8', 'gbk'):
            try:
                return path_obj.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
    except Exception as e:
        logger.warning(f"文本读取失败 {path}: {e}")
    return ""


def build_baseline(background_paths: list[str]) -> dict[str, dict]:
    """
    Step 2：从背景资料构建事实基准表。

    Returns:
        {实体名 → {source, context, confidence}}
    """
    baseline: dict[str, dict] = {}
    for bp in background_paths or []:
        text = _extract_text_from_background(bp)
        if not text.strip():
            continue
        # P1-11 修复：按行切分为段落列表再提取实体——整篇背景作为单个"段落"会导致
        # 跨段落上下文混淆、匹配效率低（长文档事实核验性能差）
        para_list = [ln.strip() for ln in text.splitlines() if ln.strip()] or [text]
        entities = extract_entities(para_list)
        for e in entities:
            prev = baseline.get(e.entity_name)
            if prev is None:
                baseline[e.entity_name] = {
                    'source': str(bp), 'context': e.context,
                    'confidence': 'high' if str(bp).lower().endswith(('.gov.cn',)) else 'medium',
                }
    return baseline


# ---------------------------------------------------------------------------
#  互联网交叉核验（可选，web 可用时）
# ---------------------------------------------------------------------------

def _web_verify(entity_name: str, entity_type: str) -> Optional[str]:
    """
    Step 3：对关键实体做互联网交叉核验（尽力而为，网络不可用返回 None）。

    N4 修复：对"姓名+职务"组合一并搜索，比对文档中的职务描述与检索结果。
    P2 修复：环境变量控制（GONGWEN_WEB_VERIFY=1 开启/0 关闭）+ 多搜索引擎降级
    （百度→必应）+ 失败日志升级为 warning。

    Returns:
        核验说明（找到官方来源）或 None（无法核验）
    """
    import os
    import urllib.parse  # B32 修复：补 import（L457 使用 urllib.parse.quote）
    import urllib.request  # B32 修复：补 import（L461-462 使用 urllib.request）
    # P2：环境变量控制——默认关闭（受限网络环境爬取易失败），显式开启才执行
    if os.environ.get("GONGWEN_WEB_VERIFY", "0") != "1":
        logger.info("互联网核验未启用（设置 GONGWEN_WEB_VERIFY=1 可开启）")
        return None

    engines = [
        ("baidu", "https://www.baidu.com/s?wd={q}"),
        ("bing", "https://www.bing.com/search?q={q}"),
    ]
    query = urllib.parse.quote(f"{entity_name} {'职务' if entity_type == 'person' else ''}")
    for engine_name, url_template in engines:
        try:
            url = url_template.format(q=query)
            # SEC-2/SEC-3 修复：改用安全抓取（SSRF 防护 + 重定向限制）
            html = _safe_fetch_url(url, timeout=8)
            if html is None:
                continue
            if entity_name in html:
                return f"互联网检索到 {entity_name} 相关信息（来源：{engine_name}，请人工核对权威性）"
        except Exception as e:
            logger.warning(f"web 核验失败（{engine_name}） {entity_name}: {e}")
    return "互联网检索未直接命中"


# N4 修复：人名+职务配对核验——识别"职务+姓名"组合（如"副主任XXX"）
# 返回 [(姓名, 职务描述), ...]
def extract_person_title_pairs(text: str) -> List[tuple[str, str]]:
    """从段落文本中提取 (姓名, 职务描述) 配对。

    模式：{职务}（2-8 字，含 主任/副主任/书记/部长 等后缀）+ {姓名}（2-3 字）。
    如"省民宗委党组成员、副主任XXX" → ("XXX", "省民宗委党组成员、副主任")。
    """
    pairs: List[tuple[str, str]] = []
    title_pattern = r'([\u4e00-\u9fa5、]{2,20}(?:' + '|'.join(_TITLE_SUFFIXES) + r'))([\u4e00-\u9fa5]{2,3})(?![、，。；])'
    for m in re.finditer(title_pattern, text):
        title = m.group(1)
        name = m.group(2)
        # 过滤常见非姓名词
        if name in ("会议", "工作", "部门", "单位", "同志", "有关", "相关"):
            continue
        pairs.append((name, title))
    return pairs


# ---------------------------------------------------------------------------
#  主流程
# ---------------------------------------------------------------------------

def run_fact_check(document_path: str | Path, background_paths: Optional[list[str]] = None,
                   web_enabled: bool = True) -> FactCheckReport:
    """
    Step 1-4：完整事实核验流程。

    Args:
        document_path: 用户文档 .docx
        background_paths: 背景资料列表（docx/pdf/md/txt/url）
        web_enabled: 是否启用互联网核验（默认开启，失败自动降级）

    Returns:
        FactCheckReport 报告
    """
    doc_path = Path(document_path)
    report = FactCheckReport(document=doc_path.name)

    # Step 1：实体提取（R3 修复：混合提取——LLM 内容理解主通道 + 规则兜底）
    from engine.core.document.parser import parse_docx
    model = parse_docx(str(doc_path))
    paragraphs = [p.text for p in model.paragraphs]
    entities = extract_entities_hybrid(paragraphs)
    if not entities:
        logger.info("未提取到关键实体")
        return report

    # N2 修复：按实体文本在文档段落中搜索定位，校正 paragraph_index（提取阶段索引可能偏差）
    for e in entities:
        if not e.entity_name:
            continue
        found_idx = None
        for idx, ptext in enumerate(paragraphs):
            if e.entity_name in ptext:
                found_idx = idx
                break
        if found_idx is not None and found_idx != e.paragraph_index:
            logger.debug(f"实体 {e.entity_name}: 段落索引 {e.paragraph_index} → 校正为 {found_idx}")
            e.paragraph_index = found_idx

    report.entities = entities

    # Step 2：基准构建
    baseline = build_baseline(background_paths or [])

    # Step 3+4：核验与标记
    # N4 修复：先做人名+职务配对识别，对"姓名+职务"组合整体核验（能发现职务写反等严重错误）
    person_title_pairs: dict[str, str] = {}
    for ptext in paragraphs:
        for name, title in extract_person_title_pairs(ptext):
            person_title_pairs.setdefault(name, title)

    for e in entities:
        if e.entity_name in baseline:
            src = baseline[e.entity_name]['source']
            e.status = '已确认'
            e.source = src
            e.note = f"据背景资料确认（{Path(src).name}）"
            report.confirmed.append(e)
        else:
            # 互联网交叉核验（仅关键项：person / org）
            note = None
            if web_enabled and e.entity_type in ('person', 'org'):
                # N4：人名+职务配对核验——搜索"姓名+职务"组合，提示人工比对
                if e.entity_type == 'person' and e.entity_name in person_title_pairs:
                    doc_title = person_title_pairs[e.entity_name]
                    note = _web_verify(f"{e.entity_name} {doc_title}", 'person')
                    if note:
                        note += f"｜文档中职务描述：{doc_title}（请人工核对该职务是否准确）"
                else:
                    note = _web_verify(e.entity_name, e.entity_type)
            if note and '相关信息' in note:
                e.status = '存疑'
                e.source = '互联网'
                e.note = note
                report.doubtful.append(e)
            else:
                e.status = '无法核验'
                e.source = ''
                e.note = '未经核验，请人工确认'
                report.unverified.append(e)

    # 写报告文件
    try:
        out = doc_path.with_name(doc_path.stem + '.fact_check.json')
        out.write_text(json.dumps(report.to_dict(), ensure_ascii=False, indent=2), encoding='utf-8')
        report._report_path = str(out)  # type: ignore[attr-defined]
    except Exception as e:
        logger.debug(f"核验报告写入失败: {e}")

    return report
