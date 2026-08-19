# -*- coding: utf-8 -*-
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# https://github.com/linhut/gongwen-skill
# Licensed under the MIT License. See the LICENSE file for details.
#
"""
Generator helper functions: table and page number field creation.
Extracted from generator.py (tier-2 split).
"""
from __future__ import annotations
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from engine.core.document.models import Paragraph, Run, Table as TableModel
except ImportError:
    from .models import Paragraph, Run, Table as TableModel

from lxml import etree
from docx.shared import Pt, RGBColor
from engine.core.document.font_utils import (
    set_run_font, BODY_FONT,
    PAGE_NUMBER_FONT, PAGE_NUMBER_LATIN_FONT, PAGE_NUMBER_SIZE_PT,
)
from engine.utils.logger import logger as _gh_logger

logger = _gh_logger


def _smart_align_cell(cell_text: str, is_header: bool, col_idx: int, total_cols: int) -> str:
    """智能判断单元格对齐方式。

    规则：
    - 表头：居中
    - 序号列（第一列）：居中
    - 短文本（≤4字符）：居中
    - 数字内容（含小数、百分比）：右对齐
    - 其他：左对齐
    """
    text = cell_text.strip()
    if not text:
        return 'left'
    if is_header:
        return 'center'
    # 序号列居中（优先级最高，避免被数字规则覆盖）
    if col_idx == 0:
        return 'center'
    # 短文本居中
    if len(text) <= 4:
        return 'center'
    # 数字右对齐（含小数、百分比、逗号分隔数字）
    import re
    if re.match(r'^[\d.,%‰]+$', text):
        return 'right'
    return 'left'


def _update_table_content(table, table_model: TableModel):
    """更新已有表格的单元格内容（带智能对齐）。"""
    total_cols = len(table.columns) if hasattr(table, 'columns') else 0
    for cell_model in table_model.cells:
        try:
            cell = table.cell(cell_model.row, cell_model.col)
            # 更新单元格中的段落内容
            if cell_model.paragraphs:
                for p_idx, para_model in enumerate(cell_model.paragraphs):
                    if p_idx < len(cell.paragraphs):
                        # 替换已有段落
                        para = cell.paragraphs[p_idx]
                        # 清除旧 runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                        # 添加新 runs
                        _add_runs_to_paragraph(para, para_model)
                        # 更新段落格式（缩进、行距、对齐等）
                        _update_pPr(para._element, para_model)
                    else:
                        # 添加新段落
                        para = cell.add_paragraph()
                        _add_runs_to_paragraph(para, para_model)
                        _apply_paragraph_format(para, para_model)
            elif cell_model.text:
                # 没有详细段落信息，直接设置文本
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    run = para.add_run(cell_model.text)
                    set_run_font(run, BODY_FONT)
                    # 智能对齐
                    is_header = cell_model.row == 0
                    align = _smart_align_cell(cell_model.text, is_header, cell_model.col, total_cols)
                    para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                                      'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        except Exception as e:
            logger.warning(f"Failed to update table cell ({cell_model.row},{cell_model.col}): {e}")


def _add_table(doc: Document, table_model: TableModel):
    """在文档中添加一个新表格，按 insert_after_index 定位到正确位置。"""
    try:
        rows = max(1, table_model.rows)
        cols = max(1, table_model.cols)
        table = doc.add_table(rows=rows, cols=cols)

        # 按 insert_after_index 移动表格到正确位置
        insert_idx = getattr(table_model, 'insert_after_index', -1)
        if insert_idx >= 0:
            body = doc.element.body
            para_count = 0
            target_elem = None
            for child in body:
                tag = etree.QName(child.tag).localname if child.tag else ''
                if tag == 'p':
                    if para_count == insert_idx:
                        target_elem = child
                        break
                    para_count += 1
            if target_elem is not None:
                tbl_elem = table._tbl
                body.remove(tbl_elem)
                target_elem.addnext(tbl_elem)

        # 设置表格边框（虚线 dashed——公文表格规范常用虚线表线，直接格式覆盖样式）
        # 无论 'Table Grid' 样式是否存在都显式写入 tblBorders，保证虚线生效
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        # 移除可能已存在的 tblBorders（避免残留实线定义）
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)
        borders = tblPr.makeelement(qn('w:tblBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'dashed',   # 虚线表线
                qn('w:sz'): '4',
                qn('w:space'): '0',
                qn('w:color'): '000000',
            })
            borders.append(border)
        tblPr.append(borders)

        # 智能对齐：表头行居中加粗，数据行按内容类型对齐
        total_cols = max(1, table_model.cols)
        for cell_model in table_model.cells:
            try:
                cell = table.cell(cell_model.row, cell_model.col)
                # 清除默认段落
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)

                if cell_model.paragraphs:
                    for p_idx, para_model in enumerate(cell_model.paragraphs):
                        if p_idx < len(cell.paragraphs):
                            para = cell.paragraphs[p_idx]
                            _add_runs_to_paragraph(para, para_model)
                            _update_pPr(para._element, para_model)
                        else:
                            para = cell.add_paragraph()
                            _add_runs_to_paragraph(para, para_model)
                            _apply_paragraph_format(para, para_model)
                    # 智能对齐（表头居中加粗，数据行按内容类型）
                    is_header = cell_model.row == 0
                    cell_text = cell_model.text or (cell_model.paragraphs[0].text if cell_model.paragraphs else '')
                    align = _smart_align_cell(cell_text, is_header, cell_model.col, total_cols)
                    align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                                 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
                    for para in cell.paragraphs:
                        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                    if is_header:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                elif cell_model.text:
                    if cell.paragraphs:
                        run = cell.paragraphs[0].add_run(cell_model.text)
                        set_run_font(run, BODY_FONT)
                        # 智能对齐
                        is_header = cell_model.row == 0
                        align = _smart_align_cell(cell_model.text, is_header, cell_model.col, total_cols)
                        align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                                     'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
                        cell.paragraphs[0].alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                        if is_header:
                            run.bold = True
            except Exception as e:
                logger.warning(f"Failed to write table cell ({cell_model.row},{cell_model.col}): {e}")

        logger.debug(f"Added table: {rows}x{cols}")
    except Exception:
        logger.exception("Failed to add table")


def _add_page_number_field(para, para_model: Paragraph) -> None:
    """
    在段落中写入 Word 页码域代码（{ PAGE } / { NUMPAGES }）。
    使用 Word XML 域代码实现动态页码，而非静态文本。

    支持格式：
    - "{PAGE}" → 当前页码
    - "{NUMPAGES}" → 总页数
    - 可包含前缀后缀，如 "- {PAGE} -" 或 "第 {PAGE} 页 共 {NUMPAGES} 页"
    """
    run_text = para_model.text if para_model.text else "{PAGE}"

    # 构建标准 Word 页码域 XML
    # 格式: 文本 + PAGE域 + 文本 + NUMPAGES域 + 文本
    parts = []
    import re
    # 拆分文本中的 {PAGE} 和 {NUMPAGES} 占位符
    remaining = run_text
    while remaining:
        m = re.search(r'\{PAGE\}|\{NUMPAGES\}', remaining)
        if not m:
            if remaining.strip():
                parts.append(('text', remaining))
            break

        # 前置文本
        prefix = remaining[:m.start()]
        if prefix.strip():
            parts.append(('text', prefix))

        # 域代码
        parts.append(('field', m.group()))
        remaining = remaining[m.end():]

    # 写入 Word XML
    for part_type, content in parts:
        if part_type == 'text':
            run_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            # P2-5 修复：页码字体按 GB/T 9704 规范（宋体 14pt），改用 font_utils 常量
            # 而非散落的字符串字面量，保证与其它页码注入路径一致
            rFonts.set(qn('w:eastAsia'), PAGE_NUMBER_FONT)
            rFonts.set(qn('w:ascii'), PAGE_NUMBER_LATIN_FONT)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(PAGE_NUMBER_SIZE_PT * 2)))  # 14pt
            rPr.append(sz)
            run_el.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = content
            run_el.append(t)
            para._element.append(run_el)
        elif part_type == 'field':
            # 创建 fldChar begin
            fld_begin = OxmlElement('w:r')
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            fld_begin.append(fldChar_begin)
            para._element.append(fld_begin)

            # 创建 instrText (域代码指令)
            instr = OxmlElement('w:r')
            rPr_instr = OxmlElement('w:rPr')
            rFonts_instr = OxmlElement('w:rFonts')
            rFonts_instr.set(qn('w:eastAsia'), PAGE_NUMBER_FONT)
            rPr_instr.append(rFonts_instr)
            instr.append(rPr_instr)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            field_name = content[1:-1]  # Remove { } → PAGE or NUMPAGES
            instrText.text = f' {field_name} '
            instr.append(instrText)
            para._element.append(instr)

            # 创建 fldChar separate
            fld_sep = OxmlElement('w:r')
            fldChar_sep = OxmlElement('w:fldChar')
            fldChar_sep.set(qn('w:fldCharType'), 'separate')
            fld_sep.append(fldChar_sep)
            para._element.append(fld_sep)

            # 创建默认显示值
            fld_default = OxmlElement('w:r')
            rPr_def = OxmlElement('w:rPr')
            rFonts_def = OxmlElement('w:rFonts')
            # P1-10 修复：页码默认显示值字体改用 PAGE_NUMBER_FONT 常量，不再硬编码 '宋体'
            rFonts_def.set(qn('w:eastAsia'), PAGE_NUMBER_FONT)
            rFonts_def.set(qn('w:ascii'), PAGE_NUMBER_LATIN_FONT)
            rPr_def.append(rFonts_def)
            sz_def = OxmlElement('w:sz')
            sz_def.set(qn('w:val'), '28')
            rPr_def.append(sz_def)
            fld_default.append(rPr_def)
            t_def = OxmlElement('w:t')
            t_def.set(qn('xml:space'), 'preserve')
            t_def.text = '1'  # 默认显示值
            fld_default.append(t_def)
            para._element.append(fld_default)

            # 创建 fldChar end
            fld_end = OxmlElement('w:r')
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            fld_end.append(fldChar_end)
            para._element.append(fld_end)


def _update_pPr(p_element, para_model: Paragraph):
    """更新 <w:p> 元素的 <w:pPr> 段落属性。
    关键原则：model 有值才替换，None 保留原文档格式不删除。"""
    fmt = para_model.format

    # 获取或创建 pPr
    pPr = p_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_element.insert(0, pPr)

    # 对齐方式：仅当 model 有值时替换
    if fmt.alignment:
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        alignment_map = {
            "left": "left", "center": "center",
            "right": "right", "justify": "both",
        }
        jc.set(qn('w:val'), alignment_map.get(fmt.alignment, "left"))
        pPr.append(jc)

    # 缩进：仅当 model 有值时替换，否则保留原文档缩进
    has_indent = (fmt.first_line_indent_pt is not None or
                  fmt.left_indent_pt is not None or
                  fmt.right_indent_pt is not None)
    if has_indent:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        ind = OxmlElement('w:ind')
        if fmt.first_line_indent_pt is not None:
            ind.set(qn('w:firstLine'), str(int(fmt.first_line_indent_pt * 20)))
            _font_size = 16.0
            if para_model.runs and para_model.runs[0].format.font_size_pt:
                _font_size = para_model.runs[0].format.font_size_pt
            chars = int(round(fmt.first_line_indent_pt / _font_size * 100))
            if chars > 0:
                ind.set(qn('w:firstLineChars'), str(chars))
        if fmt.left_indent_pt is not None:
            ind.set(qn('w:left'), str(int(fmt.left_indent_pt * 20)))
        if fmt.right_indent_pt is not None:
            ind.set(qn('w:right'), str(int(fmt.right_indent_pt * 20)))
        pPr.append(ind)

    # 行距：仅当 model 有值时替换，否则保留原文档行距
    has_spacing = (fmt.line_spacing_pt is not None or
                   fmt.space_before_pt is not None or
                   fmt.space_after_pt is not None)
    if has_spacing:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        if fmt.line_spacing_pt is not None:
            spacing_pt = max(6, min(200, fmt.line_spacing_pt))
            rule = fmt.line_spacing_rule or "exact"
            if rule == "multiple":
                line_val = int(round(spacing_pt / 16 * 240))
                spacing.set(qn('w:line'), str(line_val))
                spacing.set(qn('w:lineRule'), 'auto')
            elif rule == "atLeast":
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'atLeast')
            else:
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'exact')
        if fmt.space_before_pt is not None:
            spacing.set(qn('w:before'), str(int(fmt.space_before_pt * 20)))
        if fmt.space_after_pt is not None:
            spacing.set(qn('w:after'), str(int(fmt.space_after_pt * 20)))
        pPr.append(spacing)


def _add_runs_to_paragraph(para, para_model: Paragraph):
    """使用 python-docx API 向段落添加 runs。"""
    if para_model.runs:
        for run_model in para_model.runs:
            run = para.add_run(run_model.text)
            _apply_run_format(run, run_model)
    else:
        if para_model.text:
            run = para.add_run(para_model.text)
            # 优先使用段落 format 中的 font_name，fallback 到 BODY_FONT
            fmt_font = None
            if para_model.runs and para_model.runs[0].format:
                fmt_font = para_model.runs[0].format.font_name
            if not fmt_font and para_model.format:
                fmt_font = getattr(para_model.format, 'font_name', None)
            set_run_font(run, fmt_font or BODY_FONT)


def _apply_paragraph_format(para, para_model: Paragraph):
    """Apply formatting to a paragraph using python-docx API."""
    pf = para.paragraph_format
    fmt = para_model.format

    # Alignment
    if fmt.alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = alignment_map.get(fmt.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Indentation
    if fmt.first_line_indent_pt is not None:
        pf.first_line_indent = Pt(fmt.first_line_indent_pt)
    if fmt.left_indent_pt is not None:
        pf.left_indent = Pt(fmt.left_indent_pt)
    if fmt.right_indent_pt is not None:
        pf.right_indent = Pt(fmt.right_indent_pt)

    # Spacing
    if fmt.space_before_pt is not None:
        pf.space_before = Pt(fmt.space_before_pt)
    if fmt.space_after_pt is not None:
        pf.space_after = Pt(fmt.space_after_pt)

    # Line spacing
    if fmt.line_spacing_pt is not None:
        spacing_pt = max(6, min(200, fmt.line_spacing_pt))
        rule = fmt.line_spacing_rule or "exact"
        if rule == "multiple":
            multiple = spacing_pt / 16.0
            pf.line_spacing = multiple
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        elif rule == "atLeast":
            pf.line_spacing = Pt(spacing_pt)
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            pf.line_spacing = Pt(spacing_pt)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _apply_run_format(run, run_model: Run):
    """
    Apply formatting to a run.
    使用 font_utils.set_run_font 统一处理中文字体。
    """
    fmt = run_model.format

    # === 字体设置（统一入口） ===
    if fmt.font_name:
        set_run_font(run, fmt.font_name)
    else:
        set_run_font(run, BODY_FONT)

    # === 字号 ===
    if fmt.font_size_pt is not None:
        run.font.size = Pt(fmt.font_size_pt)

    # === 样式 ===
    if fmt.bold is not None:
        run.font.bold = fmt.bold
    if fmt.italic is not None:
        run.font.italic = fmt.italic
    if fmt.underline is not None:
        run.font.underline = fmt.underline
    if fmt.strikethrough is True:
        run.font.strike = True

    # === 颜色 ===
    if fmt.color:
        try:
            rgb_str = fmt.color.replace("#", "")
            if len(rgb_str) == 6:
                r = int(rgb_str[0:2], 16)
                g = int(rgb_str[2:4], 16)
                b = int(rgb_str[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            logger.warning(f"颜色值解析失败: {e}")
