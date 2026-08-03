# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Document generator: converts DocumentModel back into a .docx file.

核心设计变更 (v1.7.1):
- 基于源文档原地修改，不再创建空 Document()
- 正确保留表格、图片、嵌入对象、分节符等 DocumentModel 未建模的内容
- 所有字体设置必须经过 font_utils 统一入口
- 支持表格写入、页眉/页脚写入
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from core.document.models import DocumentModel, Paragraph, Run, Table as TableModel, HeaderFooter
from core.document.font_utils import (
    set_run_font, set_paragraph_font, validate_document_fonts,
    TITLE_FONT, BODY_FONT, LATIN_FONT,
    PAGE_NUMBER_FONT, PAGE_NUMBER_LATIN_FONT, PAGE_NUMBER_SIZE_PT,
    _LATIN_FONTS, _contains_cjk,
)
from utils.logger import logger


def _accept_all_revisions(body_elem) -> int:
    """接受文档中的所有修订（tracked changes），产出干净文档（P3）。

    1. 删除所有 <w:del> 元素及其内容（被删除的文本直接移除）
    2. 将 <w:ins> 子元素提升为正式内容（解包，其内 <w:r> 上移为父级直接子元素）
    3. 清理 <w:rPrChange>/<w:pPrChange> 等修订属性元素

    Args:
        body_elem: 文档 body 元素（doc.element.body）

    Returns:
        处理的修订元素数量
    """
    if body_elem is None:
        return 0
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    count = 0

    # 1. 删除 <w:del>（含 <w:delText>，即已删除内容，接受修订=最终删除）
    for del_elem in body_elem.findall(f'.//{ns}del'):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)
            count += 1

    # 2. 解包 <w:ins>：将其子元素按原顺序提升到父级（接受修订=保留插入内容）
    for ins_elem in body_elem.findall(f'.//{ns}ins'):
        parent = ins_elem.getparent()
        if parent is not None:
            idx = list(parent).index(ins_elem)
            for child in list(ins_elem):
                ins_elem.remove(child)
                parent.insert(idx, child)
                idx += 1
            parent.remove(ins_elem)
            count += 1

    # 3. 清理修订属性元素（属性级修订，接受=保留新值删除旧记录）
    for change_elem in (body_elem.findall(f'.//{ns}rPrChange')
                        + body_elem.findall(f'.//{ns}pPrChange')):
        parent = change_elem.getparent()
        if parent is not None:
            parent.remove(change_elem)
            count += 1

    if count:
        logger.info(f"_accept_all_revisions: cleaned {count} tracked-change elements")
    return count


def generate_docx(model: DocumentModel, output_path: Path | str | "io.BytesIO",
                  no_ai_declaration: bool = False) -> Path | "io.BytesIO":
    """
    Generate a .docx file from a DocumentModel.

    策略：加载源文档 → 替换段落内容（保留表格在原文位置）→ 更新表格 → 更新页眉页脚 → 保存。
    这样可以保留 DocumentModel 未建模的内容（图片、嵌入对象、分节符等）。

    图片处理说明：
    - 内联图片（<w:drawing>）：随所在 <w:p> 段落一起保留。
    - 浮动图片（锚定）：不受段落替换影响。
    - 前提：源文档保留策略生效（model.source_path 指向原 .docx）。
    - 若无源文档（如 md2docx 创建新文档），图片无法保留。

    Args:
        model: The document model to generate from
        output_path: Path where the .docx file should be saved，或文件对象（BytesIO，内存输出）
        no_ai_declaration: 为 True 时跳过 AI 声明段追加（P1，默认 False 保持原行为）

    Returns:
        Path to the generated file（或 BytesIO 对象）
    """
    import io
    is_fileobj = isinstance(output_path, io.BytesIO)

    if is_fileobj:
        out_obj = output_path
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        out_obj = str(output_path)

    logger.info(f"Generating document: {out_obj if is_fileobj else output_path}")

    # 决定源文档：优先使用源文件，若不可用则创建新文档
    source_path = model.source_path
    if source_path and Path(source_path).exists():
        try:
            doc = Document(str(source_path))
            logger.debug(f"Loaded source document: {source_path}")
        except Exception as e:
            # S8 修复：明确告知源文档加载失败将导致未建模内容（图片/分节符等）丢失
            logger.warning(
                f"源文档加载失败（{e}），将创建空文档——原文档中的图片/嵌入对象/"
                f"分节符等未建模内容将无法保留。如确认源文件损坏，请检查 {source_path}",
                exc_info=True,
            )
            doc = Document()
    else:
        doc = Document()
        logger.debug("No source document, created new Document()")

    # 0. Apply document-level font defaults (prevents Word from using MS Gothic)
    _apply_document_defaults(doc)

    # 构建段落索引缓存（_replace_paragraphs 中大量调用 _find_paragraph_object）
    _build_paragraph_index(doc)

    # 1. Apply page setup
    _apply_page_setup(doc, model)

    # 1.5 P0-2: 保存前先接受所有修订（tracked changes）——必须在段落替换之前执行，
    # 否则段落替换写入的新 run 可能被残余 <w:ins> 包裹，解包后新内容被意外提升/移除
    _accept_all_revisions(doc.element.body)

    # 2. Replace paragraphs in-place (preserving table positions)
    _replace_paragraphs(doc, model)

    # P2-1 修复：段落替换会增删段落，构建的段落索引缓存已失效，更新表格前清除重建
    _clear_paragraph_index(doc)

    # 3. Update tables with model data
    _update_tables(doc, model)

    # 4. Update headers and footers
    _update_headers_footers(doc, model)

    # 5. Update metadata
    _update_metadata(doc, model)

    # 6. Post-generation font validation & auto-fix
    font_issues = validate_document_fonts(doc)
    if font_issues:
        logger.warning(f"Found {len(font_issues)} font issues, auto-fixing...")
        _auto_fix_fonts(doc, font_issues)

    # 7. AI 声明（去重+添加，所有路径产出文档末尾统一；P1: no_ai_declaration=True 时跳过）
    # P3-5：提取为独立函数，保持 generate_docx 主流程精简
    _apply_ai_declaration(doc, no_ai_declaration)

    # 8. Save（支持路径或文件对象 BytesIO）
    doc.save(out_obj)
    # 清除段落索引缓存（文档已保存，doc 对象可能被 GC）
    _clear_paragraph_index(doc)
    logger.info(f"Document saved: {out_obj if is_fileobj else output_path} (font issues: {len(font_issues)})")

    return output_path if not is_fileobj else out_obj


# ---------------------------------------------------------------------------
#  AI Declaration（P3-5：从 generate_docx 提取，独立函数）
# ---------------------------------------------------------------------------

def _apply_ai_declaration(doc: Document, no_ai_declaration: bool) -> None:
    """在文档末尾追加/修正 AI 声明段（去重 + 统一格式）。

    Args:
        doc: 待处理的 python-docx Document
        no_ai_declaration: 为 True 时跳过（不追加也不修正）
    """
    if no_ai_declaration:
        return
    ai_variants = [
        "（内容由GongWen-skill-AI生成，仅供参考）",
        "（内容由AI生成，仅供参考）",
    ]
    ai_text = ai_variants[0]

    # 去重：在 doc 中移除多余声明段落（从后往前删避免索引偏移）
    ai_doc_indices = [i for i, p in enumerate(doc.paragraphs)
                      if any(v in (p.text or "") for v in ai_variants)]
    if len(ai_doc_indices) > 1:
        body = doc.element.body
        for idx in reversed(ai_doc_indices[:-1]):
            p_elem = doc.paragraphs[idx]._element
            body.remove(p_elem)

    # 检查是否已有声明；如有则修正其格式，无则添加
    from core.document.font_utils import set_run_font
    existing_ai_para = None
    for p in doc.paragraphs:
        if any(v in (p.text or "") for v in ai_variants):
            existing_ai_para = p
            break

    if existing_ai_para:
        # 修正已有声明的字体格式
        existing_ai_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 标记为 AI 注释段落，避免 check 误判为标题
        p_elem = existing_ai_para._element
        pPr = p_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is None:
            pPr = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        # 添加 pStyle 标记段落样式（不依赖 role，仅用于 check 时跳过）
        pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        if pStyle is None:
            pStyle = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Annotation')
        for r in existing_ai_para.runs:
            set_run_font(r, '楷体_GB2312')
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    else:
        ai_para = doc.add_paragraph()
        ai_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 通过 pStyle 标记为注释，避免 check 误判
        p_elem = ai_para._element
        pPr = p_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is None:
            pPr = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        pStyle = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Annotation')
        ai_run = ai_para.add_run(ai_text)
        set_run_font(ai_run, '楷体_GB2312')
        ai_run.font.size = Pt(9)
        ai_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
#  Document Defaults & Page Setup
# ---------------------------------------------------------------------------

def _auto_fix_fonts(doc: Document, font_issues: list[dict]):
    """自动替换检测到的无效字体（MS Gothic 等）为合规字体。"""
    from core.document.font_utils import FONT_FALLBACK_MAP, INVALID_FONT_PATTERNS

    def _get_replacement(attr: str, invalid_font: str) -> str:
        """根据属性类型和无效字体名，返回合规替换字体。"""
        # 先查 fallback map
        fallback = FONT_FALLBACK_MAP.get(invalid_font)
        if fallback:
            return fallback
        # 未命中则按属性类型返回默认值
        if attr == "eastAsia" or attr == "cs":
            return BODY_FONT
        return LATIN_FONT

    fixed_count = 0
    seen_runs = set()
    for issue in font_issues:
        run = issue.get("run_obj")
        if run is None or id(run) in seen_runs:
            continue
        # P2-6 修复：run 悬空引用防护——run 已被移出文档（无父元素）时跳过，
        # 避免对已删除的 run 设置字体引发异常或污染无关段落
        try:
            if run._element is None or run._element.getparent() is None:
                continue
        except Exception:
            continue
        seen_runs.add(id(run))
        try:
            # S9 修复：使用 FONT_FALLBACK_MAP 计算结果替换，而非一律 BODY_FONT
            invalid_font = issue.get("font_name", "") or issue.get("name", "")
            replacement = _get_replacement(issue.get("attr", "eastAsia"), invalid_font)
            set_run_font(run, replacement)
            fixed_count += 1
        except Exception as e:
            logger.debug(f"Auto-fix font failed: {e}")

    if fixed_count:
        logger.info(f"Auto-fixed {fixed_count} runs with invalid fonts")


def _apply_document_defaults(doc: Document):
    """
    设置文档级别的默认字体。
    通过在 styles.xml 中设置 docDefaults，确保即使某个 run 没有显式设置字体，
    Word 也不会使用 MS Gothic 等替代字体回退。
    """
    try:
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)

        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            doc_defaults.append(rPrDefault)

        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)

        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), BODY_FONT)
        rFonts.set(qn('w:cs'), BODY_FONT)

        logger.debug("Document default fonts applied: eastAsia=仿宋_GB2312, latin=Times New Roman")
    except Exception as e:
        logger.warning(f"Failed to set document default fonts: {e}")


def _apply_page_setup(doc: Document, model: DocumentModel):
    """Apply page setup to the document."""
    if not doc.sections:
        return
    section = doc.sections[0]
    ps = model.page_setup

    if ps.paper_width_mm is not None and 50 <= ps.paper_width_mm <= 1000:
        section.page_width = Mm(ps.paper_width_mm)
    if ps.paper_height_mm is not None and 50 <= ps.paper_height_mm <= 1000:
        section.page_height = Mm(ps.paper_height_mm)
    if ps.margin_top_mm is not None and 0 <= ps.margin_top_mm <= 100:
        section.top_margin = Mm(ps.margin_top_mm)
    if ps.margin_bottom_mm is not None and 0 <= ps.margin_bottom_mm <= 100:
        section.bottom_margin = Mm(ps.margin_bottom_mm)
    if ps.margin_left_mm is not None and 0 <= ps.margin_left_mm <= 100:
        section.left_margin = Mm(ps.margin_left_mm)
    if ps.margin_right_mm is not None and 0 <= ps.margin_right_mm <= 100:
        section.right_margin = Mm(ps.margin_right_mm)

    # 改动10：页眉/页脚距页边界从 page_setup 配置读取（省筹委会规范 footer 2.3cm / header 1.5cm）
    # fallback 保持 GB/T 9704 默认值（footer 2.5cm / header 1.5cm）
    footer_dist = getattr(ps, 'footer_distance_cm', None)
    header_dist = getattr(ps, 'header_distance_cm', None)
    section.footer_distance = Cm(footer_dist if footer_dist is not None else 2.5)
    section.header_distance = Cm(header_dist if header_dist is not None else 1.5)


# ---------------------------------------------------------------------------
#  Paragraph Replacement (in-place, preserving table positions)
# ---------------------------------------------------------------------------

def _replace_paragraphs(doc: Document, model: DocumentModel):
    """
    替换文档中的段落内容，同时保留表格和图片在原始位置。

    策略（P0-1 加固）：
    1. 交错遍历 body 直接子元素：<w:p> 按序消耗 model.paragraphs 内容，
       <w:tbl> 及表格内段落保持不动（表格不消耗 model 段落索引）
    2. model 比原文多的段落，追加到 body 末尾
    3. 原文比 model 多的 <w:p>，从 body 中移除
    4. 图片：内联图片（<w:drawing> 在 <w:p> 内）通过段落内容替换间接保留，
       浮动图片（锚定）不受影响。前提是源文档保留策略生效。

    注意：段落索引必须与模型中的 index 字段严格对齐。
    索引错位会导致内联图片跟随错误的段落移位。
    """
    body = doc.element.body
    p_tag = qn('w:p')
    model_paras = model.paragraphs
    para_idx = 0

    # 交错遍历：跳过 <w:tbl>（不消耗 model 索引），仅对 <w:p> 按序替换/移除
    for child in list(body):
        if child.tag == p_tag:
            if para_idx < len(model_paras):
                _replace_paragraph_content(doc, child, model_paras[para_idx])
                para_idx += 1
            else:
                # 原文段落多于 model → 移除多余段落
                try:
                    body.remove(child)
                except Exception:
                    pass  # 已被移除则跳过

    # model 比原文多的段落，追加到 body 末尾
    while para_idx < len(model_paras):
        new_para = doc.add_paragraph()
        _apply_paragraph_format(new_para, model_paras[para_idx])
        _add_runs_to_paragraph(new_para, model_paras[para_idx])
        para_idx += 1

    logger.debug(f"Replaced {min(len(model_paras), len([c for c in body if c.tag == p_tag]))} paragraphs, "
                 f"added {max(0, len(model_paras) - len([c for c in body if c.tag == p_tag]))}, "
                 f"removed {max(0, len([c for c in body if c.tag == p_tag]) - len(model_paras))}")


def _replace_paragraph_content(doc: Document, p_element, para_model: Paragraph):
    """
    替换一个 <w:p> 元素的内容（清除旧文本 runs，写入新 runs），保留段落属性和图片。
    """
    # 清除文本 runs，但保留含图片/绘图的 runs
    for child in list(p_element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            # 检查 run 是否包含图片（w:drawing 或 w:pict）
            has_image = False
            for sub in child:
                sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
                if sub_tag in ('drawing', 'pict'):
                    has_image = True
                    break
            if not has_image:
                p_element.remove(child)
        elif tag == 'hyperlink':
            # 保留超链接（可能包含图片）
            pass

    # 更新段落属性 (w:pPr)
    _update_pPr(p_element, para_model)

    # 获取对应的 python-docx Paragraph 对象以使用 API 添加 run
    para_obj = _find_paragraph_object(doc, p_element)
    if para_obj is not None:
        _add_runs_to_paragraph(para_obj, para_model)
    else:
        # 回退：直接操作 XML 添加 runs
        _add_runs_via_xml(p_element, para_model)


def _update_pPr(p_element, para_model: Paragraph):
    """更新 <w:p> 元素的 <w:pPr> 段落属性。
    关键原则：model 有值才替换，None 保留原文档格式不删除。"""
    fmt = para_model.format

    # 获取或创建 pPr
    pPr = p_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_element.insert(0, pPr)

    # 对齐方式：仅当 model 有值时替换
    if fmt.alignment:
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        alignment_map = {
            "left": "left", "center": "center",
            "right": "right", "justify": "both",
        }
        jc.set(qn('w:val'), alignment_map.get(fmt.alignment, "left"))
        pPr.append(jc)

    # 缩进：仅当 model 有值时替换，否则保留原文档缩进
    has_indent = (fmt.first_line_indent_pt is not None or
                  fmt.left_indent_pt is not None or
                  fmt.right_indent_pt is not None)
    if has_indent:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        ind = OxmlElement('w:ind')
        if fmt.first_line_indent_pt is not None:
            ind.set(qn('w:firstLine'), str(int(fmt.first_line_indent_pt * 20)))
            # P1-4 修复：使用段落实际字号计算 firstLineChars（此前硬编码 16pt，
            # 标题 22pt/表格 12pt 时缩进字符数计算错误）
            _font_size = 16.0
            if para_model.runs and para_model.runs[0].format.font_size_pt:
                _font_size = para_model.runs[0].format.font_size_pt
            chars = int(round(fmt.first_line_indent_pt / _font_size * 100))
            if chars > 0:
                ind.set(qn('w:firstLineChars'), str(chars))
        if fmt.left_indent_pt is not None:
            ind.set(qn('w:left'), str(int(fmt.left_indent_pt * 20)))
        if fmt.right_indent_pt is not None:
            ind.set(qn('w:right'), str(int(fmt.right_indent_pt * 20)))
        pPr.append(ind)

    # 行距：仅当 model 有值时替换，否则保留原文档行距
    has_spacing = (fmt.line_spacing_pt is not None or
                   fmt.space_before_pt is not None or
                   fmt.space_after_pt is not None)
    if has_spacing:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        if fmt.line_spacing_pt is not None:
            spacing_pt = max(6, min(200, fmt.line_spacing_pt))
            rule = fmt.line_spacing_rule or "exact"
            if rule == "multiple":
                # 倍数行距：w:line 值为 240 分之一行（如 1.5x = 360）
                # line_spacing_pt 存储的是 pt 值，需要反算回倍数
                # 公文标准字号16pt，1倍行距=240（即 16pt * 15 = 240）
                # pt → 240ths: value = spacing_pt / 16 * 240
                line_val = int(round(spacing_pt / 16 * 240))
                spacing.set(qn('w:line'), str(line_val))
                spacing.set(qn('w:lineRule'), 'auto')
            elif rule == "atLeast":
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'atLeast')
            else:
                # exact (默认，符合GB/T 9704标准)
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'exact')
        if fmt.space_before_pt is not None:
            spacing.set(qn('w:before'), str(int(fmt.space_before_pt * 20)))
        if fmt.space_after_pt is not None:
            spacing.set(qn('w:after'), str(int(fmt.space_after_pt * 20)))
        pPr.append(spacing)


# 预构建段落元素 → Paragraph 对象的映射，避免 O(N^2) 查找
# NI9 修复：缓存键为 id(doc)，配合 weakref.finalize 保证文档 GC 时自动清理，避免泄漏
_paragraph_index_cache: dict[int, dict[int, Any]] = {}


def _build_paragraph_index(doc: Document) -> dict[int, Any]:
    """为 doc 对象构建 XML 元素 id → Paragraph 对象的哈希索引（O(N) 一次性构建）。"""
    cache_key = id(doc)
    if cache_key in _paragraph_index_cache:
        return _paragraph_index_cache[cache_key]
    idx: dict[int, Any] = {id(p._element): p for p in doc.paragraphs}
    _paragraph_index_cache[cache_key] = idx
    # NI9: 文档对象被 GC 时自动清理缓存条目
    try:
        import weakref
        weakref.finalize(doc, _paragraph_index_cache.pop, cache_key, None)
    except Exception:
        pass
    return idx


def _clear_paragraph_index(doc: Document) -> None:
    """清除缓存索引（在段落数量变化后调用）。"""
    _paragraph_index_cache.pop(id(doc), None)


def _find_paragraph_object(doc: Document, p_element):
    """通过 XML 元素找到对应的 python-docx Paragraph 对象（O(1) 哈希查找）。"""
    idx = _build_paragraph_index(doc)
    return idx.get(id(p_element))


def _add_runs_to_paragraph(para, para_model: Paragraph):
    """使用 python-docx API 向段落添加 runs。"""
    if para_model.runs:
        for run_model in para_model.runs:
            run = para.add_run(run_model.text)
            _apply_run_format(run, run_model)
    else:
        if para_model.text:
            run = para.add_run(para_model.text)
            # 优先使用段落 format 中的 font_name，fallback 到 BODY_FONT
            fmt_font = None
            if para_model.runs and para_model.runs[0].format:
                fmt_font = para_model.runs[0].format.font_name
            if not fmt_font and para_model.format:
                fmt_font = getattr(para_model.format, 'font_name', None)
            set_run_font(run, fmt_font or BODY_FONT)


def _add_runs_via_xml(p_element, para_model: Paragraph):
    """直接通过 XML 添加 runs（当无法找到 python-docx Paragraph 对象时的回退方案）。"""
    if para_model.runs:
        for run_model in para_model.runs:
            r = OxmlElement('w:r')
            # 添加 run 格式属性
            rPr = OxmlElement('w:rPr')
            fmt = run_model.format
            if fmt.font_name:
                # 字体兜底保护：Latin字体 + CJK文本 → eastAsia 使用 BODY_FONT
                east_asian = fmt.font_name
                if fmt.font_name in _LATIN_FONTS and run_model.text and _contains_cjk(run_model.text):
                    east_asian = BODY_FONT
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), LATIN_FONT)
                rFonts.set(qn('w:hAnsi'), LATIN_FONT)
                rFonts.set(qn('w:eastAsia'), east_asian)
                rFonts.set(qn('w:cs'), LATIN_FONT)
                rPr.append(rFonts)
            if fmt.font_size_pt:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(fmt.font_size_pt * 2)))  # half-points
                rPr.append(sz)
            if fmt.bold:
                rPr.append(OxmlElement('w:b'))
            if fmt.italic:
                rPr.append(OxmlElement('w:i'))
            if len(rPr) > 0:
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = run_model.text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p_element.append(r)
    elif para_model.text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), BODY_FONT)
        rFonts.set(qn('w:cs'), LATIN_FONT)
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = para_model.text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p_element.append(r)


# ---------------------------------------------------------------------------
#  Table Writing
# ---------------------------------------------------------------------------

def _update_tables(doc: Document, model: DocumentModel):
    """
    更新文档中的表格内容。
    如果源文档有表格，更新其单元格内容。
    如果源文档没有表格但 model 有，在末尾添加。

    P2-11 说明：表格与段落替换的交互已由 _replace_paragraphs 的交错遍历处理
    （<w:tbl> 不消耗 model 段落索引，表格保持原位）；此处仅按索引更新/添加表格，
    不删除源文档多余表格（保留未建模内容，遵循"保留策略"）。
    """
    existing_tables = list(doc.tables)
    model_tables = model.tables

    for idx, table_model in enumerate(model_tables):
        if idx < len(existing_tables):
            # 更新已有表格的单元格内容
            _update_table_content(existing_tables[idx], table_model)
        else:
            # 添加新表格
            _add_table(doc, table_model)

    logger.debug(f"Updated {min(len(model_tables), len(existing_tables))} tables, "
                 f"added {max(0, len(model_tables) - len(existing_tables))}")


def _smart_align_cell(cell_text: str, is_header: bool, col_idx: int, total_cols: int) -> str:
    """智能判断单元格对齐方式。

    规则：
    - 表头：居中
    - 序号列（第一列）：居中
    - 短文本（≤4字符）：居中
    - 数字内容（含小数、百分比）：右对齐
    - 其他：左对齐
    """
    text = cell_text.strip()
    if not text:
        return 'left'
    if is_header:
        return 'center'
    # 序号列居中（优先级最高，避免被数字规则覆盖）
    if col_idx == 0:
        return 'center'
    # 短文本居中
    if len(text) <= 4:
        return 'center'
    # 数字右对齐（含小数、百分比、逗号分隔数字）
    import re
    if re.match(r'^[\d.,%‰]+$', text):
        return 'right'
    return 'left'


def _update_table_content(table, table_model: TableModel):
    """更新已有表格的单元格内容（带智能对齐）。"""
    total_cols = len(table.columns) if hasattr(table, 'columns') else 0
    for cell_model in table_model.cells:
        try:
            cell = table.cell(cell_model.row, cell_model.col)
            # 更新单元格中的段落内容
            if cell_model.paragraphs:
                for p_idx, para_model in enumerate(cell_model.paragraphs):
                    if p_idx < len(cell.paragraphs):
                        # 替换已有段落
                        para = cell.paragraphs[p_idx]
                        # 清除旧 runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                        # 添加新 runs
                        _add_runs_to_paragraph(para, para_model)
                        # 更新段落格式（缩进、行距、对齐等）
                        _update_pPr(para._element, para_model)
                    else:
                        # 添加新段落
                        para = cell.add_paragraph()
                        _add_runs_to_paragraph(para, para_model)
                        _apply_paragraph_format(para, para_model)
            elif cell_model.text:
                # 没有详细段落信息，直接设置文本
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    run = para.add_run(cell_model.text)
                    set_run_font(run, BODY_FONT)
                    # 智能对齐
                    is_header = cell_model.row == 0
                    align = _smart_align_cell(cell_model.text, is_header, cell_model.col, total_cols)
                    para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        except Exception as e:
            logger.warning(f"Failed to update table cell ({cell_model.row},{cell_model.col}): {e}")


def _add_table(doc: Document, table_model: TableModel):
    """在文档中添加一个新表格，按 insert_after_index 定位到正确位置。"""
    try:
        rows = max(1, table_model.rows)
        cols = max(1, table_model.cols)
        table = doc.add_table(rows=rows, cols=cols)

        # 按 insert_after_index 移动表格到正确位置
        insert_idx = getattr(table_model, 'insert_after_index', -1)
        if insert_idx >= 0:
            body = doc.element.body
            para_count = 0
            target_elem = None
            for child in body:
                tag = etree.QName(child.tag).localname if child.tag else ''
                if tag == 'p':
                    if para_count == insert_idx:
                        target_elem = child
                        break
                    para_count += 1
            if target_elem is not None:
                tbl_elem = table._tbl
                body.remove(tbl_elem)
                target_elem.addnext(tbl_elem)

        # 设置表格样式（带边框）
        try:
            table.style = 'Table Grid'
        except KeyError:
            # 文档中没有 'Table Grid' 样式时，手动添加边框
            from docx.oxml.ns import qn
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
            borders = tblPr.makeelement(qn('w:tblBorders'), {})
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = borders.makeelement(qn(f'w:{edge}'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',
                    qn('w:space'): '0',
                    qn('w:color'): '000000',
                })
                borders.append(border)
            tblPr.append(borders)

        # 智能对齐：表头行居中加粗，数据行按内容类型对齐
        total_cols = max(1, table_model.cols)
        for cell_model in table_model.cells:
            try:
                cell = table.cell(cell_model.row, cell_model.col)
                # 清除默认段落
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)

                if cell_model.paragraphs:
                    for p_idx, para_model in enumerate(cell_model.paragraphs):
                        if p_idx < len(cell.paragraphs):
                            para = cell.paragraphs[p_idx]
                            _add_runs_to_paragraph(para, para_model)
                            _update_pPr(para._element, para_model)
                        else:
                            para = cell.add_paragraph()
                            _add_runs_to_paragraph(para, para_model)
                            _apply_paragraph_format(para, para_model)
                    # 智能对齐（表头居中加粗，数据行按内容类型）
                    is_header = cell_model.row == 0
                    cell_text = cell_model.text or (cell_model.paragraphs[0].text if cell_model.paragraphs else '')
                    align = _smart_align_cell(cell_text, is_header, cell_model.col, total_cols)
                    align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
                    for para in cell.paragraphs:
                        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                    if is_header:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                elif cell_model.text:
                    if cell.paragraphs:
                        run = cell.paragraphs[0].add_run(cell_model.text)
                        set_run_font(run, BODY_FONT)
                        # 智能对齐
                        is_header = cell_model.row == 0
                        align = _smart_align_cell(cell_model.text, is_header, cell_model.col, total_cols)
                        align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
                        cell.paragraphs[0].alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                        if is_header:
                            run.bold = True
            except Exception as e:
                logger.warning(f"Failed to write table cell ({cell_model.row},{cell_model.col}): {e}")

        logger.debug(f"Added table: {rows}x{cols}")
    except Exception as e:
        logger.exception(f"Failed to add table")


# ---------------------------------------------------------------------------
#  Headers & Footers
# ---------------------------------------------------------------------------

def _update_headers_footers(doc: Document, model: DocumentModel):
    """更新页眉和页脚内容。"""
    _update_hf_list(doc, model.headers, "header")
    _update_hf_list(doc, model.footers, "footer")


def _update_hf_list(doc: Document, hf_models: list[HeaderFooter], hf_type: str):
    """更新一组页眉或页脚。"""
    for hf_model in hf_models:
        try:
            sec_idx = hf_model.section_index
            if sec_idx >= len(doc.sections):
                logger.debug(f"Section {sec_idx} not found, skipping {hf_type}")
                continue

            section = doc.sections[sec_idx]
            target = section.header if hf_type == "header" else section.footer

            if target is None:
                continue

            # 更新页眉/页脚中的段落
            if hf_model.paragraphs:
                for p_idx, para_model in enumerate(hf_model.paragraphs):
                    if p_idx < len(target.paragraphs):
                        para = target.paragraphs[p_idx]
                        # 清除旧 runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                        # 如果包含页码，写入 Word 域代码而非静态文本
                        if hf_model.has_page_number:
                            _add_page_number_field(para, para_model)
                        else:
                            # 添加新 runs
                            _add_runs_to_paragraph(para, para_model)
                        # 更新段落格式（缩进、行距、对齐等）
                        _update_pPr(para._element, para_model)
                    else:
                        para = target.add_paragraph()
                        _add_runs_to_paragraph(para, para_model)
                        _apply_paragraph_format(para, para_model)
            elif hf_model.text:
                # 简单文本模式
                if target.paragraphs:
                    para = target.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    run = para.add_run(hf_model.text)
                    set_run_font(run, BODY_FONT)

            logger.debug(f"Updated {hf_type} section {sec_idx}")
        except Exception as e:
            logger.warning(f"Failed to update {hf_type} section {hf_model.section_index}: {e}")


def _add_page_number_field(para, para_model: Paragraph) -> None:
    """
    在段落中写入 Word 页码域代码（{ PAGE } / { NUMPAGES }）。
    使用 Word XML 域代码实现动态页码，而非静态文本。
    
    支持格式：
    - "{PAGE}" → 当前页码
    - "{NUMPAGES}" → 总页数
    - 可包含前缀后缀，如 "- {PAGE} -" 或 "第 {PAGE} 页 共 {NUMPAGES} 页"
    """
    run_text = para_model.text if para_model.text else "{PAGE}"

    # 构建标准 Word 页码域 XML
    # 格式: 文本 + PAGE域 + 文本 + NUMPAGES域 + 文本
    parts = []
    import re
    # 拆分文本中的 {PAGE} 和 {NUMPAGES} 占位符
    remaining = run_text
    while remaining:
        m = re.search(r'\{PAGE\}|\{NUMPAGES\}', remaining)
        if not m:
            if remaining.strip():
                parts.append(('text', remaining))
            break
        
        # 前置文本
        prefix = remaining[:m.start()]
        if prefix.strip():
            parts.append(('text', prefix))
        
        # 域代码
        parts.append(('field', m.group()))
        remaining = remaining[m.end():]

    # 写入 Word XML
    for part_type, content in parts:
        if part_type == 'text':
            run_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            # P2-5 修复：页码字体按 GB/T 9704 规范（宋体 14pt），改用 font_utils 常量
            # 而非散落的字符串字面量，保证与其它页码注入路径一致
            rFonts.set(qn('w:eastAsia'), PAGE_NUMBER_FONT)
            rFonts.set(qn('w:ascii'), PAGE_NUMBER_LATIN_FONT)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(PAGE_NUMBER_SIZE_PT * 2)))  # 14pt
            rPr.append(sz)
            run_el.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = content
            run_el.append(t)
            para._element.append(run_el)
        elif part_type == 'field':
            # 创建 fldChar begin
            fld_begin = OxmlElement('w:r')
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            fld_begin.append(fldChar_begin)
            para._element.append(fld_begin)

            # 创建 instrText (域代码指令)
            instr = OxmlElement('w:r')
            rPr_instr = OxmlElement('w:rPr')
            rFonts_instr = OxmlElement('w:rFonts')
            rFonts_instr.set(qn('w:eastAsia'), PAGE_NUMBER_FONT)
            rPr_instr.append(rFonts_instr)
            instr.append(rPr_instr)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            field_name = content[1:-1]  # Remove { } → PAGE or NUMPAGES
            instrText.text = f' {field_name} '
            instr.append(instrText)
            para._element.append(instr)

            # 创建 fldChar separate
            fld_sep = OxmlElement('w:r')
            fldChar_sep = OxmlElement('w:fldChar')
            fldChar_sep.set(qn('w:fldCharType'), 'separate')
            fld_sep.append(fldChar_sep)
            para._element.append(fld_sep)

            # 创建默认显示值
            fld_default = OxmlElement('w:r')
            rPr_def = OxmlElement('w:rPr')
            rFonts_def = OxmlElement('w:rFonts')
            rFonts_def.set(qn('w:eastAsia'), '宋体')
            rFonts_def.set(qn('w:ascii'), 'Times New Roman')
            rPr_def.append(rFonts_def)
            sz_def = OxmlElement('w:sz')
            sz_def.set(qn('w:val'), '28')
            rPr_def.append(sz_def)
            fld_default.append(rPr_def)
            t_def = OxmlElement('w:t')
            t_def.set(qn('xml:space'), 'preserve')
            t_def.text = '1'  # 默认显示值
            fld_default.append(t_def)
            para._element.append(fld_default)

            # 创建 fldChar end
            fld_end = OxmlElement('w:r')
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            fld_end.append(fldChar_end)
            para._element.append(fld_end)


# ---------------------------------------------------------------------------
#  Metadata
# ---------------------------------------------------------------------------

def _update_metadata(doc: Document, model: DocumentModel):
    """更新文档核心属性（元数据）。"""
    try:
        meta = model.metadata
        props = doc.core_properties
        if meta.title:
            props.title = meta.title
        if meta.author:
            props.author = meta.author
        if meta.subject:
            props.subject = meta.subject
        if meta.category:
            props.category = meta.category
        logger.debug("Document metadata updated")
    except Exception as e:
        logger.warning(f"Failed to update metadata: {e}")


# ---------------------------------------------------------------------------
#  Format Helpers
# ---------------------------------------------------------------------------

def _apply_paragraph_format(para, para_model: Paragraph):
    """Apply formatting to a paragraph using python-docx API."""
    pf = para.paragraph_format
    fmt = para_model.format

    # Alignment
    if fmt.alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = alignment_map.get(fmt.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Indentation
    if fmt.first_line_indent_pt is not None:
        pf.first_line_indent = Pt(fmt.first_line_indent_pt)
    if fmt.left_indent_pt is not None:
        pf.left_indent = Pt(fmt.left_indent_pt)
    if fmt.right_indent_pt is not None:
        pf.right_indent = Pt(fmt.right_indent_pt)

    # Spacing
    if fmt.space_before_pt is not None:
        pf.space_before = Pt(fmt.space_before_pt)
    if fmt.space_after_pt is not None:
        pf.space_after = Pt(fmt.space_after_pt)

    # Line spacing
    if fmt.line_spacing_pt is not None:
        spacing_pt = max(6, min(200, fmt.line_spacing_pt))
        rule = fmt.line_spacing_rule or "exact"
        if rule == "multiple":
            # MULTIPLE 模式：python-docx 期望倍数（如 1.5），而非 Pt 值
            # 公文标准字号 16pt，行距 28.95pt → 28.95/16 ≈ 1.8
            multiple = spacing_pt / 16.0
            pf.line_spacing = multiple
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        elif rule == "atLeast":
            pf.line_spacing = Pt(spacing_pt)
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            pf.line_spacing = Pt(spacing_pt)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _apply_run_format(run, run_model: Run):
    """
    Apply formatting to a run.
    使用 font_utils.set_run_font 统一处理中文字体。
    """
    fmt = run_model.format

    # === 字体设置（统一入口） ===
    if fmt.font_name:
        set_run_font(run, fmt.font_name)
    else:
        set_run_font(run, BODY_FONT)

    # === 字号 ===
    if fmt.font_size_pt is not None:
        run.font.size = Pt(fmt.font_size_pt)

    # === 样式 ===
    if fmt.bold is not None:
        run.font.bold = fmt.bold
    if fmt.italic is not None:
        run.font.italic = fmt.italic
    if fmt.underline is not None:
        run.font.underline = fmt.underline
    if fmt.strikethrough is True:
        run.font.strike = True
    # 注意：strikethrough=False 时不写入 strike 元素
    # 否则 python-docx 会生成 <w:strike w:val="false"/>，
    # 被 parser_format.py 误判为有删除线

    # === 颜色 ===
    if fmt.color:
        try:
            rgb_str = fmt.color.replace("#", "")
            if len(rgb_str) == 6:
                r = int(rgb_str[0:2], 16)
                g = int(rgb_str[2:4], 16)
                b = int(rgb_str[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception:
            pass
