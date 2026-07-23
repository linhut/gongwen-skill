# -*- coding: utf-8 -*-
#
# 公文文档格式化 Skill —— 版头 / 版记 / 页码注入（独立发行版）
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# 项目出处：AI 公文智能优化助手 (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
#
# 本模块从原项目 backend/api/routes/optimize.py 中抽取版头/版记/页码
# 三大注入逻辑，剥离 FastAPI / 数据库依赖，使其可独立运行。
"""
GB/T 9704《党政机关公文格式》版式要素注入：

  版头（header）  —— 发文机关标志 + 发文字号 + 签发人 + 红色反线
  版记（footer）  —— 抄送机关 + 印发机关 + 印发日期 + 分隔线
  页码（pagenum） —— Word PAGE 域动态页码，支持单双页奇偶排版

三个入口函数直接在 .docx 文件上原地注入，输入输出均为文件路径。
"""
from __future__ import annotations

from utils.logger import logger


# ---------------------------------------------------------------------------
#  内部工具
# ---------------------------------------------------------------------------

def _insert_before(new_para, reference_p, body) -> None:
    """将段落插入到 reference_p 之前（若无参照则追加到末尾）。"""
    if reference_p is not None:
        body.insert(list(body).index(reference_p), new_para._element)
    else:
        body.append(new_para._element)


# ---------------------------------------------------------------------------
#  版头注入：发文机关标志 + 发文字号 + 签发人 + 红色反线
# ---------------------------------------------------------------------------

def inject_header(output_path: str, header_config: dict) -> None:
    """注入版头：发文机关标志 + 空行 + 发文字号 + 红色分隔线。

    header_config 字段：
        org_name   发文机关标志（红色 30pt 方正小标宋简体，居中）
        doc_number 发文字号（三号仿宋）
        signer     签发人姓名（上行文时与发文字号同行右对齐）

    公文版头顺序（从上到下）：机关名 → 空二行 → 发文字号/签发人 → 红色反线
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from core.document.font_utils import set_run_font, TITLE_FONT, BODY_FONT

        doc = Document(output_path)
        org_name = header_config.get('org_name', '')
        doc_number = header_config.get('doc_number', '')
        signer = header_config.get('signer', '')

        if not org_name:
            logger.warning("inject_header: org_name 为空，跳过版头注入")
            return

        body = doc.element.body
        first_p = body.find(qn('w:p'))

        # 1. 发文机关标志：红色 30pt 方正小标宋简体 居中（最顶部）
        p_org = doc.add_paragraph()
        p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_org = p_org.add_run(org_name)
        set_run_font(run_org, TITLE_FONT)
        run_org.font.size = Pt(30)
        run_org.font.color.rgb = RGBColor(0xE0, 0x00, 0x00)
        p_org.paragraph_format.line_spacing = Pt(28.95)
        p_org.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        _insert_before(p_org, first_p, body)

        # 2. 空二行（保持版头与发文字号之间的间距）
        for _ in range(2):
            p_empty = doc.add_paragraph()
            p_empty.paragraph_format.line_spacing = Pt(28.95)
            p_empty.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_empty, first_p, body)

        # 3. 计算右对齐制表位位置（基于页面实际宽度）
        page_w_mm, margin_l_mm, margin_r_mm = 210.0, 28.0, 26.0
        try:
            pw = doc.sections[0].page_width
            ml = doc.sections[0].left_margin
            mr = doc.sections[0].right_margin
            if pw and ml and mr:
                tab_pos = int((pw - ml - mr) / 635)  # EMU → twips
            else:
                tab_pos = int((page_w_mm - margin_l_mm - margin_r_mm) / 25.4 * 1440)
        except Exception:
            tab_pos = int((page_w_mm - margin_l_mm - margin_r_mm) / 25.4 * 1440)

        # 4. 发文字号 / 签发人
        if doc_number and signer:
            # 同一段落：发文号左 + 制表位右 + 签发人右
            p_num = doc.add_paragraph()
            p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p_num.paragraph_format
            pf.line_spacing = Pt(28.95)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

            tabs_el = OxmlElement('w:tabs')
            tab_el = OxmlElement('w:tab')
            tab_el.set(qn('w:val'), 'right')
            tab_el.set(qn('w:pos'), str(tab_pos))
            tab_el.set(qn('w:leader'), 'none')
            tabs_el.append(tab_el)
            pPr = p_num._element.get_or_add_pPr()
            pPr.append(tabs_el)

            run_num = p_num.add_run(doc_number)
            set_run_font(run_num, BODY_FONT)
            run_num.font.size = Pt(16)
            p_num.add_run('\t')
            run_signer = p_num.add_run(f'签发人：{signer}')
            set_run_font(run_signer, BODY_FONT)
            run_signer.font.size = Pt(16)
            _insert_before(p_num, first_p, body)
        elif doc_number:
            # 仅有发文号：居中
            p_num = doc.add_paragraph()
            p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_num = p_num.add_run(doc_number)
            set_run_font(run_num, BODY_FONT)
            run_num.font.size = Pt(16)
            p_num.paragraph_format.line_spacing = Pt(28.95)
            p_num.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_num, first_p, body)
        elif signer:
            # 仅有签发人：右对齐
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_signer = p_signer.add_run(f'签发人：{signer}')
            set_run_font(run_signer, BODY_FONT)
            run_signer.font.size = Pt(16)
            p_signer.paragraph_format.line_spacing = Pt(28.95)
            p_signer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_signer, first_p, body)

        # 5. 红色反线（底部边框）— 紧贴发文字号下方
        p_border = doc.add_paragraph()
        p_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p_border._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '15')
        bottom.set(qn('w:color'), 'E00000')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '60')
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)
        _insert_before(p_border, first_p, body)

        doc.save(output_path)
        logger.info(f"版头已注入: {output_path}")
    except Exception as e:
        logger.error(f"版头注入失败: {e}", exc_info=True)
        raise


# ---------------------------------------------------------------------------
#  版记注入：抄送 + 印发机关 + 印发日期 + 分隔线
# ---------------------------------------------------------------------------

def inject_footer(output_path: str, footer_config: dict) -> None:
    """注入版记：分隔线 + 抄送 + 印发机关/日期 + 分隔线。

    footer_config 字段：
        cc         抄送机关
        printer    印发机关
        print_date 印发日期（亦兼容 printDate）

    GB/T 9704 §7.4：版记用三号仿宋（16pt），左空一字，0.35mm 黑色实线分隔。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from core.document.font_utils import set_run_font, BODY_FONT

        doc = Document(output_path)
        cc = footer_config.get('cc', '')
        printer = footer_config.get('printer', '')
        print_date = footer_config.get('printDate', footer_config.get('print_date', ''))

        if not (cc or printer or print_date):
            logger.warning("inject_footer: 抄送/印发信息均为空，跳过版记注入")
            return

        # === 版记换页逻辑：末页剩余空间不足则强制分页 ===
        try:
            section = doc.sections[0]
            ph, tm, bm = section.page_height, section.top_margin, section.bottom_margin
            if ph and tm and bm:
                content_height_emu = ph - tm - bm
                line_height_emu = 367665  # ≈ 28.95pt
                lines_per_page = int(content_height_emu / line_height_emu)
                para_count = len(doc.paragraphs)
                MIN_LINES_FOR_FOOTER = 3
                last_page_paras = para_count % lines_per_page or lines_per_page
                remaining_lines = lines_per_page - last_page_paras
                if remaining_lines < MIN_LINES_FOR_FOOTER:
                    page_break_para = doc.add_paragraph()
                    run_pb = page_break_para.add_run('')
                    br = OxmlElement('w:br')
                    br.set(qn('w:type'), 'page')
                    run_pb._element.append(br)
                    logger.info(f"版记空间不足（剩余 {remaining_lines} 行），强制分页")
        except Exception as e:
            logger.warning(f"版记换页计算失败: {e}")

        def _add_border_para(border_size='8', border_color='000000'):
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '0')
            ind.set(qn('w:right'), '0')
            ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), border_size)
            top.set(qn('w:color'), border_color)
            top.set(qn('w:space'), '0')
            pBdr.append(top)
            pPr.append(pBdr)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '200')
            spacing.set(qn('w:lineRule'), 'atLeast')
            pPr.append(spacing)
            return p

        # 1. 上分隔线（0.35mm 黑色实线），前置 200pt 间距将版记推到末页底部
        p_top_line = doc.add_paragraph()
        pPr_top = p_top_line._element.get_or_add_pPr()
        ind_top = OxmlElement('w:ind')
        ind_top.set(qn('w:left'), '0')
        ind_top.set(qn('w:right'), '0')
        pPr_top.append(ind_top)
        spacing_top = OxmlElement('w:spacing')
        spacing_top.set(qn('w:before'), str(200 * 20))
        spacing_top.set(qn('w:after'), '0')
        spacing_top.set(qn('w:line'), '200')
        spacing_top.set(qn('w:lineRule'), 'atLeast')
        pPr_top.append(spacing_top)
        pBdr_top = OxmlElement('w:pBdr')
        top_line = OxmlElement('w:top')
        top_line.set(qn('w:val'), 'single')
        top_line.set(qn('w:sz'), '8')
        top_line.set(qn('w:color'), '000000')
        top_line.set(qn('w:space'), '0')
        pBdr_top.append(top_line)
        pPr_top.append(pBdr_top)

        # 2. 抄送行（三号仿宋，左空一字，悬挂缩进）
        if cc:
            p_cc = doc.add_paragraph()
            p_cc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_cc = p_cc.add_run(f'抄送：{cc}。')
            set_run_font(run_cc, BODY_FONT)
            run_cc.font.size = Pt(16)
            p_cc.paragraph_format.line_spacing = Pt(28.95)
            p_cc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pPr_cc = p_cc._element.get_or_add_pPr()
            ind_cc = OxmlElement('w:ind')
            ind_cc.set(qn('w:left'), '320')
            ind_cc.set(qn('w:hanging'), '320')
            pPr_cc.append(ind_cc)
            if printer or print_date:
                _add_border_para('8')

        # 3. 印发机关 + 印发日期（同一行，左右空一字，日期制表位推到最右）
        if printer or print_date:
            p_info = doc.add_paragraph()
            p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pPr_info = p_info._element.get_or_add_pPr()
            ind_info = OxmlElement('w:ind')
            ind_info.set(qn('w:left'), '320')
            ind_info.set(qn('w:right'), '320')
            pPr_info.append(ind_info)
            tabs_el = OxmlElement('w:tabs')
            tab_el = OxmlElement('w:tab')
            tab_el.set(qn('w:val'), 'right')
            tab_el.set(qn('w:pos'), '8500')
            tab_el.set(qn('w:leader'), 'none')
            tabs_el.append(tab_el)
            pPr_info.append(tabs_el)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '579')
            spacing.set(qn('w:lineRule'), 'exact')
            pPr_info.append(spacing)
            if printer:
                run_printer = p_info.add_run(printer)
                set_run_font(run_printer, BODY_FONT)
                run_printer.font.size = Pt(16)
            if printer and print_date:
                p_info.add_run('\t')
            if print_date:
                run_date = p_info.add_run(f'{print_date}印发')
                set_run_font(run_date, BODY_FONT)
                run_date.font.size = Pt(16)

        # 4. 下分隔线
        _add_border_para('8')

        doc.save(output_path)
        logger.info(f"版记已注入: {output_path}")
    except Exception as e:
        logger.error(f"版记注入失败: {e}", exc_info=True)
        raise


# ---------------------------------------------------------------------------
#  页码注入：Word PAGE 域动态页码
# ---------------------------------------------------------------------------

def _add_page_run(para_elem, text: str, font_name: str, size_pt: int) -> None:
    """添加一个普通文本 run 到段落元素。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    para_elem.append(r)


def _build_page_number_xml(fmt: str, font_name: str, size_pt: int) -> list:
    """解析页码格式字符串，返回 (类型, 内容, 字体, 字号) 元素列表。"""
    import re
    elements = []
    remaining = fmt
    while remaining:
        m = re.search(r'\{PAGE\}|\{NUMPAGES\}', remaining)
        if not m:
            if remaining.strip():
                elements.append(('text', remaining, font_name, size_pt))
            break
        prefix = remaining[:m.start()]
        if prefix:
            elements.append(('text', prefix, font_name, size_pt))
        field_name = m.group()[1:-1]
        elements.append(('field', field_name, font_name, size_pt))
        remaining = remaining[m.end():]
    return elements


def _apply_page_number_elements(para_elem, elements: list) -> None:
    """将 _build_page_number_xml 返回的元素列表写入段落 XML。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for elem in elements:
        etype = elem[0]
        if etype == 'text':
            _add_page_run(para_elem, elem[1], elem[2], elem[3])
        elif etype == 'field':
            field_name, font_name, size_pt = elem[1], elem[2], elem[3]

            r1 = OxmlElement('w:r')
            fc1 = OxmlElement('w:fldChar')
            fc1.set(qn('w:fldCharType'), 'begin')
            r1.append(fc1)
            para_elem.append(r1)

            r2 = OxmlElement('w:r')
            rPr2 = OxmlElement('w:rPr')
            rf2 = OxmlElement('w:rFonts')
            rf2.set(qn('w:eastAsia'), font_name)
            rf2.set(qn('w:ascii'), 'Times New Roman')
            rPr2.append(rf2)
            r2.append(rPr2)
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = f' {field_name} '
            r2.append(instr)
            para_elem.append(r2)

            r3 = OxmlElement('w:r')
            fc3 = OxmlElement('w:fldChar')
            fc3.set(qn('w:fldCharType'), 'separate')
            r3.append(fc3)
            para_elem.append(r3)

            r4 = OxmlElement('w:r')
            rPr4 = OxmlElement('w:rPr')
            rf4 = OxmlElement('w:rFonts')
            rf4.set(qn('w:eastAsia'), font_name)
            rf4.set(qn('w:ascii'), 'Times New Roman')
            rPr4.append(rf4)
            sz4 = OxmlElement('w:sz')
            sz4.set(qn('w:val'), str(int(size_pt * 2)))
            rPr4.append(sz4)
            r4.append(rPr4)
            t4 = OxmlElement('w:t')
            t4.set(qn('xml:space'), 'preserve')
            t4.text = '1'
            r4.append(t4)
            para_elem.append(r4)

            r5 = OxmlElement('w:r')
            fc5 = OxmlElement('w:fldChar')
            fc5.set(qn('w:fldCharType'), 'end')
            r5.append(fc5)
            para_elem.append(r5)


def _inject_even_page_footer_direct(output_path: str, fmt: str, font_name: str, size_pt: int) -> None:
    """直接操作 ZIP，添加偶数页页脚和 evenAndOddHeaders（单双页奇偶排版）。"""
    import zipfile, io
    from lxml import etree

    even_ftr_xml = (
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:p><w:pPr>'
        '<w:jc w:val="left"/>'
        '<w:ind w:left="280"/>'
        '</w:pPr></w:p></w:ftr>'
    )
    even_ftr = etree.fromstring(even_ftr_xml.encode())
    even_p = even_ftr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    elems = _build_page_number_xml(fmt, font_name, size_pt)
    _apply_page_number_elements(even_p, elems)

    buf = io.BytesIO()
    with zipfile.ZipFile(output_path, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    root = etree.fromstring(data)
                    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    for sp in root.findall('.//{%s}sectPr' % ns_w):
                        even_fr = etree.SubElement(sp, '{%s}footerReference' % ns_w)
                        even_fr.set('{%s}type' % ns_w, 'even')
                        even_fr.set('{%s}id' % ns_r, 'rIdFtrEven')
                        for i, child in enumerate(sp):
                            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                            if tag == 'footerReference' and child.get('{%s}type' % ns_w) == 'default':
                                sp.remove(even_fr)
                                sp.insert(i + 1, even_fr)
                                break
                    zout.writestr(item.filename,
                        etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True))
                elif item.filename == 'word/settings.xml':
                    root = etree.fromstring(data)
                    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    existing = root.find('{%s}evenAndOddHeaders' % ns_w)
                    if existing is None:
                        eo = etree.SubElement(root, '{%s}evenAndOddHeaders' % ns_w)
                        root.insert(0, eo)
                    zout.writestr(item.filename,
                        etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True))
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels = data.decode('utf-8')
                    if 'rIdFtrEven' not in rels:
                        rels = rels.replace('</Relationships>',
                            '<Relationship Id="rIdFtrEven" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer2.xml"/></Relationships>')
                    zout.writestr(item.filename, rels.encode())
                else:
                    zout.writestr(item.filename, data)
            zout.writestr('word/footer2.xml',
                etree.tostring(even_ftr, xml_declaration=True, encoding='UTF-8', standalone=True))

    with open(output_path, 'wb') as f:
        f.write(buf.getvalue())
    logger.info("偶数页页脚已注入（直接 ZIP 方式）")


def inject_page_number(output_path: str, page_number_config: dict) -> None:
    """注入页码：使用 Word PAGE 域实现动态页码。

    page_number_config 字段：
        enabled   是否启用（默认 True）
        font      字体（默认 宋体）
        size      字号（默认 14）
        alignment 对齐（center / left / right；right 表示单右双左奇偶排版）
        format    格式（默认 '— {PAGE} —'，可用 {PAGE} / {NUMPAGES}）

    GB/T 9704：页码用四号宋体（14pt），单页码居中，或单右双左奇偶排版。
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        enabled = page_number_config.get('enabled')
        if enabled is None:
            enabled = page_number_config.get('show', True)
        if not enabled:
            return

        fmt = page_number_config.get('format', '— {PAGE} —')
        font_name = page_number_config.get('font', '宋体')
        size_pt = page_number_config.get('size', 14)
        align = page_number_config.get('alignment') or page_number_config.get('position', 'center')
        if align == 'right-left':
            align = 'right'

        doc = Document(output_path)

        is_odd_even = align in ('left', 'right', 'right-left')
        if is_odd_even:
            doc.settings.odd_and_even_pages_header_footer = True

        for section in doc.sections:
            try:
                footer = section.footer
            except Exception:
                footer = None
            if footer is None:
                continue
            for p in list(footer.paragraphs):
                p._element.getparent().remove(p._element)
            para = footer.add_paragraph()
            pPr = para._element.get_or_add_pPr()
            align_map = {'center': 'center', 'left': 'left', 'right': 'right'}
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), align_map.get(align, 'center'))
            pPr.append(jc)
            if is_odd_even:
                ind = OxmlElement('w:ind')
                if align == 'right':
                    ind.set(qn('w:right'), str(int(14 * 20)))
                else:
                    ind.set(qn('w:left'), str(int(14 * 20)))
                pPr.append(ind)
            try:
                bm = section.bottom_margin
                if bm:
                    bm_mm = bm / 36000
                    fd_mm = max(2, bm_mm - 7)
                    section.footer_distance = int(fd_mm * 36000)
                else:
                    section.footer_distance = Pt(22)
            except Exception:
                section.footer_distance = Pt(22)
            elems = _build_page_number_xml(fmt, font_name, size_pt)
            _apply_page_number_elements(para._element, elems)

        doc.save(output_path)

        if is_odd_even:
            _inject_even_page_footer_direct(output_path, fmt, font_name, size_pt)

        logger.info(f"页码已注入: {output_path}")
    except Exception as e:
        logger.error(f"页码注入失败: {e}", exc_info=True)
        raise
