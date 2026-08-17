# -*- coding: utf-8 -*-
"""gongwen/_legacy.py cmd_optimize_content 测试——通过 main() 直接调用。"""
import gongwen._legacy as legacy
import sys
import json
from pathlib import Path
from unittest.mock import patch

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from gongwen._bootstrap import _ENGINE_DIR  # noqa: F401
sys.path.insert(0, str(_ENGINE_DIR))


def _make_docx(tmp_path):
    """创建测试 docx 文件。"""
    from docx import Document
    doc = Document()
    doc.add_paragraph("测试正文段落一，内容比较长。")
    doc.add_paragraph("测试正文段落二，内容也比较长。")
    path = tmp_path / "input.docx"
    doc.save(str(path))
    return path


def _make_changes(tmp_path, docx_path):
    """创建 changes JSON。"""
    from docx import Document
    doc = Document(str(docx_path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    changes = []
    for i, text in enumerate(paras):
        changes.append({
            "paragraph_index": i,
            "original_text": text,
            "optimized_text": text.replace("测试", "优化后"),
            "reason": "测试修改",
        })
    path = tmp_path / "changes.json"
    path.write_text(json.dumps(changes, ensure_ascii=False), encoding="utf-8")
    return path


class TestOptimizeContentTracked:
    """cmd_optimize_content --mode tracked"""

    def test_tracked_mode(self, tmp_path):
        docx = _make_docx(tmp_path)
        changes = _make_changes(tmp_path, docx)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '--changes', str(changes), '-o', str(out),
                                        '--apply', '--mode', 'tracked', '-t', 'notice']):
            result = legacy.main()
            assert result in (0, None)
            assert out.exists()


class TestOptimizeContentInline:
    """cmd_optimize_content --mode inline"""

    def test_inline_mode(self, tmp_path):
        docx = _make_docx(tmp_path)
        changes = _make_changes(tmp_path, docx)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '--changes', str(changes), '-o', str(out),
                                        '--apply', '--mode', 'inline', '-t', 'notice']):
            result = legacy.main()
            assert result in (0, None)
            assert out.exists()


class TestOptimizeContentCommentMode:
    """cmd_optimize_content --comment-mode"""

    def test_comment_mode(self, tmp_path):
        docx = _make_docx(tmp_path)
        changes = _make_changes(tmp_path, docx)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '--changes', str(changes), '-o', str(out),
                                        '--apply', '--comment-mode', '-t', 'notice']):
            result = legacy.main()
            assert result in (0, None)
            assert out.exists()


class TestOptimizeContentDefaultMode:
    """cmd_optimize_content 默认模式（tracked）"""

    def test_default_mode(self, tmp_path):
        docx = _make_docx(tmp_path)
        changes = _make_changes(tmp_path, docx)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '--changes', str(changes), '-o', str(out),
                                        '--apply', '-t', 'notice']):
            result = legacy.main()
            assert result in (0, None)
            assert out.exists()


class TestOptimizeContentNoChanges:
    """cmd_optimize_content 无 --changes 应报错"""

    def test_no_changes(self, tmp_path):
        docx = _make_docx(tmp_path)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '-o', str(out), '--apply', '-t', 'notice']):
            try:
                legacy.main()
                assert False, "应报错"
            except SystemExit:
                pass  # 预期退出码非 0


class TestOptimizeContentInvalidMode:
    """cmd_optimize_content 无效模式应报错"""

    def test_invalid_mode(self, tmp_path):
        docx = _make_docx(tmp_path)
        changes = _make_changes(tmp_path, docx)
        out = tmp_path / "out.docx"
        with patch.object(sys, 'argv', ['gongwen', 'optimize-content', str(docx),
                                        '--changes', str(changes), '-o', str(out),
                                        '--apply', '--mode', 'invalid', '-t', 'notice']):
            try:
                legacy.main()
                assert False
            except SystemExit:
                pass  # argparse 拒绝无效 mode
