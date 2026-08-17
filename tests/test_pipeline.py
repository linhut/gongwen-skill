"""Smoke tests for the parse → generate round-trip pipeline."""
import pytest
from pathlib import Path
import tempfile
import os


@pytest.mark.integration
class TestParseGenerateSmoke:
    """Integration tests that require python-docx and real file I/O."""

    def _create_minimal_docx(self, path: Path) -> None:
        """Create a .docx with title/body/signature/date for testing (P3-36: 增强测试文档)."""
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        p = doc.add_paragraph("关于召开2026年工作会议的通知")
        for r in p.runs:
            r.font.name = "方正小标宋简体"
            r.font.size = Pt(22)
        doc.add_paragraph("各部门：")
        doc.add_paragraph("为确保2026年工作顺利开展，经研究决定召开工作会议。现将有关事项通知如下：")
        doc.add_paragraph("一、会议时间")
        doc.add_paragraph("2026年7月28日上午9时。")
        doc.add_paragraph("特此通知。")
        doc.add_paragraph("XX局办公室")
        doc.add_paragraph("二〇二六年七月二十八日")
        doc.save(str(path))

    def test_parse_generate_roundtrip(self, tmp_path):
        """Parse a minimal docx and regenerate it without errors (P3-26: 断言深度增强)."""
        from core.document.parser import parse_docx
        from core.document.generator import generate_docx

        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        self._create_minimal_docx(src)

        # Parse
        model = parse_docx(str(src))
        assert model is not None
        assert len(model.paragraphs) >= 4, f"Expected >=4 paragraphs, got {len(model.paragraphs)}"

        # Check that key content survived parsing
        texts = [p.text for p in model.paragraphs]
        assert any("通知" in t for t in texts)
        assert any("特此通知" in t for t in texts)

        # 段落索引连续（P3-26 深度断言）
        indexes = [p.index for p in model.paragraphs]
        assert indexes == sorted(indexes), "段落索引应保持顺序"
        # 标题应被识别（方正小标宋 + 22pt）
        assert any(p.is_heading for p in model.paragraphs), "应识别出标题段落"

        # Generate
        result = generate_docx(model, str(out))
        assert result.exists()
        assert result.stat().st_size > 0

        # round-trip 后重新解析，内容不应丢失（P3-26 深度断言）
        from docx import Document as DocxReader
        reopened = DocxReader(str(out))
        reopened_texts = [p.text for p in reopened.paragraphs]
        assert any("特此通知" in t for t in reopened_texts), "round-trip 后内容应保留"

    def test_parse_with_rules_check(self, tmp_path):
        """Parse a docx and run the rule engine check without errors."""
        from core.document.parser import parse_docx
        from core.rules.engine import RuleEngine

        src = tmp_path / "input.docx"
        self._create_minimal_docx(src)

        model = parse_docx(str(src))
        engine = RuleEngine()
        issues = engine.check(model, "notice")

        # issues should be a list (possibly empty for a minimal docx)
        assert isinstance(issues, list)

    def test_parse_fix_generate(self, tmp_path):
        """Full pipeline: parse → check_and_fix → generate."""
        from core.document.parser import parse_docx
        from core.document.generator import generate_docx
        from core.rules.engine import RuleEngine

        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        self._create_minimal_docx(src)

        model = parse_docx(str(src))
        engine = RuleEngine()
        issues, fixed = engine.check_and_fix(model, "notice")
        result = generate_docx(fixed, str(out))
        assert result.exists()
        assert result.stat().st_size > 0
