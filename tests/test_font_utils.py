"""Tests for font_utils — core font handling (P3-24/P3-32 补充测试)."""
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "engine"))

from core.document.font_utils import (
    set_run_font, get_effective_font, validate_font_name,
    get_font_fallback, _contains_cjk, BODY_FONT, LATIN_FONT,
)


class TestContainsCjk:
    def test_chinese_true(self):
        assert _contains_cjk("正文内容") is True

    def test_ascii_false(self):
        assert _contains_cjk("Hello World") is False

    def test_mixed_true(self):
        assert _contains_cjk("中文ABC") is True


class TestSetRunFont:
    def test_set_basic_font(self):
        doc = Document()
        run = doc.add_paragraph().add_run("测试")
        set_run_font(run, "仿宋_GB2312")
        rFonts = run._element.rPr.rFonts
        assert rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia') == "仿宋_GB2312"

    def test_empty_font_keeps_default(self):
        doc = Document()
        run = doc.add_paragraph().add_run("测试")
        set_run_font(run, "")
        assert get_effective_font(run) in (None, LATIN_FONT, BODY_FONT)


class TestGetEffectiveFont:
    def test_returns_east_asia_font(self):
        doc = Document()
        run = doc.add_paragraph().add_run("中文")
        set_run_font(run, "黑体")
        eff = get_effective_font(run)
        assert eff == "黑体"


class TestValidateFontName:
    def test_valid_font(self):
        assert validate_font_name("仿宋_GB2312") is True

    def test_invalid_empty(self):
        # 空值视为无效（未设置字体）
        assert validate_font_name("") is False

    def test_none(self):
        assert validate_font_name(None) is False

    def test_invalid_ms_gothic(self):
        assert validate_font_name("MS Gothic") is False


class TestGetFontFallback:
    def test_known_fallback(self):
        assert get_font_fallback("方正小标宋简体") == "SimSun"

    def test_unknown_returns_input(self):
        assert get_font_fallback("未知字体") == "未知字体"
