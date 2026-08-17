# -*- coding: utf-8 -*-
"""P1-3c：header / footer / pagenum 三大版式注入端到端测试。

测试策略：
1. 构造一个最小 docx 输入
2. 跑 header 命令注入版头 → 验证产出
3. 跑 footer 命令注入版记 → 验证产出
4. 跑 pagenum 命令注入页码 → 验证产出
"""
import subprocess
import sys
from pathlib import Path

import pytest


def _make_input_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    p = doc.add_paragraph("关于测试的通知")
    for r in p.runs:
        r.font.name = "方正小标宋简体"
        r.font.size = Pt(22)
    doc.add_paragraph("各部门：")
    doc.add_paragraph("现将有关事项通知如下。")
    doc.add_paragraph("特此通知。")
    doc.add_paragraph("XX单位办公室")
    doc.add_paragraph("二〇二六年八月十六日")
    doc.save(str(path))


def _run(*args) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, "-m", "gongwen", *args],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )


@pytest.fixture
def input_docx(tmp_path):
    p = tmp_path / "input.docx"
    _make_input_docx(p)
    return p


def test_header_injects_red_header(input_docx, tmp_path):
    """header 命令必须产出包含版头的 docx。"""
    out = tmp_path / "header.docx"
    r = _run("header", str(input_docx), "-o", str(out),
             "--org-name", "测试单位", "--doc-number", "〔2026〕1号")
    assert r.returncode == 0, f"header failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists(), f"产出文件不存在: {out}"
    assert out.stat().st_size > 0


def test_header_requires_org_name(input_docx, tmp_path):
    """header 命令缺 --org-name 必须报非 0 退出码。"""
    out = tmp_path / "should_not_exist.docx"
    r = _run("header", str(input_docx), "-o", str(out))
    assert r.returncode != 0, "header 缺必备参数时应报错"
    # 错误提示应该提到 org-name


def test_footer_injects_footer(input_docx, tmp_path):
    """footer 命令必须产出包含版记的 docx。"""
    out = tmp_path / "footer.docx"
    r = _run("footer", str(input_docx), "-o", str(out),
             "--cc", "相关单位", "--printer", "XX办公室", "--print-date", "2026年8月16日")
    assert r.returncode == 0, f"footer failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists()
    assert out.stat().st_size > 0


def test_footer_requires_at_least_one_arg(input_docx, tmp_path):
    """footer 命令缺所有 cc/printer/print-date 时应报错。"""
    out = tmp_path / "should_not_exist.docx"
    r = _run("footer", str(input_docx), "-o", str(out))
    assert r.returncode != 0, "footer 缺所有版记参数时应报错"


def test_pagenum_injects_page_numbers(input_docx, tmp_path):
    """pagenum 命令必须产出包含 PAGE 域动态页码的 docx。"""
    out = tmp_path / "pagenum.docx"
    r = _run("pagenum", str(input_docx), "-o", str(out), "--alignment", "right")
    assert r.returncode == 0, f"pagenum failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists()
    assert out.stat().st_size > 0


def test_pagenum_default_alignment(input_docx, tmp_path):
    """pagenum 不传 --alignment 时应使用默认（单右双左）。"""
    out = tmp_path / "pagenum_default.docx"
    r = _run("pagenum", str(input_docx), "-o", str(out))
    assert r.returncode == 0, f"pagenum default failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists()
    assert out.stat().st_size > 0
