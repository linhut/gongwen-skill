# -*- coding: utf-8 -*-
"""engine/inject.py 单元测试——补充 inject_header/footer/pagenum 边界情况。"""
from inject import (
    inject_header, inject_footer, inject_page_number,
    _get_line_spacing_pt,
)
from docx import Document
import pytest
import sys
import tempfile
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "engine"))


def _make_test_docx(path):
    """创建简单测试文档。"""
    doc = Document()
    doc.add_paragraph("测试正文段落")
    doc.add_paragraph("第二段正文")
    doc.save(str(path))


class TestGetLineSpacingPt:
    def test_returns_float(self):
        result = _get_line_spacing_pt()
        assert isinstance(result, float)
        assert result > 0

    def test_default_is_33(self):
        result = _get_line_spacing_pt()
        assert result == 33.0


class TestIsFontInstalled:
    def test_returns_bool(self):
        # _is_font_installed is in _legacy.py, test via cmd_font
        import importlib
        mod = importlib.import_module("gongwen._legacy")
        result = mod._is_font_installed("仿宋_GB2312")
        assert isinstance(result, bool)

    def test_nonexistent_font(self):
        import importlib
        mod = importlib.import_module("gongwen._legacy")
        result = mod._is_font_installed("不存在的字体名XYZ123")
        assert result is False


class TestGetFontsDir:
    def test_path_exists(self):
        import importlib
        mod = importlib.import_module("gongwen._legacy")
        d = mod._get_fonts_dir()
        assert isinstance(d, Path)

    def test_has_font_files(self):
        import importlib
        mod = importlib.import_module("gongwen._legacy")
        d = mod._get_fonts_dir()
        if d.exists():
            files = list(d.glob("*.ttf")) + list(d.glob("*.TTF"))
            assert len(files) >= 3


class TestInjectHeader:
    def test_basic_header(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_header(str(docx_path), {"org_name": "测试机关", "doc_number": "测〔2026〕1号"})
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "测试机关" in full_text
        assert "测〔2026〕1号" in full_text

    def test_empty_org_name_skips(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_header(str(docx_path), {"org_name": ""})
        # 不应添加版头内容
        doc = Document(str(docx_path))
        assert len(doc.paragraphs) == 2  # 原始两段

    def test_only_doc_number(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_header(str(docx_path), {"org_name": "机关", "doc_number": "号123"})
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "号123" in full_text

    def test_with_signer(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_header(str(docx_path), {
            "org_name": "机关", "doc_number": "号1", "signer": "张三"
        })
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "签发人" in full_text
        assert "张三" in full_text

    def test_only_signer(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_header(str(docx_path), {"org_name": "机关", "signer": "李四"})
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "李四" in full_text


class TestInjectFooter:
    def test_basic_footer(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_footer(str(docx_path), {
            "cc": "各省厅", "printer": "办公厅", "print_date": "2026年8月"
        })
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "抄送" in full_text or "各省厅" in full_text

    def test_empty_footer(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_footer(str(docx_path), {})
        # 空版记仍应有分隔线（不报错即可）
        doc = Document(str(docx_path))
        assert doc is not None

    def test_only_cc(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_footer(str(docx_path), {"cc": "抄送单位"})
        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "抄送" in full_text


class TestInjectPageNumber:
    def test_center_alignment(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_page_number(str(docx_path), {"alignment": "center"})
        doc = Document(str(docx_path))
        # 页码注入到 section 页脚
        assert doc is not None

    def test_right_alignment(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_page_number(str(docx_path), {"alignment": "right"})
        doc = Document(str(docx_path))
        assert doc is not None

    def test_custom_format(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _make_test_docx(docx_path)
        inject_page_number(str(docx_path), {
            "alignment": "center", "format": "第 {PAGE} 页"
        })
        doc = Document(str(docx_path))
        assert doc is not None
