# -*- coding: utf-8 -*-
"""
桌签功能测试（方案七 P2-5：零覆盖 → 核心逻辑 100% + CLI 集成）。

覆盖：
- _calc_font_size：字号按名长动态调整（对齐 WPS 模板实测值）
- _format_name：两字名 2 空格、多字原样、带空格重格式化
- parse_name_list：分隔符 / 注释 / 空行
- _replace_placeholder_in_xml：XML 占位替换 + 字号调整
- _duplicate_body_for_combined：合并模式 body 复制
- CLI：--template 必填 / --placeholder 参数
"""
import sys
from pathlib import Path

import pytest

_ENGINE_DIR = Path(__file__).resolve().parent.parent / "engine"
if str(_ENGINE_DIR) not in sys.path:
    sys.path.insert(0, str(_ENGINE_DIR))

from table_sign_generator import (  # noqa: E402
    _calc_font_size, _format_name, parse_name_list,
    _replace_placeholder_in_xml, _duplicate_body_for_combined,
    _prepare_docx_from_template, generate_table_signs,
    generate_table_signs_combined,
)

# XML 命名空间（与 table_sign_generator 一致）
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _make_xml_with_placeholder(placeholder="Jose AI"):
    """构造包含占位文本的 document.xml 片段。"""
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>{placeholder}</w:t></w:r><w:r><w:rPr><w:sz w:val="260"/></w:rPr><w:t>other</w:t></w:r></w:p>
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/></w:sectPr>
  </w:body>
</w:document>""".encode('utf-8')


# ---------------------------------------------------------------------------
#  核心逻辑测试
# ---------------------------------------------------------------------------

def test_calc_font_size_2chars():
    """2 字名 → sz=312 (156pt)，对齐 WPS 模板。"""
    assert _calc_font_size(2) == 312
    assert _calc_font_size(1) == 312  # 1 字名同 2 字


def test_calc_font_size_3chars():
    """3 字名 → sz=240 (120pt)。"""
    assert _calc_font_size(3) == 240


def test_calc_font_size_4chars():
    """4 字名 → sz=200 (100pt)。"""
    assert _calc_font_size(4) == 200


def test_calc_font_size_5chars():
    """5 字及以上 → sz=160 (80pt)。"""
    assert _calc_font_size(5) == 160
    assert _calc_font_size(6) == 160


def test_format_name_2chars():
    """"张三" → "张  三"（2 个空格）。"""
    assert _format_name("张三") == "张  三"


def test_format_name_3chars():
    """三字名保持原样。"""
    assert _format_name("张三四") == "张三四"


def test_format_name_with_spaces():
    """带空格的名字（"张 三"）重格式化 → "张  三"。"""
    assert _format_name("张 三") == "张  三"


def test_parse_name_list():
    """基本名单解析：每行一人。"""
    names = parse_name_list("张三\n李四\n王五")
    assert names == ["张三", "李四", "王五"]


def test_parse_name_list_delimiters():
    """逗号/顿号/空格分隔。"""
    names = parse_name_list("张三,李四、王五 赵六")
    assert names == ["张三", "李四", "王五", "赵六"]


def test_parse_name_list_comments():
    """# 注释行和空行忽略。"""
    names = parse_name_list("# 会议名单\n张三\n\n李四")
    assert names == ["张三", "李四"]


def test_replace_placeholder_in_xml():
    """XML 占位替换 + 字号调整（2 字名 → sz=312）。"""
    xml_bytes = _make_xml_with_placeholder("Jose AI")
    result = _replace_placeholder_in_xml(xml_bytes, "Jose AI", "张三")
    assert "张" in result.decode('utf-8')
    assert "Jose AI" not in result.decode('utf-8')
    # 字号已调整为 312（156pt）
    assert 'w:val="312"' in result.decode('utf-8')


def test_replace_placeholder_custom():
    """自定义占位符替换。"""
    xml_bytes = _make_xml_with_placeholder("NAME_HERE")
    result = _replace_placeholder_in_xml(xml_bytes, "NAME_HERE", "李四")
    decoded = result.decode('utf-8')
    # 注意：两字名经 _format_name 插入 2 空格 → "李  四"
    assert "李  四" in decoded
    assert "NAME_HERE" not in decoded


def test_duplicate_body_combined():
    """合并模式：body 复制 N 次并替换占位符。"""
    xml_bytes = _make_xml_with_placeholder()
    result = _duplicate_body_for_combined(xml_bytes, ["张三", "李四"])
    text = result.decode('utf-8')
    # 两字名经 _format_name 插入 2 空格 → "张  三"/"李  四"
    assert "张  三" in text and "李  四" in text
    # N 人 → N-1 个分页符
    assert text.count('w:type="page"') == 1


def test_duplicate_body_combined_custom_placeholder():
    """合并模式：自定义占位符。"""
    xml_bytes = _make_xml_with_placeholder("PLACE")
    result = _duplicate_body_for_combined(xml_bytes, ["张三"], placeholder="PLACE")
    decoded = result.decode('utf-8')
    assert "张  三" in decoded
    assert "PLACE" not in decoded


def test_default_template_exists():
    """方案六：内置默认模板存在。"""
    from table_sign_template import DEFAULT_TEMPLATE_PATH
    assert DEFAULT_TEMPLATE_PATH.exists(), "内置默认模板应存在"


def test_generate_with_default_template(tmp_path):
    """使用内置默认模板生成桌签（零配置路径）。"""
    from table_sign_generator import DEFAULT_TEMPLATE
    assert DEFAULT_TEMPLATE is not None and DEFAULT_TEMPLATE.exists()
    files = generate_table_signs(["张三"], tmp_path, prefix="座签",
                                 template_path=DEFAULT_TEMPLATE)
    assert len(files) == 1
    assert files[0].exists()
    assert "张三" in files[0].name


def test_default_template_layout_matches_sample():
    """默认模板排版参数与示例模板（E:/座签制作/示新座签.docx）一致。

    校验：A4 页面/页边距/表格 2 行 1 列/表格宽 10838 twips/行高 9.65cm/
    边框 single sz=4 color=auto/华文新魏 156pt 不加粗。
    """
    from table_sign_template import build_default_template, DEFAULT_TEMPLATE_PATH
    from docx import Document

    build_default_template()  # 确保按当前代码重新生成
    doc = Document(str(DEFAULT_TEMPLATE_PATH))
    s = doc.sections[0]
    t = doc.tables[0]
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tblW = t._tbl.tblPr.find(W + 'tblW')
    borders = t._tbl.tblPr.find(W + 'tblBorders')
    trH = t.rows[0]._tr.trPr.find(W + 'trHeight')
    run = t.cell(0, 0).paragraphs[0].runs[0]

    # 页面：A4 纵向 + 边距（示例实测）
    assert abs(s.page_width.cm - 20.99) < 0.05
    assert abs(s.page_height.cm - 29.70) < 0.05
    assert abs(s.top_margin.cm - 2.69) < 0.05 and abs(s.bottom_margin.cm - 2.69) < 0.05
    assert abs(s.left_margin.cm - 2.79) < 0.05 and abs(s.right_margin.cm - 2.79) < 0.05

    # 表格：2 行 1 列（正反各 1 行）、宽 10838 twips、行高 5472 twips atLeast
    assert len(t.rows) == 2 and len(t.columns) == 1
    assert tblW is not None and tblW.get(W + 'w') == '10838'
    assert trH is not None and abs(int(trH.get(W + 'val')) - 5472) <= 1
    assert trH.get(W + 'hRule') == 'atLeast'

    # 边框：single sz=4 color=auto
    assert borders is not None
    b = borders[0]
    assert b.get(W + 'val') == 'single' and b.get(W + 'sz') == '4' and b.get(W + 'color') == 'auto'

    # 字体字号：华文新魏 156pt 不加粗，占位符两行均有
    assert run.font.name == '华文新魏'
    assert run.font.size.pt == 156
    assert run.font.bold is False
    assert 'Jose AI' in t.cell(0, 0).paragraphs[0].text
    assert 'Jose AI' in t.cell(1, 0).paragraphs[0].text


def test_generate_combined_with_default_template(tmp_path):
    """使用内置默认模板生成合并桌签。"""
    from table_sign_generator import DEFAULT_TEMPLATE
    out = tmp_path / "combined.docx"
    result = generate_table_signs_combined(["张三", "李四"], out,
                                           template_path=DEFAULT_TEMPLATE)
    assert result.exists()
    assert out.stat().st_size > 0


def test_missing_template_raises(tmp_path):
    """缺少模板时报 ValueError。"""
    from table_sign_generator import generate_table_signs
    with pytest.raises((ValueError, FileNotFoundError)):
        generate_table_signs(["张三"], tmp_path, template_path=Path("不存在的模板.dotx"))


# ---------------------------------------------------------------------------
#  CLI 集成测试
# ---------------------------------------------------------------------------

def test_table_signs_cli_requires_template():
    """P0-7：--template 已改为非必填（内置默认模板兜底），缺省不再报 argparse 错误。"""
    import subprocess
    r = subprocess.run([sys.executable, "-m", "gongwen", "table-signs", "--help"],
                       capture_output=True, text=True, encoding="utf-8", timeout=60)
    assert r.returncode == 0
    assert "--template" in (r.stdout or "")
    assert "默认使用内置模板" in (r.stdout or "")  # 帮助文本提示零配置


def test_table_signs_cli_placeholder_registered():
    """--placeholder 参数已注册（help 显示；utf-8 编码防 Windows GBK 解码失败）。"""
    import subprocess
    r = subprocess.run([sys.executable, "-m", "gongwen", "table-signs", "--help"],
                       capture_output=True, text=True, encoding="utf-8", timeout=60)
    assert r.returncode == 0
    assert "--placeholder" in (r.stdout or "")
    assert "--template" in (r.stdout or "")
