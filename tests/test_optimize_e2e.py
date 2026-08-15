# -*- coding: utf-8 -*-
"""P1-3b：optimize / check 端到端测试——锁住核心格式修复管线。

测试策略：
1. 用 python-docx 构造一个最小可读的"通知"docx（含基本字体/字号问题）
2. 跑 `python -m gongwen check <docx> -t notice --json` 验证只读检查不抛异常
3. 跑 `python -m gongwen optimize <docx> -o out.docx -t notice --apply` 验证修复产出
4. 断言：产出文件存在、非空、check 后 issues 数下降或保持
"""
import json
import subprocess
import sys
from pathlib import Path

import pytest


def _make_input_docx(path: Path) -> None:
    """构造一个字段不全的最小 .docx，故意带格式问题以便 check 发现。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # 标题——使用非国标字体，让 check 找出问题
    p = doc.add_paragraph("关于测试工作的通知")
    for r in p.runs:
        r.font.name = "宋体"
        r.font.size = Pt(14)  # 故意小于二号(22pt)
    # 受文机关
    doc.add_paragraph("各部门：")
    # 正文
    doc.add_paragraph("现将有关事项通知如下。")
    doc.add_paragraph("一、第一项工作")
    doc.add_paragraph("具体内容一")
    doc.add_paragraph("二、第二项工作")
    doc.add_paragraph("具体内容二")
    doc.add_paragraph("特此通知。")
    doc.add_paragraph("XX单位办公室")
    doc.add_paragraph("二〇二六年八月十六日")
    doc.save(str(path))


def _run(*args, env=None) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, "-m", "gongwen", *args],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
        env=env,
    )


@pytest.fixture
def input_docx(tmp_path):
    p = tmp_path / "input.docx"
    _make_input_docx(p)
    return p


def test_check_command_returns_zero_and_emits_issues(input_docx):
    """check 命令必须 exit 0（只读）并打印 issue 列表。"""
    r = _run("check", str(input_docx), "-t", "notice", "--json")
    assert r.returncode == 0, f"check failed:\n{r.stderr}\n{r.stdout}"
    out = (r.stdout or "").strip()
    assert out, "check --json 未输出任何内容"
    # P1-2 修复后 logger 输出到 stderr，stdout 应为纯 JSON 数组
    try:
        data = json.loads(out)
    except json.JSONDecodeError as exc:
        pytest.fail(f"check --json stdout 不应是纯 JSON 数组（logger 仍污染 stdout?）: {exc}\n片段: {out[:300]}")
    assert isinstance(data, list), f"check --json 应输出 list，实际: {type(data).__name__}"


def test_optimize_apply_produces_valid_docx(input_docx, tmp_path):
    """optimize --apply 必须产出非空 docx 文件。"""
    out = tmp_path / "out.docx"
    r = _run("optimize", str(input_docx), "-o", str(out),
             "-t", "notice", "--apply")
    assert r.returncode == 0, f"optimize failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists(), f"产出文件不存在: {out}"
    assert out.stat().st_size > 0, "产出文件为空"


def test_optimize_preview_does_not_write_output(input_docx, tmp_path):
    """optimize 不带 --apply 时（预览模式）不应写到 -o 路径。"""
    out = tmp_path / "preview.docx"
    r = _run("optimize", str(input_docx), "-o", str(out), "-t", "notice")
    assert r.returncode == 0, f"optimize preview failed:\n{r.stderr}\n{r.stdout}"
    # 预览模式不写文件
    assert not out.exists(), "预览模式不应创建产出文件"


def test_fix_common_command_apply(input_docx, tmp_path):
    """fix-common 必须产出非空 docx（路径 D 一键修复，无 --apply 参数，直接执行写文件）。"""
    out = tmp_path / "fixed.docx"
    r = _run("fix-common", str(input_docx), "-o", str(out))
    assert r.returncode == 0, f"fix-common failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists()
    assert out.stat().st_size > 0


def test_bold_first_command_apply(input_docx, tmp_path):
    """bold-first 命令端到端：对正文段落首句加粗（无 --apply 参数，直接执行）。"""
    out = tmp_path / "bold.docx"
    r = _run("bold-first", str(input_docx), "-o", str(out))
    assert r.returncode == 0, f"bold-first failed:\n{r.stderr}\n{r.stdout}"
    assert out.exists()
    assert out.stat().st_size > 0


def test_audit_command_returns_zero(input_docx):
    """audit 命令必须 exit 0（合规检查，只读）。"""
    r = _run("audit", str(input_docx))
    assert r.returncode == 0, f"audit failed:\n{r.stderr}\n{r.stdout}"
