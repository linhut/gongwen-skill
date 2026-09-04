# -*- coding: utf-8 -*-
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# https://github.com/linhut/gongwen-skill
# Licensed under the MIT License. See the LICENSE file for details.
#
# 本模块仅服务于 gongwen 的 md2docx 子命令。
"""md2docx 渲染器：直接用 python-docx 按 GB/T 9704 生成公文初稿。

既定方案（改造说明）
====================
原 cmd_md2docx 走「DocumentModel → generate_docx」通用管线：
  1) 该管线跨模块引用（如 engine.core.document.parser 等），存在
     「技能版本不一致」导致的导入错误风险；
  2) generate_docx 的 docDefaults 未显式设置字号，正文回退为 11pt，
     不符合 GB/T 9704「正文用三号仿宋_GB2312（16pt）」；
  3) 初稿本身并不合规，需完全依赖后续 optimize 兜底。

本模块改为用 python-docx 直接构建初稿：页面 A4 + 省筹委会页边距、大标题二号
方正小标宋、黑体一级标题、楷体二级标题、仿宋 16pt 正文、落款/日期、AI 声明段。
初稿按 rules 配置逐项排版，输出后仍由外部 optimize 做合规兜底修复。
"""
from __future__ import annotations

import logging
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Cm, Pt, RGBColor

from engine.core.document.font_utils import set_run_font, BODY_FONT, LATIN_FONT

_logger = logging.getLogger(__name__)

# V2.3：导语/过渡词开头 → 不是主送机关/称呼段（与 modifier.detect_paragraph_type 的
# _INTRODUCTION_RE/_TRANSITION_RE 逻辑对齐，避免"为深入贯彻落实…通知如下："短导语被
# 误判为称呼段而丢失首行缩进）
_SALUTATION_EXCLUDE_RE = re.compile(
    r'^\s*(按照|根据|遵照|依据|为了|为贯彻|为落实|为认真|为深入|为切实|为全面|'
    r'经|据|奉|针对|基于|鉴于|综上|为此|对此|结合|围绕|因此|故|由此可见|从上述)'
)

# ---------------------------------------------------------------------------
#  小工具
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _to_pt(value: Any, default: float = 16.0) -> float:
    """把 '16pt' / 16 / '16磅' 统一转成磅值。"""
    if value is None:
        return default
    if isinstance(value, bool):
        return default
    if isinstance(value, (int, float)):
        return float(value)
    m = re.match(r"^\s*([\d.]+)\s*(?:pt|磅)?\s*$", str(value))
    if m:
        return float(m.group(1))
    return default


def _align(align: str | None):
    return _ALIGN_MAP.get((align or "").lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)


# 中文公文编号标题模式（与 engine parser 的内容信号启发式对齐）
_H1_RE = re.compile(r"^[一二三四五六七八九十]+、[^、].{0,30}$")
_H2_RE = re.compile(r"^（[一二三四五六七八九十]+）.{0,40}$")
_H3_RE = re.compile(r"^\d+[.、].{0,40}$")
_H4_RE = re.compile(r"^（\d+）.{0,40}$")


def _detect_heading_level(text: str) -> int | None:
    """对未用 # 标记的正文段落做标题层级内容检测。

    返回 1/2/3/4 标题层级；若文本过长或不像标题则返回 None。
    与 parser._detect_heading_heuristic 的"内容信号"路径保持一致，
    使 md2docx 初稿即使不写 # 也能套用黑体/楷体等标题格式。
    """
    t = (text or "").strip()
    if not t or len(t) > 40:
        return None
    if _H1_RE.match(t):
        return 1
    if _H2_RE.match(t):
        return 2
    if _H3_RE.match(t):
        return 3
    if _H4_RE.match(t):
        return 4
    return None


def _set_exact_line_spacing(para, pt: float) -> None:
    """固定行距（GB/T 9704 统一 33 磅）。"""
    pf = para.paragraph_format
    pf.line_spacing = Pt(max(6, min(200, pt)))
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _apply_doc_defaults(doc: Document, body_pt: float = 16.0) -> None:
    """设置文档级默认字体与字号，避免正文回退 11pt。

    docDefaults 需同时写入 rFonts（中文字体）与 sz（字号，半磅值）。
    """
    try:
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_element.insert(0, doc_defaults)

        rPrDefault = doc_defaults.find(qn("w:rPrDefault"))
        if rPrDefault is None:
            rPrDefault = OxmlElement("w:rPrDefault")
            doc_defaults.append(rPrDefault)
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPrDefault.append(rPr)

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), LATIN_FONT)
        rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        rFonts.set(qn("w:eastAsia"), BODY_FONT)
        rFonts.set(qn("w:cs"), LATIN_FONT)

        # 字号（半磅）：三号 = 16pt → 32
        sz = rPr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rPr.append(sz)
        sz.set(qn("w:val"), str(int(body_pt * 2)))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(body_pt * 2)))
    except Exception as e:
        # 失败不影响主流程，仍按逐段显式格式输出
        _logger.debug("设置段落默认字体失败，按逐段显式格式输出: %s", e)


def _apply_section(doc: Document, page_setup) -> None:
    """应用页面设置（A4 + 页边距 + 页眉/页脚距离）。"""
    if not doc.sections:
        return
    sec = doc.sections[0]
    w = getattr(page_setup, "paper_width_mm", None) or 210
    h = getattr(page_setup, "paper_height_mm", None) or 297
    sec.page_width = Mm(w)
    sec.page_height = Mm(h)
    sec.top_margin = Mm(getattr(page_setup, "margin_top_mm", None) or 28.0)
    sec.bottom_margin = Mm(getattr(page_setup, "margin_bottom_mm", None) or 28.0)
    sec.left_margin = Mm(getattr(page_setup, "margin_left_mm", None) or 27.0)
    sec.right_margin = Mm(getattr(page_setup, "margin_right_mm", None) or 27.0)
    sec.header_distance = Cm(getattr(page_setup, "header_distance_cm", None) or 1.5)
    sec.footer_distance = Cm(getattr(page_setup, "footer_distance_cm", None) or 2.3)


# ---------------------------------------------------------------------------
#  段落排版（按角色 / 标题层级套用 GB/T 9704 默认）
# ---------------------------------------------------------------------------

def _role_style(rules: dict, para) -> dict:
    """根据段落角色/标题层级返回默认排版字典。"""
    cfg = rules or {}
    if getattr(para, "is_heading", False) and para.heading_level == 0:
        c = cfg.get("doc_title", {})
        return {
            "font": c.get("font", "方正小标宋简体"),
            "size": _to_pt(c.get("size", 22)),
            "bold": bool(c.get("bold", False)),
            "align": c.get("align", "center"),
            "indent_em": 0,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if getattr(para, "is_heading", False) and para.heading_level == 1:
        c = cfg.get("heading_1", {})
        return {
            "font": c.get("font", "黑体"),
            "size": _to_pt(c.get("size", 16)),
            "bold": bool(c.get("bold", False)),
            "align": "left",
            "indent_em": 2,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if getattr(para, "is_heading", False) and para.heading_level == 2:
        c = cfg.get("heading_2", {})
        return {
            "font": c.get("font", "楷体_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": bool(c.get("bold", False)),
            "align": "left",
            "indent_em": 2,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if getattr(para, "is_heading", False) and para.heading_level == 3:
        c = cfg.get("heading_3", {})
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": bool(c.get("bold", True)),
            "align": "left",
            "indent_em": 2,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }

    role = para.role
    if role == "recipient":
        c = cfg.get("salutation", {})
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": False,
            "align": c.get("align", "left"),
            "indent_em": 0,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if role == "salutation":
        c = cfg.get("salutation", {})
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": False,
            "align": c.get("align", "left"),
            "indent_em": 0,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if role == "signature":
        c = cfg.get("signature", {})
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 18), 18),
            "bold": bool(c.get("bold", False)),
            "align": c.get("align", "center"),
            "indent_em": 0,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if role == "date":
        c = cfg.get("date", {})
        # P2-19 修复：GB/T 9704 成文日期"右空四字"——日期右对齐且右侧留 4 字空位
        # 规则值可能是数字 4 或字符串 "4em"，统一解析为 em 数值
        _ri_raw = c.get("right_indent", 4)
        try:
            _ri_em = float(str(_ri_raw).replace("em", "").strip())
        except (ValueError, TypeError):
            _ri_em = 4.0
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": False,
            "align": c.get("align", "right"),
            "indent_em": 0,
            "right_indent_em": _ri_em,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }
    if role == "attachment":
        c = cfg.get("body", {})
        return {
            "font": c.get("font", "仿宋_GB2312"),
            "size": _to_pt(c.get("size", 16)),
            "bold": False,
            "align": c.get("align", "justify"),
            "indent_em": 2,
            "line": _to_pt(c.get("line_spacing", 33), 33),
        }

    # 默认正文
    c = cfg.get("body", {})
    return {
        "font": c.get("font", "仿宋_GB2312"),
        "size": _to_pt(c.get("size", 16)),
        "bold": False,
        "align": c.get("align", "justify"),
        "indent_em": 2,
        "line": _to_pt(c.get("line_spacing", 33), 33),
    }


def _add_paragraph(doc: Document, para, rules: dict):
    """把模型段落写入 docx（返回 python-docx Paragraph 对象）。"""
    style = _role_style(rules, para)

    # 未用 # 标记的编号标题内容检测（一、/（一）/1./（1））：
    # 即便 markdown 未写 ##，也套用黑体/楷体等标题格式，与 check 的启发式一致
    if not getattr(para, "is_heading", False) and para.role in (None, "body", "recipient"):
        _detected = _detect_heading_level(para.text)
        if _detected == 1:
            _c = (rules or {}).get("heading_1", {})
            style = {"font": _c.get("font", "黑体"), "size": _to_pt(_c.get("size", 16)),
                     "bold": False, "align": "left", "indent_em": 2,
                     "line": _to_pt(_c.get("line_spacing", 33), 33)}
        elif _detected == 2:
            _c = (rules or {}).get("heading_2", {})
            style = {"font": _c.get("font", "楷体_GB2312"), "size": _to_pt(_c.get("size", 16)),
                     "bold": False, "align": "left", "indent_em": 2,
                     "line": _to_pt(_c.get("line_spacing", 33), 33)}
        elif _detected == 3:
            _c = (rules or {}).get("heading_3", {})
            style = {"font": _c.get("font", "仿宋_GB2312"), "size": _to_pt(_c.get("size", 16)),
                     "bold": True, "align": "left", "indent_em": 2,
                     "line": _to_pt(_c.get("line_spacing", 33), 33)}
        elif _detected == 4:
            _c = (rules or {}).get("heading_4", {})
            style = {"font": _c.get("font", "仿宋_GB2312"), "size": _to_pt(_c.get("size", 16)),
                     "bold": False, "align": "left", "indent_em": 2,
                     "line": _to_pt(_c.get("line_spacing", 33), 33)}

    p = doc.add_paragraph()
    pf = p.paragraph_format

    # 段前分页标记（--- 附件分页）：该段从新页开始
    if getattr(para, "page_break", False):
        pf.page_break_before = True

    # 主送机关/称呼段：顶格左对齐、无首行缩进（GB/T 9704）
    # 通过 role 或文本特征（短文本以 ：/: 结尾）识别，兼容 recipient 放正文的草稿。
    # V2.3 修复：排除导语/过渡词开头的短段——"为深入贯彻落实…通知如下："等导语段
    # 也以冒号结尾且常不超过 50 字，会被误判为主送机关而左对齐、丢失首行缩进。
    _txt = (para.text or "").strip()
    _is_salutation = para.role in ("recipient", "salutation") or (
        para.role in (None, "body")
        and len(_txt) <= 50
        and _txt.endswith(("：", ":"))
        and not _SALUTATION_EXCLUDE_RE.match(_txt)
    )

    # 对齐：主送机关/称呼段强制左对齐；其余优先模型显式值
    if _is_salutation:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
    elif para.format.alignment:
        p.alignment = _align(para.format.alignment)
        # 首行缩进：2em = 2 × 字号
        if para.format.first_line_indent_pt is not None and para.format.first_line_indent_pt > 0:
            pf.first_line_indent = Pt(para.format.first_line_indent_pt)
        elif style["indent_em"] > 0 and para.text.strip():
            pf.first_line_indent = Pt(style["indent_em"] * style["size"])
        else:
            pf.first_line_indent = Pt(0)
    else:
        p.alignment = _align(style["align"])
        if style["indent_em"] > 0 and para.text.strip():
            pf.first_line_indent = Pt(style["indent_em"] * style["size"])
        else:
            pf.first_line_indent = Pt(0)

    # 左缩进（列表等）
    if para.format.left_indent_pt is not None:
        pf.left_indent = Pt(para.format.left_indent_pt)

    # P2-19：右缩进（成文日期"右空四字"等）
    _right_em = style.get("right_indent_em")
    if _right_em:
        pf.right_indent = Pt(float(_right_em) * float(style["size"]))

    # 行距
    if para.format.line_spacing_pt is not None:
        _set_exact_line_spacing(p, para.format.line_spacing_pt)
    else:
        _set_exact_line_spacing(p, style["line"])

    # runs
    runs = para.runs if para.runs else []
    if not runs and para.text:
        # 无 run 时按段落默认创建单 run
        run = p.add_run(para.text)
        set_run_font(run, style["font"])
        run.font.size = Pt(style["size"])
        if style["bold"]:
            run.font.bold = True
    else:
        for r in runs:
            if not r.text:
                continue
            run = p.add_run(r.text)
            fmt = r.format
            font = fmt.font_name or style["font"]
            size = fmt.font_size_pt if fmt.font_size_pt is not None else style["size"]
            set_run_font(run, font)
            run.font.size = Pt(size)
            if fmt.bold is not None:
                run.font.bold = fmt.bold
            elif style["bold"]:
                run.font.bold = True
            if fmt.italic is not None:
                run.font.italic = fmt.italic
            if fmt.strikethrough is not None:
                # True 设置删除线 / False 显式清除已有删除线（避免 strikethrough=False 时旧删除线残留）
                run.font.strike = fmt.strikethrough
            if fmt.color:
                try:
                    rgb = str(fmt.color).lstrip("#")
                    run.font.color.rgb = RGBColor(int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))
                except Exception as e:
                    _logger.debug("颜色解析失败，跳过该 run 的字体颜色: %s", e)
    return p


def _set_cell_shading(cell, fill: str | None) -> None:
    """给表格单元格写入 w:shd 底色（如浅蓝灰 D9E2F3）。"""
    if not fill:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_table_cell_margins(table, margin: dict | None) -> None:
    """写入 w:tblCellMar（单元格边距，twips）。"""
    if not margin:
        return
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = table._tbl._add_tblPr()
    cell_mar = tblPr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = tblPr.makeelement(qn("w:tblCellMar"), {})
        tblPr.append(cell_mar)
    for edge in ("top", "left", "bottom", "right"):
        val = margin.get(edge)
        if val is None:
            continue
        el = cell_mar.find(qn("w:" + edge))
        if el is None:
            el = cell_mar.makeelement(qn("w:" + edge), {})
            cell_mar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")


# V2.3：md2docx 表格样式与 _common.yaml table 配置对齐——应用表头底色、
# 单元格边距、数字列右对齐（此前仅 Table Grid 边框+字体，配置块未被读取）
_NUM_CELL_RE = re.compile(r"^[\d.,%‰]+$")


def _add_table(doc: Document, tbl, anchor_elem, rules=None) -> None:
    """创建 Word 表格并插到锚点段落之后。"""
    if tbl.rows <= 0 or tbl.cols <= 0:
        return
    table = doc.add_table(rows=tbl.rows, cols=tbl.cols)
    table.style = "Table Grid"
    # 把表格移动到锚点元素之后
    if anchor_elem is not None:
        anchor_elem.addnext(table._tbl)

    # V2.3：从规则读取表格样式（表头底色、单元格边距、数字智能对齐）
    tcfg = (rules or {}).get("table", {}) or {}
    _hdr_cfg = tcfg.get("header", {}) or {}
    _fill = _hdr_cfg.get("fill")
    _margin = tcfg.get("cell_margin")
    if _margin:
        _set_table_cell_margins(table, _margin)

    cells_by = {}
    for c in tbl.cells:
        cells_by[(c.row, c.col)] = c

    for row in range(tbl.rows):
        for col in range(tbl.cols):
            cell_model = cells_by.get((row, col))
            text = cell_model.text if cell_model else ""
            cell = table.cell(row, col)
            # 清空默认段落，写入一个 run
            cell.text = ""
            para = cell.paragraphs[0]
            is_header = row == 0
            if is_header and _fill:
                _set_cell_shading(cell, _fill)
            run = para.add_run(text)
            if is_header:
                set_run_font(run, "黑体")
                run.font.size = Pt(12)
                run.font.bold = True
                para.alignment = _align("center")
            else:
                set_run_font(run, "仿宋_GB2312")
                run.font.size = Pt(12)
                # V2.3：数字列右对齐（与 optimize _smart_align_cell 一致）
                if _NUM_CELL_RE.match((text or "").strip()):
                    para.alignment = _align("right")
                else:
                    para.alignment = _align("left")


def _add_ai_declaration(doc: Document) -> None:
    """文档末尾追加 AI 声明段（楷体 9pt 居中，pStyle 标记为 Annotation）。"""
    ai_text = "（内容由GongWen-skill-AI生成，仅供参考）"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_elem = para._element
    pPr = p_elem.get_or_add_pPr()
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), "Annotation")
    run = para.add_run(ai_text)
    set_run_font(run, "楷体_GB2312")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
#  主入口
# ---------------------------------------------------------------------------

def render_model_to_docx(model, output_path, rules: dict | None = None,
                         no_ai_declaration: bool = False):
    """用 python-docx 直接按 GB/T 9704 把 DocumentModel 渲染为 .docx。

    替代 generate_docx（不依赖 engine.core.document.parser 等管线模块，
    直接显式排版，保证初稿即合规）。
    """
    import os
    doc = Document()

    _apply_doc_defaults(doc, body_pt=16.0)
    _apply_section(doc, model.page_setup)

    # 清理源文中的 AI 声明段落（如 "*（内容由AI生成，仅供参考）*"），
    # 避免正文中遗留重复声明，统一由 _add_ai_declaration 追加格式化版本
    _AI_PATTERNS = ("由GongWen-skill-AI生成", "由AI生成", "仅供参考")
    for pi in range(len(model.paragraphs) - 1, -1, -1):
        txt = (model.paragraphs[pi].text or "").replace("*", "").strip()
        if any(m in txt for m in _AI_PATTERNS):
            del model.paragraphs[pi]

    # 表格按锚点段落索引分组
    table_map: dict[int, list] = {}
    for t in model.tables:
        table_map.setdefault(t.insert_after_index, []).append(t)

    # 渲染段落；表格插到其锚点段落之后
    last_elem = None
    for i, para in enumerate(model.paragraphs):
        p = _add_paragraph(doc, para, rules)
        last_elem = p._element
        for t in table_map.get(i, []):
            _add_table(doc, t, last_elem, rules)

    # 锚点为 -1（文档开头）的表格
    for t in table_map.get(-1, []):
        if doc.paragraphs:
            _add_table(doc, t, doc.paragraphs[0]._element, rules)
        else:
            _add_table(doc, t, None, rules)

    # 锚点超出末尾（最后一个段落之后）
    for t in table_map.get(len(model.paragraphs), []):
        _add_table(doc, t, last_elem, rules)

    # AI 声明
    if not no_ai_declaration:
        _add_ai_declaration(doc)

    if isinstance(output_path, str):
        out = output_path
        parent = os.path.dirname(os.path.abspath(out))
        if parent:
            os.makedirs(parent, exist_ok=True)
    else:
        out = output_path

    # 统一文档作者（用户要求：生成文档的作者写为 Jose AI）
    try:
        doc.core_properties.author = "Jose AI"
    except Exception as e:
        _logger.debug("设置文档作者失败: %s", e)
    doc.save(out)
    return output_path
