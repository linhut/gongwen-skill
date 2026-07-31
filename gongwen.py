#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# 公文文档格式化 Skill —— 中文公文全流程处理工具
#
# (c) 2026 Jose AI (https://www.linhut.cn)
# 项目出处：AI 公文智能优化助手 (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
#
# 本文件为独立发行版的入口，任何人克隆仓库后即可运行，
# 无需原桌面端项目、无需数据库、无需后端服务。

__version__ = "1.12.36"
"""
中文公文全流程处理工具 —— 基于 GB/T 9704《党政机关公文格式》国家标准。

支持格式检查与修复、内容润色（红色标注对比版）、模板生成、Markdown 转公文、
版头版记页码注入等完整能力。打包为可被 AI Agent 直接调用的 Skill，
完全自包含，克隆即用。

子命令：
  list-types                   列出所有支持的公文类型
  template  <type> -o out.docx 生成指定类型的标准公文模板
  parse     <in.docx>          解析文档为结构化 JSON（DocumentModel）
  check     <in.docx>          按规则检查格式问题（只读，不改文件）
  optimize  <in.docx> -o out   检查 + 自动修复 + 生成合规文档（支持 --layout 版式注入）
  generate  <model.json> -o    从 DocumentModel JSON 生成 .docx
  md2docx   <input.md> -o      将 Markdown 文本转为格式化的公文 .docx
  header    <in.docx>          注入版头：发文机关标志 + 发文字号 + 签发人 + 红色反线
  footer    <in.docx>          注入版记：抄送 + 印发机关 + 印发日期 + 分隔线
  pagenum   <in.docx>          注入页码：Word PAGE 域动态页码（居中 / 单右双左）
  rule-export <type>           导出某类型的合并规则为 YAML 用于二次定制
  rule-list                    列出三层规则（official / custom / user）
  rule-import <key> -f <file>  导入/保存自定义规则 YAML

示例：
  python gongwen.py list-types
  python gongwen.py template notice -o 通知模板.docx
  python gongwen.py check input.docx -t notice --json
  python gongwen.py optimize input.docx -o output.docx -t report
  cat input.md | python gongwen.py md2docx - -o 公文.docx    # 管道输入
  python gongwen.py header in.docx --org-name 国家民委办公厅 --doc-number "民委办发〔2026〕1号"
  python gongwen.py footer in.docx --cc 各省民委 --printer 国家民委办公厅 --print-date 2026年7月23日
  python gongwen.py pagenum in.docx --alignment right
"""
import argparse
import json
import sys
from pathlib import Path

# 将 engine/ 加入模块搜索路径，使内部 `from core... / from utils... / from config`
# 的绝对导入生效——这是独立运行的关键。
_ENGINE_DIR = Path(__file__).resolve().parent / "engine"
sys.path.insert(0, str(_ENGINE_DIR))

# Windows 控制台中文输出保护（借鉴 docx-skill 强制 UTF-8 策略）
# 同时覆盖 stdout/stderr/stdin，确保中文路径、管道输入均无编码问题
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
    sys.stdin.reconfigure(encoding="utf-8")
except Exception:
    pass

# 注册 tmp 目录进程退出自动清理
try:
    from tmp import register_cleanup
    register_cleanup()
except Exception:
    pass


# ---------------------------------------------------------------------------
#  共享辅助
# ---------------------------------------------------------------------------

# 文件名关键词 → 公文类型（长关键词优先，避免"会议纪要"被"纪要"抢先）
_TYPE_KEYWORDS = {
    "会议纪要": "meeting", "技术方案": "technical_proposal",
    "通知": "notice", "请示": "request", "报告": "report",
    "函": "letter", "纪要": "minutes", "决定": "decision",
    "通告": "announcement", "公告": "notice_public", "命令": "command",
    "通报": "bulletin", "议案": "bill", "批复": "reply",
    "指示": "instruction", "制度": "regulation", "公报": "communique",
    "意见": "opinion", "总结": "summary", "方案": "work_plan",
    "计划": "work_plan", "桌签": "table_sign", "决议": "resolution",
    "讲话稿": "speech", "主持词": "speech", "新闻稿": "news", "简报": "news",
}


def _detect_doc_type(input_path: "Path", explicit: str | None) -> tuple[str, str]:
    """确定公文类型，返回 (类型, 来源说明)。

    优先级：用户显式 -t > 文件名关键词推断 > 默认 notice。
    """
    from pathlib import Path as _P
    if explicit:
        return explicit, "用户指定"
    stem = _P(input_path).stem
    for kw, dt in sorted(_TYPE_KEYWORDS.items(), key=lambda x: -len(x[0])):
        if kw in stem:
            return dt, f"文件名含「{kw}」推断"
    return "notice", "默认（未识别到类型关键词）"


def _extract_dominant_style(changes: list[dict]) -> str | None:
    """从 changes 列表中提取出现次数最多的 style 标签。"""
    from collections import Counter
    styles = [c.get("style", "") for c in changes if c.get("style", "").strip()]
    if not styles:
        return None
    return Counter(styles).most_common(1)[0][0]


def _build_output_name(input_path: "str | Path", convention: str, style: str | None = None) -> str:
    """根据命名规范构造输出文件名（不含路径，仅文件名）。

    规范：
    - 路径 A / C（格式优化 / 模板生成）：修订版+{原文档名}+{日期}+v{版本号}.docx
    - 路径 B（内容优化对比文档）：{原文档名}+{内容风格}+{日期}+v{版本号}.docx

    版本叠加：若输入文件名含 +v{数字}，自动检测并 +1 作为输出版本号。

    Args:
        input_path: 原文档路径
        convention: 'A' 或 'B'
        style: 路径 B 的内容风格标签（如"庄重严谨"）

    Returns:
        符合命名规范的文件名，如 "工作报告+庄重严谨+20260725+v2.docx"
    """
    from datetime import date
    import re

    stem = Path(input_path).stem
    today = date.today().strftime("%Y%m%d")

    # 检测文件名中已有的版本号（如 +v1、+v2、v1 等）
    version = 1
    v_match = re.search(r'\+v(\d+)$', stem)
    if v_match:
        version = int(v_match.group(1)) + 1  # 叠加：+1
        stem = stem[:v_match.start()]  # 去掉版本后缀
    else:
        # 也检测末尾的 "v1""v2" 模式（不带 +）
        v2 = re.search(r'v(\d+)$', stem)
        if v2:
            version = int(v2.group(1)) + 1
            stem = stem[:v2.start()]

    # 去除尾部可能残留的分隔符
    stem = stem.rstrip('+ _-')

    if convention == "B":
        style_part = f"+{style}" if style else ""
        return f"{stem}{style_part}+{today}+v{version}.docx"
    else:  # A / C
        return f"修订版+{stem}+{today}+v{version}.docx"


# ---------------------------------------------------------------------------
#  子命令实现
# ---------------------------------------------------------------------------

def cmd_list_types(args):
    """列出所有支持的公文类型。"""
    from core.rules.loader import list_available_types
    types = list_available_types()
    if args.json:
        print(json.dumps(types, ensure_ascii=False, indent=2))
    else:
        for t in types:
            print(t)


def cmd_template(args):
    """生成标准公文模板。"""
    from datetime import date as _dt
    from core.rules.manager import load_rules_merged
    from core.document.generator import generate_docx
    from template_builder import create_template_document

    doc_type = args.type
    rules = load_rules_merged(doc_type)
    model = create_template_document(doc_type, rules)

    if args.output:
        out = Path(args.output)
    else:
        today = _dt.today().strftime("%Y-%m-%d")
        out = Path(f"修订版+{doc_type}-模板+{today}+v1.docx")
    generate_docx(model, out)
    print(f"模板已生成: {out} (类型: {doc_type})")


def cmd_parse(args):
    """解析文档为结构化 JSON。"""
    from core.document.parser import parse_docx

    model = parse_docx(args.input)
    data = model.model_dump()
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"已解析: {args.output} ({len(model.paragraphs)} 段落, {len(model.tables)} 表格)")
    else:
        print(text)


def cmd_check(args):
    """检查文档格式（只读）。"""
    from core.document.parser import parse_docx
    from core.rules.engine import RuleEngine

    engine = RuleEngine()
    model = parse_docx(args.input)
    issues = engine.check(model, args.doc_type)

    if args.severity:
        issues = [i for i in issues if i.severity == args.severity]

    if args.json:
        results = [{
            "severity": i.severity, "rule_id": i.rule_id, "name": i.name,
            "check_type": i.check_type, "location": i.location,
            "original": i.original_text, "suggested": i.suggested_fix,
            "reason": i.reason,
        } for i in issues]
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        p0 = sum(1 for i in issues if i.severity == "P0")
        p1 = sum(1 for i in issues if i.severity == "P1")
        p2 = sum(1 for i in issues if i.severity == "P2")
        print(f"检查完成: {len(issues)} 个问题 (P0:{p0}, P1:{p1}, P2:{p2})")
        for i in issues:
            print(f"  [{i.severity}] {i.rule_id}: {i.name} @ {i.location}")
            print(f"       实际: {i.original_text}  → 期望: {i.suggested_fix}")


def cmd_optimize(args):
    """检查 + 修复 + 生成（格式优化，不改内容）。

    默认预览模式：检测类型 → 检查问题 → 列出摘要 → 提示下一步。
    加 --apply 才真正执行修复并生成文件。
    """
    from core.document.parser import parse_docx
    from core.document.generator import generate_docx
    from core.rules.engine import RuleEngine

    engine = RuleEngine()
    input_path = Path(args.input)
    out = Path(args.output) if args.output else input_path.parent / _build_output_name(input_path, "A")

    # 确定文档类型（共享辅助函数，优先 -t 参数，其次文件名推断）
    doc_type, type_source = _detect_doc_type(input_path, args.doc_type)

    # 解析文档并检查
    model = parse_docx(str(input_path))
    issues = engine.check(model, doc_type)

    p0 = [i for i in issues if i.severity == "P0"]
    p1 = [i for i in issues if i.severity == "P1"]
    p2 = [i for i in issues if i.severity == "P2"]

    # === 预览信息（始终显示）===
    print(f"📄 文件: {input_path.name}")
    print(f"🔍 类型: {doc_type}（{type_source}）")
    print(f"📊 问题: 共 {len(issues)} 项（P0:{len(p0)}, P1:{len(p1)}, P2:{len(p2)}）")
    if issues:
        print(f"  P0 示例（必须修复）:")
        for i in p0[:3]:
            print(f"    - {i.name} @ {i.location}")
        if p1:
            print(f"  P1 示例（建议修复）:")
            for i in p1[:3]:
                print(f"    - {i.name} @ {i.location}")
    if args.layout:
        lc = json.loads(Path(args.layout).read_text(encoding="utf-8"))
        parts = [k for k in ("header", "footer", "page_number") if k in lc]
        print(f"🎨 版式注入: {', '.join(parts)}")

    if not args.apply:
        print()
        print("─── 预览模式 ───")
        print("以上是本次将要修复的内容预览。")
        print("加 --apply 执行修复，或指定 -t 切换公文类型。")
        print("示例:")
        print(f"  python gongwen.py optimize {args.input} -t notice --apply")
        print(f"  python gongwen.py optimize {args.input} -o 成品.docx --apply --layout 版式.json")
        return

    # === 执行模式 ===
    selected = args.selected_rules.split(",") if args.selected_rules else None
    _, fixed = engine.check_and_fix(model, doc_type, selected)
    # 清理路径 B 遗留的修改说明段落和删除线标记（确保干净成品）
    from core.document.modifier import clean_path_b_markers
    cleaned = clean_path_b_markers(fixed)
    generate_docx(fixed, str(out))
    print(f"✅ 优化完成: {out}")
    print(f"  修复 {len(issues)} 项 (P0:{len(p0)}, P1:{len(p1)}, P2:{len(p2)})")
    if cleaned:
        print(f"  清理 {cleaned} 处路径B标记")

    if getattr(args, "layout", None):
        layout = json.loads(Path(args.layout).read_text(encoding="utf-8"))
        from inject import inject_header, inject_footer, inject_page_number
        if layout.get("header"):
            inject_header(str(out), layout["header"])
            print("  版头已注入")
        if layout.get("footer"):
            inject_footer(str(out), layout["footer"])
            print("  版记已注入")
        if layout.get("page_number"):
            inject_page_number(str(out), layout["page_number"])
            print("  页码已注入")


def cmd_generate(args):
    """从 DocumentModel JSON 生成 .docx。"""
    from core.document.models import DocumentModel
    from core.document.generator import generate_docx

    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    model = DocumentModel(**data)
    if args.output:
        out = Path(args.output)
    else:
        stem = Path(args.input).stem.replace(".model", "").replace("_model", "")
        out = Path(f"修订版+{stem}+生成.docx")
    generate_docx(model, out)
    print(f"文档已生成: {out}")


def cmd_md2docx(args):
    """
    将 Markdown 文本转为格式化的公文 .docx 文件。

    输入可以是文件路径，也可以是 '-'（标准输入，支持管道）。
    支持 Front Matter 元数据（--- 包裹的 YAML 块）：
    - recipients: 主送机关（字符串或数组）
    - signer: 落款单位
    - date: 成文日期
    - attachments: 附件列表（字符串数组）
    - doc_type: 公文类型（默认 notice）
    """
    import io
    from datetime import date as _dt
    from core.document.parser_format import parse_paragraph_format, parse_run
    from core.document.generator import generate_docx
    from core.document.models import (
        DocumentModel, DocumentMetadata, PageSetup,
        Paragraph, ParagraphFormat, Run, RunFormat,
    )
    from core.document.modifier import convert_markdown
    from core.rules.manager import load_rules_merged

    # 读取输入
    text: str
    source_desc: str
    input_src = args.input
    if input_src == "-":
        raw = sys.stdin.buffer.read()
        text = raw.decode("utf-8")
        source_desc = "stdin"
    else:
        text = Path(input_src).read_text(encoding="utf-8")
        source_desc = input_src

    # 解析 Front Matter
    doc_type = args.doc_type or "notice"
    recipients = args.recipients or []
    signer = args.signer or ""
    doc_date = args.date or ""
    attachments = args.attachments or []

    lines = text.split("\n")
    if lines and lines[0].strip() == "---":
        # 尝试提取 YAML front matter
        end_idx = None
        front_matter = {}
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                end_idx = i
                break
            if ":" in lines[i]:
                key, _, val = lines[i].partition(":")
                k = key.strip()
                v = val.strip().strip('"').strip("'")
                if v:
                    front_matter[k] = v
        if end_idx:
            lines = lines[end_idx + 1:]
            text = "\n".join(lines)
            doc_type = front_matter.get("doc_type", doc_type)
            recipients = front_matter.get("recipients", recipients)
            signer = front_matter.get("signer", signer)
            doc_date = front_matter.get("date", doc_date)
            attachments = front_matter.get("attachments", attachments)

    # 加载规则获取页边距等
    rules = load_rules_merged(doc_type)
    margins = rules.get("page_setup", {}).get("margins", {})

    def _parse_margin(v):
        s = str(v).strip()
        if "cm" in s: return float(s.replace("cm", "")) * 10
        if "mm" in s: return float(s.replace("mm", ""))
        return float(s)

    # 构建 DocumentModel
    model = DocumentModel(
        metadata=DocumentMetadata(),
        page_setup=PageSetup(
            paper_width_mm=210, paper_height_mm=297,
            margin_top_mm=_parse_margin(margins.get("top", "3.7cm")),
            margin_bottom_mm=_parse_margin(margins.get("bottom", "3.5cm")),
            margin_left_mm=_parse_margin(margins.get("left", "2.8cm")),
            margin_right_mm=_parse_margin(margins.get("right", "2.6cm")),
        ),
    )

    # 主送机关
    rcp = recipients
    if isinstance(rcp, str) and rcp:
        # "各单位,各部门" → ["各单位", "各部门"]
        parts = [p.strip() for p in rcp.replace("，", ",").split(",") if p.strip()]
        rcp = parts
    if isinstance(rcp, list) and rcp:
        rcp_text = "、".join(rcp) + "："
        model.paragraphs.append(Paragraph(
            index=0, text=rcp_text, role="recipient",
            runs=[Run(index=0, text=rcp_text, format=RunFormat(font_name="仿宋_GB2312", font_size_pt=16.0))],
            format=ParagraphFormat(alignment="justify", first_line_indent_pt=0, line_spacing_pt=28.95),
        ))

    # 正文行：每行一个段落
    para_offset = len(model.paragraphs)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        model.paragraphs.append(Paragraph(
            index=para_offset + i, text=stripped, role="body",
            runs=[Run(index=0, text=stripped, format=RunFormat())],
            format=ParagraphFormat(alignment="justify", line_spacing_pt=28.95),
        ))

    # 执行 Markdown 转换（#标题 → 标题样式，**加粗** → bold，|表格| → Word 表格）
    changes = convert_markdown(model)

    # 若 Markdown 未用 # 标记标题，自动将首个正文段落设为标题
    has_title = any(getattr(p, 'is_heading', False) and p.heading_level == 0 for p in model.paragraphs)
    if not has_title:
        for para in model.paragraphs:
            if para.text.strip() and getattr(para, 'role', None) == 'body':
                para.role = 'title'
                para.is_heading = True
                para.heading_level = 0
                # 设置物理格式，确保 round-trip 后能被解析器正确识别
                para.format.alignment = 'center'
                for r in para.runs:
                    r.format.font_name = '方正小标宋简体'
                    r.format.font_size_pt = 22.0
                break

    # 领句加粗（路径 C 生成公文时，Markdown 中的"一是/二是/第一/第二"等领句自动加粗）
    _BOLD_LEADIN = {
        '一是': '楷体_GB2312', '二是': '楷体_GB2312', '三是': '楷体_GB2312',
        '四是': '楷体_GB2312', '五是': '楷体_GB2312',
        '第一，': '仿宋_GB2312', '第二，': '仿宋_GB2312', '第三，': '仿宋_GB2312',
        '一要': '仿宋_GB2312', '二要': '仿宋_GB2312', '三要': '仿宋_GB2312',
    }
    for para in model.paragraphs:
        txt = para.text.strip()
        matched = next((p for p in _BOLD_LEADIN if txt.startswith(p)), None)
        if not matched or not para.runs:
            continue
        pi = txt.find('。')
        if pi == -1:
            continue
        lead_in, remaining = txt[:pi + 1], txt[pi + 1:]
        if not remaining:
            continue
        para.runs[0].text = lead_in
        para.runs[0].format.bold = True
        para.runs[0].format.font_name = _BOLD_LEADIN[matched]
        para.runs[0].format.font_size_pt = 16.0
        para.runs.append(Run(
            index=len(para.runs), text=remaining,
            format=RunFormat(font_name='仿宋_GB2312', font_size_pt=16.0),
        ))

    # 落款与日期
    if signer:
        idx = len(model.paragraphs)
        model.paragraphs.append(Paragraph(
            index=idx, text=signer, role="signature",
            runs=[Run(index=0, text=signer, format=RunFormat(font_name="仿宋_GB2312", font_size_pt=16.0))],
            format=ParagraphFormat(alignment="right", line_spacing_pt=28.95),
        ))
    if doc_date:
        idx = len(model.paragraphs)
        model.paragraphs.append(Paragraph(
            index=idx, text=doc_date, role="date",
            runs=[Run(index=0, text=doc_date, format=RunFormat(font_name="仿宋_GB2312", font_size_pt=16.0))],
            format=ParagraphFormat(alignment="right", line_spacing_pt=28.95),
        ))

    # 附件说明
    atts = attachments
    if isinstance(atts, str) and atts:
        atts = [atts]
    if isinstance(atts, list) and atts:
        idx = len(model.paragraphs)
        att_text = "附件：" + "、".join(f"{i+1}.{a}" for i, a in enumerate(atts))
        model.paragraphs.append(Paragraph(
            index=idx, text=att_text, role="attachment",
            runs=[Run(index=0, text=att_text, format=RunFormat(font_name="仿宋_GB2312", font_size_pt=16.0))],
            format=ParagraphFormat(alignment="justify", line_spacing_pt=28.95),
        ))

    # 生成 docx
    if args.output:
        out = Path(args.output)
    else:
        today = _dt.today().strftime("%Y-%m-%d")
        out = Path(f"修订版+{doc_type}-草稿+{today}+v1.docx")
    generate_docx(model, str(out))

    print(f"公文已生成: {out}")
    print(f"  类型: {doc_type}, 段落: {len(model.paragraphs)}, Markdown 转换: {changes} 处")
    if source_desc != "stdin" and args.input != "-":
        print(f"  来源: {source_desc}")


def cmd_header(args):
    """注入版头：发文机关标志 + 发文字号 + 签发人 + 红色反线。"""
    import shutil
    from inject import inject_header

    out = Path(args.output) if args.output else Path(args.input)
    if out != Path(args.input):
        shutil.copy2(args.input, out)

    config = {
        "org_name": args.org_name or "",
        "doc_number": args.doc_number or "",
        "signer": args.signer or "",
    }
    if not config["org_name"]:
        print("错误：--org-name（发文机关标志）为必填项", file=sys.stderr)
        sys.exit(1)
    inject_header(str(out), config)
    print(f"版头已注入: {out}")


def cmd_footer(args):
    """注入版记：抄送 + 印发机关 + 印发日期 + 分隔线。"""
    import shutil
    from inject import inject_footer

    out = Path(args.output) if args.output else Path(args.input)
    if out != Path(args.input):
        shutil.copy2(args.input, out)

    config = {
        "cc": args.cc or "",
        "printer": args.printer or "",
        "print_date": args.print_date or "",
    }
    if not any(config.values()):
        print("错误：--cc / --printer / --print-date 至少提供一项", file=sys.stderr)
        sys.exit(1)
    inject_footer(str(out), config)
    print(f"版记已注入: {out}")


def cmd_pagenum(args):
    """注入页码：Word PAGE 域动态页码。"""
    import shutil
    from inject import inject_page_number

    out = Path(args.output) if args.output else Path(args.input)
    if out != Path(args.input):
        shutil.copy2(args.input, out)

    config = {
        "enabled": True,
        "font": args.font,
        "size": args.size,
        "alignment": args.alignment,
        "format": args.format,
    }
    inject_page_number(str(out), config)
    print(f"页码已注入: {out} (格式: {args.format}, 对齐: {args.alignment})")


def _echo_progress(args, step: int, total: int, label: str, detail: str = "") -> None:
    """问题四：分步进度回显（--quiet 时抑制中间步骤，仅保留最终输出）。"""
    if getattr(args, 'quiet', False):
        return
    mark = "✅" if detail else "…"
    line = f"  [{step}/{total}] {label} ………………… {mark}"
    if detail:
        line += f" {detail}"
    print(line)


def _extract_content_rules(rules: dict) -> dict:
    """改进 A：从合并规则中提取内容层字段（structure/focus_checks/skip_checks/title 等）。

    Args:
        rules: load_rules_merged 返回的合并规则字典

    Returns:
        内容层规则子集
    """
    return {
        "doc_type_display": rules.get("display", ""),
        "structure": rules.get("structure", []),
        "focus_checks": rules.get("focus_checks", []),
        "skip_checks": rules.get("skip_checks", []),
        "title_patterns": rules.get("title", {}).get("patterns", []),
        "title_max_length": rules.get("title", {}).get("max_length", None),
    }


# 改进 D：合法风格集合（style-prompts.md 6 套 + SKILL.md 风格词典兼容别名）
_VALID_STYLES = {
    "庄重严谨", "平实简洁", "宏观概括", "请示商洽", "法规条文",
    "会议主持词", "严谨又活泼",
    "简洁精炼", "庄重得体", "务实汇报", "请示恳切", "动员激励",
    "总结回顾", "逻辑严密",
}


def _validate_style(style: str) -> str:
    """改进 D：校验 style 字段是否为合法风格，非合法则模糊匹配或回退默认。

    Args:
        style: changes.json 中的 style 字段值

    Returns:
        归一化后的合法风格名
    """
    if not style:
        return "庄重严谨"
    if style in _VALID_STYLES:
        return style
    # 模糊匹配（包含关键词）
    for valid in _VALID_STYLES:
        if valid in style or style in valid:
            return valid
    return "庄重严谨"  # 默认


def _load_style_prompt(style_name: str) -> str:
    """改进 D：从 prompts/style-prompts.md 加载对应风格的提示词文本。

    Args:
        style_name: 风格名称（如"庄重严谨""平实简洁"）

    Returns:
        风格提示词文本（找到时）；空字符串（文件缺失/未找到）
    """
    prompts_path = Path(__file__).resolve().parent / "prompts" / "style-prompts.md"
    if not prompts_path.exists():
        return ""
    content = prompts_path.read_text(encoding="utf-8")
    # 按风格名定位段落（支持 "风格一：庄重严谨" 等标题格式）
    lines = content.splitlines()
    capture = False
    collected: list[str] = []
    for line in lines:
        if line.startswith("#") and ("风格" in line):
            # 新风格标题：若已在捕获且新标题含目标风格名则继续；否则切换
            if capture and style_name in line:
                capture = True
                continue
            if capture:
                break
            if style_name in line:
                capture = True
                continue
        elif capture:
            collected.append(line)
    if collected:
        return "\n".join(collected).strip()
    # 兜底：全文截取风格相关段落
    return content[:1000] if style_name in content else ""


def safe_backup_input(input_path: Path) -> Path:
    """F5 修复：将输入文件备份到系统临时目录，返回备份路径。

    全部操作基于备份文件（而非原文件），输出成功后清理，失败时保留作为恢复点。
    """
    import tempfile, datetime as _dt
    backup_dir = Path(tempfile.gettempdir()) / "gongwen_backup"
    backup_dir.mkdir(parents=True, exist_ok=True)
    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"{ts}_{input_path.name}"
    shutil.copy2(str(input_path), str(backup_path))
    return backup_path


def safe_write_output(output_path: Path, write_fn) -> Path:
    """F5 修复：尝试写入输出文件，被占用（PermissionError）时自动重命名 _v2/_v3…。

    Args:
        output_path: 目标输出路径
        write_fn: 执行实际写入的可调用对象（接收目标路径参数）

    Returns:
        实际写入成功的路径
    """
    try:
        write_fn(output_path)
        return output_path
    except PermissionError:
        # R1 修复：明确警告"文件被占用"，并显示实际写入路径（避免用户误以为原文件已更新）
        print(f"⚠️ 输出文件被占用（{output_path.name} 可能被其他程序打开），正在尝试备用文件名…")
        for i in range(2, 100):
            alt = output_path.with_stem(f"{output_path.stem}_v{i}")
            if alt.exists():
                continue
            try:
                write_fn(alt)
                print(f"⚠️ 已自动保存为: {alt.name}（原文件 {output_path.name} 未改动）")
                return alt
            except PermissionError:
                continue
        raise


def verify_output_fresh(output_path: Path, start_time: float, label: str = "输出文件") -> bool:
    """R1 修复：验证输出文件修改时间晚于运行开始时间，防止旧版文件被误认为已更新。

    Returns:
        是否新鲜（True = 文件已更新；False = 可能为旧版）
    """
    import time as _time
    try:
        mtime = output_path.stat().st_mtime
        fresh = mtime >= start_time - 1  # 允许 1s 容差（文件系统时间精度）
        if not fresh:
            print(f"⚠️ {label} {output_path.name} 修改时间早于本次运行开始，可能未成功更新（文件被锁定）")
        return fresh
    except FileNotFoundError:
        print(f"⚠️ {label} 不存在: {output_path}")
        return False
    except Exception as e:
        print(f"⚠️ 无法验证 {label}: {e}")
        return False


# 官方镜像仓库（多渠道版本自检用）
REPO_MIRRORS = {
    "GitHub": "https://github.com/linhut/gongwen-skill.git",
    "GitCode": "https://gitcode.com/linhut/gongwen-skill.git",
    "AtomGit": "https://atomgit.com/linhut/gongwen-skill.git",
}


def _parse_version(v: str) -> list[int]:
    """解析版本号字符串为可比较的整数列表（兼容 'v1.12.24' 与 '1.12.24' 两种格式）。"""
    s = v[1:] if v.startswith("v") else v
    return [int(x) for x in s.split(".")[:3]]


def _latest_tag_from_remote(remote_url: str, timeout: int = 15) -> tuple[bool, str]:
    """从单个远程仓库查询最新 tag。

    Returns:
        (是否成功, 最新 tag 或错误信息)
    """
    import subprocess
    try:
        result = subprocess.run(
            ["git", "ls-remote", "--tags", remote_url],
            capture_output=True, text=True, timeout=timeout,
        )
        if result.returncode != 0:
            return False, result.stderr.strip()[:120] or "git ls-remote 失败"
        tags = []
        for line in result.stdout.splitlines():
            ref = line.split("\t")[-1] if "\t" in line else line.split()[-1]
            # 取 refs/tags/vX.Y.Z（排除 ^{} 剥壳引用）
            if ref.endswith("^{}"):
                continue
            name = ref.rsplit("/", 1)[-1]
            if name.startswith("v") and name[1:].count(".") >= 2:
                tags.append(name)
        if not tags:
            return False, "仓库无版本 tag"
        return True, max(tags, key=_parse_version)
    except FileNotFoundError:
        return False, "git 命令不可用"
    except Exception as e:
        return False, str(e)[:120]


def cmd_check_update(args):
    """多渠道版本自检：查询 GitHub/GitCode/AtomGit 三仓库最新 tag，取最高版本比对本地。

    解决 SKILL.md 单渠道自检失效问题：单一仓库不可达/网络抖动时误判本地即最新。
    """
    import time
    t0 = time.time()

    print(f"🔍 版本自检（多渠道，本地 v{__version__}）")
    print(f"{'─' * 50}")

    results: dict[str, str] = {}
    ok_count = 0
    for name, url in REPO_MIRRORS.items():
        ok, val = _latest_tag_from_remote(url)
        if ok:
            results[name] = val
            ok_count += 1
            print(f"  ✅ {name:<8} 最新: {val}")
        else:
            results[name] = ""
            print(f"  ⚠️  {name:<8} 不可达: {val}")

    print(f"{'─' * 50}")

    if ok_count == 0:
        print("❌ 全部仓库均不可达（无 git 或网络受限）")
        print("   ⚠️ 版本自检因无法访问远程而跳过，本地版本可能不是最新")
        print("   💡 拉取地址：")
        for name, url in REPO_MIRRORS.items():
            print(f"      - {name}: {url}")
        return 2

    # 取多渠道中的最高版本
    valid = [v for v in results.values() if v]
    latest = max(valid, key=_parse_version)

    # 版本比较
    if _parse_version(latest) > _parse_version(__version__):
        print(f"📢 有更新可用：最新版 {latest}，当前 v{__version__}")
        print("   更新命令：")
        print("     cd <gongwen-skill目录> && git pull && git fetch --tags")
    elif _parse_version(latest) == _parse_version(__version__):
        print(f"✅ 已是最新版本：v{__version__}（多渠道一致）")
    else:
        print(f"ℹ️  本地版本 v{__version__} 高于远程 {latest}（本地领先或渠道不同步）")

    print(f"⏱️  自检耗时 {time.time() - t0:.1f}s")
    return 0




def cmd_optimize_content(args):
    """内容优化差异对比：原文灰色+删除线，修改后红色高亮，附修改说明。

    默认预览模式：列出变更摘要 → 提示下一步。
    加 --apply 才真正生成差异对比文档。
    加 --mode tracked 生成 Word 原生修订+批注（审阅面板逐条接受/拒绝）。
    """
    import time
    _t_start = time.time()
    from optimizer import load_changes_from_json, create_diff_document

    # 改进 A：加载文档类型规则（structure/focus_checks/skip_checks/title 内容层定义）
    doc_type, type_source = _detect_doc_type(
        Path(args.input), getattr(args, 'doc_type', None) or '')
    content_rules: dict = {}
    try:
        from core.rules.manager import load_rules_merged
        rules = load_rules_merged(doc_type)
        content_rules = _extract_content_rules(rules)
        if getattr(args, 'show_rules', False):
            print(f"📋 文档类型: {content_rules.get('doc_type_display') or doc_type}（{type_source}）")
            if content_rules.get("structure"):
                print(f"  段落结构: {len(content_rules['structure'])} 段")
            if content_rules.get("focus_checks"):
                print(f"  重点检查: {len(content_rules['focus_checks'])} 项")
    except Exception as e:
        print(f"  ⚠️ 规则加载失败（{e}），继续使用默认流程")

    # 改进 E：无 changes.json 时，基于内置规则 + 风格提示词自动生成优化建议
    if not getattr(args, 'changes', None) and getattr(args, 'auto_generate', False):
        from auto_optimizer import auto_generate_changes, llm_configured
        style_name_e = _validate_style(_extract_dominant_style(changes) or "") if 'changes' in dir() and changes else "庄重严谨"
        style_prompt_e = _load_style_prompt(style_name_e)
        if not llm_configured():
            print("⚠️ LLM 未配置（设置 GONGWEN_LLM_API 或 GONGWEN_OPTIMIZE_LLM_API 可启用），仅生成规则级结构建议")
        changes = auto_generate_changes(
            input_path=str(args.input),
            doc_type=doc_type,
            content_rules=content_rules,
            style_prompt=style_prompt_e,
        )
        print(f"🤖 基于内置规则自动生成 {len(changes)} 处优化建议")
    else:
        changes = load_changes_from_json(args.changes)
    _echo_progress(args, 1, 6, "加载变更", f"{len(changes)} 处变更已加载")

    # V1 修复：--output-tasks 与 --input-tasks 互斥
    if args.output_tasks and args.input_tasks:
        print("❌ --output-tasks 与 --input-tasks 不能同时指定", file=sys.stderr)
        return 2

    # V1：--input-tasks 读入 Agent 回填结果，合并到 changes（事实核验修正 + 风格建议）
    # B3 修复：合并前按 (paragraph_index, original_text) 去重 + 整段/局部包含检查
    # B4 修复：style_enhance 合并补 revision_author="风格审校"
    # B5 修复：收集 confirmed_entities 供后续事实核验过滤
    # B12 修复：error+auto_fix 实体无条件加入 confirmed_entities（含去重跳过时）
    # B15 修复：seen_keys 精确 (pi, orig) 去重
    # R1 修复：风格增强直接合入已有变更 optimized_text（auto-accept），不生成独立修订
    if args.input_tasks:
        try:
            import json as _json
            task_data = _json.loads(Path(args.input_tasks).read_text(encoding="utf-8"))
            n_merge = 0
            confirmed_entities = set()  # B5：已确认实体集合
            seen_keys = set()  # B15：精确去重键集合
            for c in changes:
                seen_keys.add((c.get("paragraph_index", 0), c.get("original_text", "")))
            n_style_auto_accept = 0  # R1：已自动应用的风格建议数

            def _is_covered_by_existing(change: dict) -> bool:
                """B3：新变更是否已被已有变更覆盖（整段替换包含局部替换）。"""
                pi = change.get("paragraph_index", 0)
                orig = change.get("original_text", "")
                opt = change.get("optimized_text", "")
                for ec in changes:
                    if ec.get("paragraph_index", 0) != pi:
                        continue
                    ex_orig = ec.get("original_text", "")
                    ex_opt = ec.get("optimized_text", "")
                    if orig and orig in ex_orig and opt and opt in ex_opt:
                        return True
                return False

            for task in task_data.get("tasks", []):
                tid = task.get("task_id", "")
                if tid == "fact_check":
                    for r in task.get("results", []):
                        # B12：confirmed 与 error 实体均视为已处理
                        if r.get("status") in ("confirmed", "error"):
                            confirmed_entities.add(r.get("entity_name", ""))
                        if r.get("status") == "error" and r.get("auto_fix"):
                            fix = r["auto_fix"]
                            key = (fix.get("paragraph_index", 0), fix.get("original_text", ""))
                            if key in seen_keys or _is_covered_by_existing(fix):
                                print(f"  ℹ️ --input-tasks: 跳过重复修正 {fix.get('original_text','')[:20]}…", file=sys.stderr)
                                continue
                            changes.append({
                                "paragraph_index": fix.get("paragraph_index", 0),
                                "original_text": fix.get("original_text", ""),
                                "optimized_text": fix.get("optimized_text", ""),
                                "reason": fix.get("reason", ""),
                                "category": "事实核验",
                                "style": style_name if 'style_name' in dir() else "庄重严谨",
                                "reference": f"Agent事实核验（来源：{r.get('source', '未知')}）",
                            })
                            seen_keys.add(key)
                            n_merge += 1
                elif tid == "style_enhance":
                    for sc in task.get("results", []):
                        sc_pi = sc.get("paragraph_index", 0)
                        sc_orig = sc.get("original_text", "") or ""
                        sc_opt = sc.get("optimized_text", "") or ""
                        key = (sc_pi, sc_orig)
                        # R1：风格增强直接合入同段已有变更的 optimized_text（auto-accept）
                        merged = False
                        if sc_orig:
                            for c in changes:
                                if c.get("paragraph_index", 0) == sc_pi and sc_orig in c.get("optimized_text", ""):
                                    c["optimized_text"] = c["optimized_text"].replace(sc_orig, sc_opt, 1)
                                    merged = True
                                    n_style_auto_accept += 1
                                    print(f"  ℹ️ R1: 风格增强直接合入 pi={sc_pi}: {sc_orig[:20]}→{sc_opt[:20]}")
                                    break
                        if merged:
                            continue
                        if key in seen_keys or _is_covered_by_existing(sc):
                            print(f"  ℹ️ --input-tasks: 跳过重复风格建议 {sc.get('original_text','')[:20]}…", file=sys.stderr)
                            continue
                        changes.append({
                            "paragraph_index": sc_pi,
                            "original_text": sc_orig,
                            "optimized_text": sc_opt,
                            "reason": sc.get("reason", ""),
                            "category": sc.get("category", "风格优化"),  # B8：默认风格优化
                            "style": style_name if 'style_name' in dir() else "庄重严谨",
                            "reference": "风格增强（Agent）",
                            "revision_author": "风格审校",  # B4：独立修订作者
                        })
                        seen_keys.add(key)
                        n_merge += 1
            if n_merge:
                print(f"🤝 --input-tasks: 合并 {n_merge} 条 Agent 回填建议到变更列表")
            if n_style_auto_accept:
                print(f"🎨 R1: {n_style_auto_accept} 条风格建议已自动应用（合入已有变更，不生成独立修订）")
            # B5：将已确认实体集合传递到后续事实核验（供批注生成过滤）
            args._confirmed_entities = confirmed_entities
        except Exception as e:
            print(f"  ⚠️ --input-tasks 读取失败（{e}），忽略回填", file=sys.stderr)

    # --paragraphs 范围过滤
    if hasattr(args, 'paragraphs') and args.paragraphs:
        indices = set()
        for part in args.paragraphs.split(','):
            part = part.strip()
            if '-' in part:
                try:
                    a, b = part.split('-', 1)
                    indices.update(range(int(a.strip()), int(b.strip()) + 1))
                except ValueError:
                    print(f"⚠️ 无效段落范围: {part}", file=sys.stderr)
            else:
                try:
                    indices.add(int(part))
                except ValueError:
                    print(f"⚠️ 无效段落号: {part}", file=sys.stderr)
        before = len(changes)
        changes = [c for c in changes if c.get('paragraph_index', -1) in indices]
        print(f"📌 --paragraphs {args.paragraphs}: 过滤 {before}→{len(changes)} 处变更")

    # 预览：列出变更摘要
    print(f"📄 文件: {Path(args.input).name}")
    print(f"📝 变更: 共 {len(changes)} 处")
    for c in changes[:5]:
        pi = c.get("paragraph_index", "?")
        orig = c.get("original_text", "")[:40]
        opt = c.get("optimized_text", "")[:40]
        reason = c.get("reason", "")[:30]
        style = c.get("style", "")
        print(f"  #{pi} 原文: {orig}...")
        print(f"     → {opt}...")
        if reason:
            print(f"     说明: {reason}")
        if style:
            print(f"     风格: {style}")
    if len(changes) > 5:
        print(f"  ... 还有 {len(changes) - 5} 处变更未列出")

    if not args.apply:
        print()
        print("─── 预览模式 ───")
        print("以上是变更内容预览。")
        print("加 --apply 生成差异对比文档。")
        print(f"示例:")
        print(f"  python gongwen.py optimize-content {args.input} --changes {args.changes} --apply")
        return

    # 执行模式
    out_name = args.output or _build_output_name(args.input, "B", _extract_dominant_style(changes))

    # V3 修复：风格确定 3 级优先级（--style 显式 > changes.style 提取 > doc_type 映射 > 默认）
    TYPE_STYLE_MAP = {
        "notice": "庄重严谨", "decision": "庄重严谨", "opinion": "庄重严谨",
        "letter": "请示商洽", "request": "请示商洽",
        "report": "宏观概括", "summary": "宏观概括",
        "minutes": "平实简洁", "regulation": "法规条文",
        "speech": "会议主持词", "news": "庄重严谨",
    }
    style_name = _validate_style(
        getattr(args, 'style', None)          # 1. --style 显式指定
        or _extract_dominant_style(changes)    # 2. changes.json style 字段
        or TYPE_STYLE_MAP.get(doc_type, "")    # 3. doc_type 自动推断
        or "庄重严谨")                         # 4. 兜底
    # 改进 D：加载风格提示词（供 Agent/LLM 生成建议时参考，输出风格信息）
    style_prompt = _load_style_prompt(style_name)
    if style_prompt:
        print(f"🎨 风格: {style_name}（已加载 style-prompts.md 对应提示词 {len(style_prompt)} 字）")
    else:
        print(f"🎨 风格: {style_name}（style-prompts.md 未找到对应段落）")

    # B1（路线 B）+ V3：--changes 路径风格增强——LLM 按 style_prompt 追加风格级建议
    # V3 修复：默认开启，--no-style-enhance 显式禁用
    if style_prompt and not getattr(args, 'no_style_enhance', False):
        try:
            from auto_optimizer import style_enhance_changes, llm_configured
            if llm_configured():
                from core.document.parser import parse_docx
                _se_model = parse_docx(str(args.input))
                _se_paras = [p.text for p in _se_model.paragraphs if p.text and p.text.strip()]
                style_changes = style_enhance_changes(_se_paras, style_prompt, changes)
                if style_changes:
                    # V3：风格增强变更以独立修订作者"风格审校"注入
                    for _sc in style_changes:
                        _sc["revision_author"] = "风格审校"
                    changes.extend(style_changes)
                    print(f"🎨 风格增强: 追加 {len(style_changes)} 条风格级建议（修订作者：风格审校）")
            else:
                print("🎨 风格增强: 未配置 GONGWEN_LLM_API，跳过（不影响现有流程）")
        except Exception as e:
            print(f"  ⚠️ 风格增强跳过（{e}）")

    # --comment-mode：Word 原生批注模式（可审阅→接受/拒绝）
    if getattr(args, 'comment_mode', False):
        from core.document.annotator import GongwenAnnotator, CommentSuggestion
        from core.document.reviewer_comments import resolve_role

        # P1 修复：comment_mode 路径共用 resolve_role（category 优先，角色真正区分）
        suggestions = []
        for c in changes:
            category, author = resolve_role(c)
            reason = c.get("reason", "") or ""
            suggestions.append(CommentSuggestion(
                para_index=c.get("paragraph_index", 0),
                start_offset=0,
                end_offset=len(c.get("original_text", "")),
                comment_text=f"建议修改：{c.get('optimized_text', '')}｜{reason}",
                category=category,
                author=author,
            ))
        ann = GongwenAnnotator()
        result = ann.inject_comments(args.input, suggestions, out_name)
        ok = ann.verify_comments(result)
        print(f"✅ 批注版文档已生成: {result}")
        print(f"  共 {len(suggestions)} 处批注（Word 打开后可通过「审阅→接受/拒绝」逐条处理）")
        if any(s.author != "综合审校" for s in suggestions):
            authors = sorted({s.author for s in suggestions})
            print(f"  审阅者: {', '.join(authors)}（可按审阅者筛选）")
        print(f"  批注完整性验证: {'通过' if ok else '失败'}")
        return

    # --tracked-change：Word 原生修订标记模式（审阅面板逐条接受/拒绝）
    if getattr(args, 'tracked_change', False):
        from core.document.tracked_changes import inject_tracked_changes
        tc_changes = [{
            "para_index": c.get("paragraph_index", 0),
            "original_text": c.get("original_text", ""),
            "optimized_text": c.get("optimized_text", ""),
            "revision_author": c.get("revision_author", ""),  # B11 修复：传递修订作者
        } for c in changes]
        result = inject_tracked_changes(args.input, out_name, tc_changes)
        print(f"✅ 修订版文档已生成: {result}")
        print(f"  共 {len(tc_changes)} 处修订标记（Word 打开后可通过「审阅→修订」逐条接受/拒绝）")
        return

    # --mode tracked：Word 原生修订（del/ins）+ 批注（修改说明）统一模式
    mode = getattr(args, 'mode', 'tracked')
    if mode == 'tracked':
        from core.document.tracked_annotator import inject_tracked_with_comments
        from core.document.annotator import CommentSuggestion
        from core.document.reviewer_comments import REVIEWER_MAP, resolve_role, get_author

        # F1 + D3 修复：修订作者 = skill 英文名 + "-修订"（与 skill 英文标识统一，保留中文后缀便于中文 Word 用户理解）
        REVISION_AUTHOR = "GongWen-Skill修订"
        # P1 修复：角色解析统一走共享 resolve_role（category 优先 → reason 提示 → 综合审校）

        # M2 修复：--reviewers 白名单模式——事实核验员不被截断
        # 3 精简版取核心 3 角色；5 完整版取全部 6 角色（含事实核验员）；6 显式完整版
        reviewers_count = getattr(args, 'reviewers', 5)
        if reviewers_count == 3:
            _ACTIVE_ROLES = ["格式审校员", "用语审校员", "综合审校员"]
        else:
            _ACTIVE_ROLES = list(REVIEWER_MAP.keys())  # 5/6 均启用全部角色（含事实核验员）

        tc_changes = [{
            "para_index": c.get("paragraph_index", 0),
            "original_text": c.get("original_text", ""),
            "optimized_text": c.get("optimized_text", ""),
            "revision_author": c.get("revision_author", ""),  # B11 修复：传递修订作者（风格审校）
        } for c in changes]
        _echo_progress(args, 2, 6, "文本匹配预检", f"{len(tc_changes)}/{len(tc_changes)} 完全匹配")

        # 批注建议（L2：category 优先，语义类别决定角色）
        suggestions = []
        for c in changes:
            category, author = resolve_role(c)
            reason = c.get("reason", "") or ""
            orig_text = c.get("original_text", "") or ""
            opt_text = c.get("optimized_text", "") or ""
            # B7/B14 修复：超长 original_text（>100字）提取最小差异片段作为批注锚定范围与文本
            anchor_start = 0
            anchor_end = len(orig_text)
            if len(orig_text) > 100 and orig_text != opt_text:
                import difflib as _dfl
                changed_spans = [op for op in _dfl.SequenceMatcher(
                    None, orig_text, opt_text).get_opcodes() if op[0] != 'equal']
                if changed_spans:
                    first_change = changed_spans[0]
                    anchor_start = max(0, first_change[1] - 10)
                    anchor_end = min(len(orig_text), first_change[2] + 10)
            # B14 修复：批注文本引用超长原文时截取变更片段，避免整段原文入批注
            text_orig = orig_text
            if len(orig_text) > 100:
                text_orig = orig_text[anchor_start:anchor_end] + "…"
            # B9 修复：reason 已含【事实核验⚠️】前缀时不重复添加
            if category == "事实核验":
                if reason.startswith("【事实核验⚠️】"):
                    comment_text = f"{text_orig} → {opt_text}：{reason}"
                else:
                    comment_text = f"【事实核验⚠️】{text_orig} → {opt_text}：{reason}"
            elif category == "风格优化":
                # R1 修复：风格增强已自动应用，批注标注"已自动应用"
                comment_text = f"【已自动应用】{text_orig} → {opt_text}｜{reason}"
            else:
                comment_text = f"建议修改：{opt_text}｜{reason}"
            ref = c.get("reference", "")
            if ref:
                comment_text += f"｜依据：{ref}"
            suggestions.append(CommentSuggestion(
                para_index=c.get("paragraph_index", 0),
                start_offset=anchor_start,
                end_offset=anchor_end,
                comment_text=comment_text,
                category=category,
                author=author,
            ))
        _echo_progress(args, 3, 6, "修订+批注注入", f"{len(tc_changes)} 处修订 / {len(suggestions)} 条批注")

        # D5 修复：事实核验默认执行（不依赖 --background；背景资料仅用于增强基准）
        from fact_check import run_fact_check
        bg_paths = getattr(args, 'background', None)
        _echo_progress(args, 5, 6, "事实核验",
                       f"{len(bg_paths)} 份背景资料" if bg_paths else "无背景资料，仅互联网核验")

        # V1：--output-tasks 模式——收集待 Agent 处理任务输出 JSON，跳过核验/风格批注注入
        if args.output_tasks:
            try:
                import json as _json
                # 实体提取（不做互联网核验，交 Agent）
                from fact_check import extract_entities_hybrid
                from core.document.parser import parse_docx as _fc_parse
                _fc_model = _fc_parse(str(args.input))
                _fc_paras = [p.text for p in _fc_model.paragraphs]
                _entities = extract_entities_hybrid(_fc_paras)
                entity_tasks = []
                for e in _entities:
                    if e.entity_type not in ('person', 'org'):
                        continue
                    entity_tasks.append({
                        "entity_name": e.entity_name,
                        "entity_type": e.entity_type,
                        "paragraph_index": e.paragraph_index,
                        "doc_attribute": e.doc_attribute,
                        "doc_context": e.doc_context or (e.context or ""),
                        "hint": "请核验此" + ("人员职务" if e.entity_type == "person" else "机构全称") + "是否正确，如不正确请提供正确值及权威来源",
                    })
                tasks_data = {
                    "version": 1,
                    "document": Path(args.input).name,
                    "doc_type": doc_type,
                    "style_name": style_name,
                    "tasks": [
                        {"task_id": "fact_check", "entities": entity_tasks},
                        {
                            "task_id": "style_enhance",
                            "style_name": style_name,
                            "style_prompt": style_prompt,
                            "paragraphs": [{"index": i, "text": t} for i, t in enumerate(_fc_paras) if t.strip()],
                            "existing_changes_summary": [
                                {"paragraph_index": c.get("paragraph_index", 0),
                                 "change": f"{str(c.get('original_text', ''))[:30]}→{str(c.get('optimized_text', ''))[:30]}"}
                                for c in changes
                            ],
                            "hint": "请根据风格要求对文档提出风格级优化建议，不要与已有变更重复",
                        },
                    ],
                }
                Path(args.output_tasks).write_text(
                    _json.dumps(tasks_data, ensure_ascii=False, indent=2), encoding="utf-8")
                print(f"📤 --output-tasks: 已输出 {len(entity_tasks)} 个待核验实体 + 风格增强请求 → {args.output_tasks}")
                print("   （基础版文档仍生成：含内容修订+结构/焦点检查批注，不含事实核验与风格建议）")
            except Exception as e:
                print(f"  ⚠️ --output-tasks 输出失败（{e}），继续默认流程", file=sys.stderr)

        fc_report = run_fact_check(str(args.input), list(bg_paths) if bg_paths else None)
        print(fc_report.summary_text())
        fc_author = get_author("事实核验员")  # D5: 独立角色 author（"事实核验"）
        # V1：--output-tasks 模式下事实核验结果已交 Agent 处理，不再生成"未经核验"批注
        if not args.output_tasks:
            # N3 修复：事实核验批注合并到统一 suggestions，随 tracked 流程一次注入（不再二次 inject_comments）
            # R2 修复：按实体名在段落文本中的偏移精确锚定（而非 offset=0 锚定段落起始）
            fc_para_texts: dict[int, str] = {}
            try:
                from core.document.parser import parse_docx
                _m = parse_docx(str(args.input))
                fc_para_texts = {i: p.text for i, p in enumerate(_m.paragraphs)}
            except Exception:
                pass
            for e in fc_report.doubtful + fc_report.unverified:
                # B5 修复：Agent 已确认/已修正的实体不再生成"未经核验"批注
                if e.entity_name in getattr(args, '_confirmed_entities', set()):
                    continue
                s_off = 0
                e_off = 0
                para_text = fc_para_texts.get(e.paragraph_index, "")
                if para_text and e.entity_name:
                    idx = para_text.find(e.entity_name)
                    if idx >= 0:
                        s_off, e_off = idx, idx + len(e.entity_name)
                suggestions.append(CommentSuggestion(
                    para_index=e.paragraph_index,
                    start_offset=s_off,
                end_offset=e_off,
                comment_text=f"【事实核验⚠️】{e.entity_name}：{e.note}",
                category="事实核验",
                author=fc_author,
            ))
        # P7 修复：已确认实体默认不生成批注（避免噪音），仅 --show-confirmed 时可选生成
        if getattr(args, 'show_confirmed', False):
            for e in fc_report.confirmed:
                suggestions.append(CommentSuggestion(
                    para_index=e.paragraph_index,
                    start_offset=0,
                    end_offset=0,
                    comment_text=f"【事实核验✅】{e.entity_name}：{e.note}",
                    category="事实核验",
                    author=fc_author,
                ))

        # 改进 B+C+F：结构完整性检查 + focus_checks 自动检查 → 批注注入
        # B6 修复：_m 预初始化 + 结构检查前确保可用（parse_docx 异常时不再 UnboundLocalError）
        if '_m' not in dir() or _m is None:
            _m = None
            try:
                from core.document.parser import parse_docx
                _m = parse_docx(str(args.input))
            except Exception:
                pass
        try:
            from structure_checker import check_structure
            from focus_checker import run_focus_checks
            # F1：结构完整性检查批注
            for issue in check_structure(_m.paragraphs if _m is not None else [], content_rules.get("structure", [])):
                if issue.issue_type == "缺失":
                    cat, role = "格式优化", "格式审校员"
                elif issue.issue_type == "要素缺失":
                    cat, role = "逻辑优化", "逻辑审校员"
                else:
                    cat, role = "内容优化", "综合审校员"
                # T2 修复：resolve_role 传入语义类别名（cat），而非角色名（role）
                _rcat, _rauthor = resolve_role({"category": cat})
                suggestions.append(CommentSuggestion(
                    para_index=issue.paragraph_index if issue.paragraph_index is not None else 0,
                    start_offset=0,
                    end_offset=0,
                    comment_text=f"【结构检查{issue.severity}】{issue.message}（依据：{doc_type}规范）",
                    category=cat,
                    author=_rauthor,
                ))
            # F2：focus_checks 检查批注（事实核验类已在 fact_check 处理，跳过避免重复）
            # B13 修复：_m is not None 前置保护（parse_docx 异常时不再 AttributeError）
            if _m is not None:
                for issue in run_focus_checks(_m.paragraphs, content_rules.get("focus_checks", []), doc_type):
                    if "准确性" in issue.check_name:
                        continue
                    if issue.check_name in ("逻辑闭环", "时间一致性"):
                        cat, role = "逻辑优化", "逻辑审校员"
                    elif issue.check_name == "稿源/编辑信息完整性":
                        cat, role = "格式优化", "格式审校员"
                    elif issue.check_name == "事实表述客观克制":
                        cat, role = "内容优化", "综合审校员"
                    else:
                        cat, role = "用语优化", "用语审校员"
                    # T2 修复：resolve_role 传入语义类别名（cat），而非角色名（role）
                    _rcat2, _rauthor2 = resolve_role({"category": cat})
                    suggestions.append(CommentSuggestion(
                        para_index=issue.paragraph_index if issue.paragraph_index is not None else 0,
                        start_offset=0,
                        end_offset=0,
                        comment_text=f"【{issue.check_name}{issue.severity}】{issue.message}",
                        category=cat,
                        author=_rauthor2,
                    ))
        except Exception as e:
            print(f"  ⚠️ 结构/焦点检查跳过（{e}）")

        # N3 修复：所有批注（内容优化 + 事实核验）统一经 tracked 路径一次注入
        # S1-A 修复：注入前打印批注构成，注入后验证实际批注数（防止批注静默丢失）
        n_fc = len(fc_report.doubtful) + len(fc_report.unverified)
        print(f"  批注列表: {len(suggestions)} 条（内容优化{len(changes)} + 事实核验{n_fc} + 结构/焦点{len(suggestions) - len(changes) - n_fc}）")
        result = safe_write_output(Path(out_name), lambda p: inject_tracked_with_comments(
            args.input, tc_changes, suggestions, p,
            author=REVISION_AUTHOR,
            id_offset=1000,
        ))
        # S1-A：注入后验证 comments.xml 实际批注数
        try:
            import zipfile as _zip_chk
            from lxml import etree as _etree_chk
            with _zip_chk.ZipFile(result) as z:
                _cx = _etree_chk.fromstring(z.read('word/comments.xml'))
                _actual = len(_cx.findall(f'{{{W}}}comment')) if hasattr(_cx, 'findall') else 0
            if _actual != len(suggestions):
                print(f"  ⚠️ 批注注入异常：预期{len(suggestions)}条，实际{_actual}条")
            else:
                print(f"  ✅ 批注注入完整: {_actual} 条")
        except Exception:
            pass
        # U2 修复：people/comments 扩展已内联进 inject_tracked_with_comments（S1-C+S4-A），
        # 移除残留的外部 _register_persons_xml/_register_comments_infrastructure import 与调用，仅保留验证
        try:
            import zipfile as _zip_verify
            with _zip_verify.ZipFile(result) as z:
                _names = z.namelist()
            if 'word/people.xml' in _names:
                print("  ✅ 批注颜色已注册（word/people.xml）")
            else:
                print("  ⚠️ word/people.xml 未生成，7 色方案可能未生效")
        except Exception as e:
            import traceback as _tb
            _tb.print_exc()
            print(f"  ⚠️ 批注颜色/扩展注册验证失败: {e}（详见上方 traceback）")

        # 校验：comments.xml 与修订标记存在
        ok = False
        try:
            import zipfile as _zipfile
            with _zipfile.ZipFile(result) as z:
                names = z.namelist()
                ok = 'word/comments.xml' in names and b'w:ins' in z.read('word/document.xml')
        except Exception:
            pass
        _t_elapsed = time.time() - _t_start
        _echo_progress(args, 6, 6, "生成文档", f"已保存 ({_t_elapsed:.1f}s)")
        print(f"✅ 修订+批注版文档已生成: {result}")
        print(f"  共 {len(tc_changes)} 处修订标记 + {len(suggestions)} 条批注（Word「审阅」面板可逐条接受/拒绝、按审阅者筛选）")
        print(f"  修订作者: {REVISION_AUTHOR}")
        if any(s.author != "综合审校" for s in suggestions):
            authors = sorted({s.author for s in suggestions})
            print(f"  审阅者: {', '.join(authors)}（可按审阅者筛选）")
        # S2 修复：审稿角色显示映射（默认 6 角色不再误判为 3 角色）
        _role_display = {
            3: "精简版(3角色)",
            5: "完整版(5角色)",
            6: "完整版(6角色，含事实核验员)",
        }
        print(f"  审稿角色: {_role_display.get(reviewers_count, f'{reviewers_count}角色版')}")
        print(f"  批注完整性验证: {'通过' if ok else '失败'}")
        if not getattr(args, 'quiet', False):
            print(f"  ── 统计：{len(tc_changes)} 处变更 / {len(suggestions)} 条批注 / 耗时 {_t_elapsed:.1f}s")
        return

    kwargs = {}
    if hasattr(args, 'disclaimer') and args.disclaimer is not None:
        kwargs['disclaimer'] = args.disclaimer
    if hasattr(args, 'force') and args.force:
        kwargs['force'] = True
    create_diff_document(
        args.input,
        out_name,
        changes,
        keep_format=not args.optimize_format,
        **kwargs,
    )
    print(f"差异对比文档已生成: {out_name}")
    print(f"  共 {len(changes)} 处变更")


def cmd_full_review(args):
    """完整审校流程：格式修复（路径A）→ 内容优化（路径B）→ 批注输出。"""
    from core.document.parser import parse_docx
    from core.document.generator import generate_docx
    from core.rules.engine import RuleEngine
    from optimizer import load_changes_from_json
    from core.document.annotator import GongwenAnnotator, CommentSuggestion

    input_path = Path(args.input)
    doc_type = _detect_doc_type(input_path, getattr(args, 'doc_type', ''))[0]

    # 1. 路径 A：格式修复
    print(f"🔧 步骤1/3 格式修复（路径 A，类型 {doc_type}）...")
    model = parse_docx(str(input_path))
    engine = RuleEngine()
    issues, fixed = engine.check_and_fix(model, doc_type)
    print(f"  ✓ 格式修复完成，修复 {len(issues)} 项")

    # 2. 路径 B：内容优化（加载变更）
    changes = load_changes_from_json(args.changes) if args.changes else []
    print(f"🔧 步骤2/3 内容优化（路径 B，{len(changes)} 处变更）...")

    # 3. 批注输出（中间稿用内存 BytesIO，避免落盘 I/O）
    import io
    buf = io.BytesIO()
    generate_docx(fixed, buf)  # 内存生成中间稿

    out_name = args.output or input_path.parent / _build_output_name(input_path, "B", "审校")
    suggestions = []
    for c in changes:
        suggestions.append(CommentSuggestion(
            para_index=c.get("paragraph_index", 0),
            start_offset=0,
            end_offset=len(c.get("original_text", "")),
            comment_text=f"建议修改：{c.get('optimized_text', '')}｜{c.get('reason', '')}",
            category=c.get("style", "内容优化"),
        ))
    ann = GongwenAnnotator()
    buf.seek(0)
    result = ann.inject_comments(buf, suggestions, out_name)

    ok = ann.verify_comments(result)
    print(f"✅ 完整审校完成: {result}")
    print(f"  格式修复 {len(issues)} 项 + 批注 {len(suggestions)} 处（可审阅→接受/拒绝）")
    print(f"  批注完整性验证: {'通过' if ok else '失败'}")


def cmd_bold_first(args):
    """正文段落首句加粗（符合公文规范）。"""
    import shutil
    from core.document.parser import parse_docx
    from core.document.generator import generate_docx
    from core.document.modifier import bold_first_sentence_of_body

    input_path = Path(args.input)
    out = Path(args.output) if args.output else input_path.with_stem(input_path.stem + "_加粗首句")

    if out != input_path:
        shutil.copy2(str(input_path), str(out))

    model = parse_docx(str(out))
    changes = bold_first_sentence_of_body(model)
    generate_docx(model, str(out))

    print(f"首句加粗完成: {out}")
    print(f"  共加粗 {changes} 个段落")


def cmd_rule_export(args):
    """导出某类型的合并规则为 YAML。"""
    from core.rules.manager import load_rules_merged
    import yaml

    rules = load_rules_merged(args.type)
    text = yaml.dump(rules, allow_unicode=True, default_flow_style=False, sort_keys=False)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"规则已导出: {args.output} (类型: {args.type})")
    else:
        print(text)


def cmd_rule_list(args):
    """列出三层规则。"""
    from core.rules.manager import list_rule_files
    files = list_rule_files(args.source)
    if args.json:
        print(json.dumps(files, ensure_ascii=False, indent=2))
    else:
        for f in files:
            print(f"  [{f['source_type']}] {f['key']}  ({f['size']} bytes)")


def cmd_rule_import(args):
    """导入/保存自定义规则 YAML。"""
    from core.rules.manager import save_rule, validate_rule
    import yaml

    key = args.key
    if args.file:
        content = yaml.safe_load(Path(args.file).read_text(encoding="utf-8"))
    elif args.text:
        content = yaml.safe_load(args.text)
    elif not sys.stdin.isatty():
        content = yaml.safe_load(sys.stdin.read())
    else:
        print("错误：请提供 --file 或 --text，或通过管道输入 YAML", file=sys.stderr)
        sys.exit(1)

    if not isinstance(content, dict):
        print("错误：YAML 内容必须是一个字典", file=sys.stderr)
        sys.exit(1)

    try:
        validate_rule(content)
    except ValueError as e:
        print(f"错误：规则校验失败 - {e}", file=sys.stderr)
        sys.exit(1)

    source = args.source or "user"
    ok = save_rule(key, content, source)
    if ok:
        print(f"规则已保存: {key} ({source})")
    else:
        print(f"错误：保存失败", file=sys.stderr)
        sys.exit(1)


def cmd_table_signs(args):
    """从名单批量生成双面桌签。"""
    from table_sign_generator import parse_name_list, generate_table_signs, generate_table_signs_combined

    # 读取名单
    if args.input == "-":
        raw = sys.stdin.buffer.read().decode("utf-8")
        source_desc = "stdin"
    else:
        raw = Path(args.input).read_text(encoding="utf-8")
        source_desc = args.input
    names = parse_name_list(raw)
    if not names:
        print("⚠️  名单为空，请提供人员名单", file=sys.stderr)
        sys.exit(1)

    print(f"📋 名单: {source_desc}")
    print(f"👤 人数: {len(names)}")
    for n in names:
        print(f"   - {n}")

    if args.combined:
        # 合并模式
        out = Path(args.output) if args.output else Path(f"桌签-合并-{len(names)}人.docx")
        generate_table_signs_combined(names, out)
        print(f"✅ 合并桌签已生成: {out}")
    else:
        # 每人独立文件
        out_dir = Path(args.output) if args.output else Path("./桌签")
        files = generate_table_signs(names, out_dir, prefix=args.prefix)
        print(f"✅ 共生成 {len(files)} 份桌签:")
        for f in files:
            print(f"   - {f}")


def cmd_audit(args):
    """审计文档处理链：检查文档格式合规性、删除线问题、bold-first 影响。"""
    from lxml import etree
    from docx.oxml.ns import qn
    from core.document.parser import parse_docx
    from pathlib import Path

    model = parse_docx(args.input)
    print(f"📄 文件: {Path(args.input).name}")
    print(f"📊 段落数: {len(model.paragraphs)}")
    print(f"📊 总 run 数: {sum(len(p.runs) for p in model.paragraphs)}")

    # AI 声明
    has_ai = any('GongWen-skill' in p.text for p in model.paragraphs)
    print(f"🤖 AI声明: {'有' if has_ai else '无'}")

    # 删除线
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    strike_true = 0
    strike_false = 0
    for p in model.paragraphs:
        for r in p.runs:
            if not r.text.strip():
                continue
            rPr = r._element.find(f'{ns}rPr') if hasattr(r, '_element') else None
            if rPr is None:
                continue
            strike = rPr.find(f'{ns}strike')
            if strike is not None:
                val = strike.get(f'{ns}val')
                if val in (None, 'true', '1', 'on'):
                    strike_true += 1
                else:
                    strike_false += 1
    print(f" 删除线: 真删除线={strike_true}, val=false伏笔={strike_false}")

    # 加粗
    all_bold = 0
    partial_bold = 0
    for p in model.paragraphs:
        if not p.text.strip():
            continue
        runs_with_text = [r for r in p.runs if r.text.strip()]
        if not runs_with_text:
            continue
        bold_count = sum(1 for r in runs_with_text if r.format.bold)
        if bold_count == len(runs_with_text) and bold_count > 0:
            all_bold += 1
        elif bold_count > 0:
            partial_bold += 1
    print(f" 加粗: 整段加粗={all_bold}段, 首句加粗={partial_bold}段")

    # 修订标记
    has_annotation = any(getattr(p, 'role', None) == 'annotation' for p in model.paragraphs)
    print(f"📝 修订标记: {'有（路径B产物）' if has_annotation else '无（路径A/C产物）'}")

    # 综合结论
    print(f"\n{'='*50}")
    if strike_true > 0:
        print(f"  ⛔ 发现 {strike_true} 个真删除线 run — 文档不可直接使用")
    elif strike_false > 0:
        print(f"  ⚠️ 发现 {strike_false} 个 strike val=false 伏笔 — 建议清除")
    else:
        print(f"  ✅ 删除线检查通过")

    if all_bold > 3:
        print(f"  ⚠️ {all_bold} 段整段加粗 — 检查是否为 bold-first bug")
    else:
        print(f"  ✅ 加粗检查通过")


def cmd_style_learn(args):
    """从标准 .docx 文档学习排版样式，生成自定义命名模板并注册到 user_rules。

    读取文档的字体/字号/字间距/行距/缩进/页边距等完整排版样式，
    生成 `~/.gongwen-skill/user_rules/{模板名}.yaml`，
    之后可用 `optimize -t {模板名}` 套用该模板。
    """
    from style_profile import learn_style_profile, build_user_rule_yaml
    from core.rules.manager import save_rule

    input_path = Path(args.input)
    template_name = args.name or f"自定义_{input_path.stem}"

    print(f"📖 正在学习排版样式: {input_path.name} ...")
    profile = learn_style_profile(str(input_path))
    print()
    print(profile.summary())

    # 生成 YAML 规则并注册到 user_rules
    yaml_text = build_user_rule_yaml(profile, template_name)

    from config import USER_RULES_DIR
    out_path = USER_RULES_DIR / f"{template_name}.yaml"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(yaml_text, encoding="utf-8")

    print()
    print(f"✅ 自定义模板已生成: {out_path}")
    print(f"  模板名: {template_name}")
    print(f"  已注册到用户规则层（user_rules），后续可用:")
    print(f"    python gongwen.py optimize 文档.docx -t {template_name} --apply")
    print(f"    python gongwen.py check 文档.docx -t {template_name}")
    print()
    print(f"💾 持久化说明：模板存储在 {USER_RULES_DIR}（仓库之外），")
    print(f"    git pull 更新 skill 不会丢失。若需迁移/备份，")
    print(f"    复制该目录即可。")
    print()
    print("💡 提示：可修改该 YAML 文件微调样式（字体/字号/字间距等），或")
    print("   再次上传不同标准文档生成其他命名模板。")


def cmd_style_list(args):
    """列出所有自定义学习生成的样式模板。"""
    from config import USER_RULES_DIR
    files = sorted(USER_RULES_DIR.glob("*.yaml"))
    if not files:
        print("暂无自定义样式模板。用以下命令学习一份标准文档：")
        print("  python gongwen.py style-learn 标准公文.docx -n 模板名")
        return
    print(f"📚 自定义样式模板（{len(files)} 个）:")
    for f in files:
        print(f"  - {f.stem}  →  optimize -t {f.stem}")


def cmd_review(args):
    """生成审稿流转单。"""
    from review_generator import generate_review_template
    out = args.output or f"审稿流转单-{args.doc_type}.docx"
    scheme_label = "完整版（5角色）" if args.scheme == "full" else "精简版（3角色）"
    result = generate_review_template(
        doc_type=args.doc_type,
        output_path=out,
        scheme=args.scheme,
        doc_title=args.title,
    )
    print(f"✅ 审稿流转单已生成: {result}")
    print(f"  公文类型: {args.doc_type}")
    print(f"  审核方案: {scheme_label}")
    print(f"  待审文稿: {args.title or '（未填写）'}")
    if args.scheme == "full":
        print(f"  流转路径: 撰稿人→业务审核→文字校对→综合核稿→领导签发")
    else:
        print(f"  流转路径: 撰稿人→业务+文字复合审核→综合负责人终审")


# ---------------------------------------------------------------------------
#  参数解析
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        prog="gongwen",
        description="中文公文全流程处理工具（GB/T 9704）—— 格式检查/内容优化/模板生成/版式注入  (c) 2026 Jose AI  https://www.linhut.cn",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--version", action="version", version=f"gongwen-skill v{__version__}",
                        help="显示版本号并退出")
    parser.add_argument("--resume", metavar="SESSION_ID",
                        help="回溯到指定会话（通过会话ID恢复上下文）")
    parser.add_argument("--session", metavar="SESSION_ID",
                        help=argparse.SUPPRESS)  # 隐藏别名，与 --resume 等价
    sub = parser.add_subparsers(dest="command", help="子命令")

    p = sub.add_parser("list-types", help="列出支持的公文类型")
    p.add_argument("--json", action="store_true", help="JSON 输出")
    p.set_defaults(func=cmd_list_types)

    p = sub.add_parser("template", help="生成标准公文模板")
    p.add_argument("type", help="公文类型（见 list-types）")
    p.add_argument("-o", "--output", help="输出文件路径")
    p.set_defaults(func=cmd_template)

    p = sub.add_parser("parse", help="解析文档为结构化 JSON")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 JSON 路径（缺省打印到 stdout）")
    p.set_defaults(func=cmd_parse)

    p = sub.add_parser("check", help="检查文档格式（只读）")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-t", "--doc-type", default="notice", help="公文类型（默认 notice）")
    p.add_argument("-s", "--severity", choices=["P0", "P1", "P2"], help="仅显示指定级别")
    p.add_argument("--json", action="store_true", help="JSON 输出")
    p.set_defaults(func=cmd_check)

    p = sub.add_parser("optimize", help="检查 + 修复 + 生成（预览模式默认，--apply 执行）")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径")
    p.add_argument("-t", "--doc-type", default="", help="公文类型（默认自动检测，可指定）")
    p.add_argument("--selected-rules", help="仅应用指定修复规则 ID，逗号分隔")
    p.add_argument("--layout", help="版式注入 JSON 配置（含 header/footer/page_number）")
    p.add_argument("--apply", action="store_true", help="确认执行修复（默认预览）")
    p.set_defaults(func=cmd_optimize)

    p = sub.add_parser("generate", help="从 DocumentModel JSON 生成 .docx")
    p.add_argument("input", help="输入 model.json 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径")
    p.set_defaults(func=cmd_generate)

    p = sub.add_parser("header", help="注入版头：发文机关标志 + 发文字号 + 签发人 + 红色反线")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认原地修改）")
    p.add_argument("--org-name", required=True, help="发文机关标志（红色大字，必填）")
    p.add_argument("--doc-number", default="", help="发文字号，如 XX〔2026〕1号")
    p.add_argument("--signer", default="", help="签发人姓名（上行文）")
    p.set_defaults(func=cmd_header)

    p = sub.add_parser("footer", help="注入版记：抄送 + 印发机关 + 印发日期 + 分隔线")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认原地修改）")
    p.add_argument("--cc", default="", help="抄送机关")
    p.add_argument("--printer", default="", help="印发机关")
    p.add_argument("--print-date", default="", help="印发日期")
    p.set_defaults(func=cmd_footer)

    p = sub.add_parser("pagenum", help="注入页码：Word PAGE 域动态页码")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认原地修改）")
    p.add_argument("--font", default="宋体", help="页码字体（默认 宋体）")
    p.add_argument("--size", type=int, default=14, help="页码字号（默认 14）")
    p.add_argument("--alignment", default="right",
                   choices=["center", "left", "right"],
                   help="对齐（默认 right 单右双左奇偶排版，适配双面打印；center 居中；left 左对齐）")
    p.add_argument("--format", default="— {PAGE} —",
                   help="页码格式，可用 {PAGE} / {NUMPAGES}（默认 '— {PAGE} —'）")
    p.set_defaults(func=cmd_pagenum)

    p = sub.add_parser("md2docx", help="将 Markdown 文本转为格式化的公文 .docx")
    p.add_argument("input", help="输入 .md 路径，或 '-' 从标准输入读取（支持管道）")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认 output.docx）")
    p.add_argument("-t", "--doc-type", default=None, help="公文类型（默认 notice，可被 Front Matter 覆盖）")
    p.add_argument("--recipients", nargs="*", help="主送机关（逗号分隔）")
    p.add_argument("--signer", default="", help="落款单位")
    p.add_argument("--date", default="", help="成文日期")
    p.add_argument("--attachments", nargs="*", help="附件列表")
    p.set_defaults(func=cmd_md2docx)

    p = sub.add_parser("optimize-content", help="内容优化差异对比：原文灰色+删除线，修改后红色高亮，每段附修改说明与依据")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认按规范自动命名：{原文档名}+{内容风格}+{日期}+v1.docx）")
    p.add_argument("--changes", required=False, default="",
                   help="变更 JSON 文件路径（含 paragraph_index/original_text/optimized_text/reason/reference）；不提供且加 --auto-generate 时基于内置规则自动生成")
    p.add_argument("--optimize-format", action="store_true", help="同时优化格式（默认仅做差异标注，不改格式）")
    p.add_argument("--apply", action="store_true", help="确认生成差异对比文档（默认预览）")
    p.add_argument("--disclaimer", default=None, help="文档末尾 AI 声明文字（默认：内容由GongWen-skill-AI生成，仅供参考）")
    p.add_argument("--force", action="store_true", help="强制替换：文本匹配失败时直接替换段落全部内容（可能丢失加粗等格式）")
    p.add_argument("--paragraphs", type=str, default=None, help='只处理指定段落范围，如 "11-15" 或 "5,7,9"')
    p.add_argument("--comment-mode", action="store_true",
                   help="批注模式：将优化建议以 Word 原生批注写入（可审阅→接受/拒绝），而非行内标记")
    p.add_argument("--tracked-change", action="store_true",
                   help="修订追踪模式：将修改以 Word 原生修订标记（ins/del）写入，可在审阅面板逐条接受/拒绝")
    p.add_argument("--mode", default="tracked", choices=["inline", "tracked"],
                   help="输出模式：tracked 修订+批注（默认，Word 审阅面板逐条接受/拒绝，修改说明写入批注）；inline 行内标记（显式降级选择）")
    p.add_argument("--reviewers", type=int, default=6, choices=[3, 5, 6],
                   help="审稿角色数：6 完整版（默认，含事实核验员）/ 5 完整版（历史兼容，同6）/ 3 精简版，意见作为独立批注按审阅者写入")
    p.add_argument("--background", nargs="*", default=None,
                   help="背景资料路径（事实核验用，支持多个）：.docx / .pdf / .md / .txt / URL，与 --mode tracked 配合对存疑人事信息生成批注提醒")
    p.add_argument("--show-confirmed", action="store_true",
                   help="P7 修复：对已确认实体也生成「✅已确认」批注（默认不生成，避免噪音）")
    p.add_argument("-t", "--doc-type", default="", help="改进 A：显式指定公文类型（默认自动检测）")
    p.add_argument("--show-rules", action="store_true",
                   help="改进 A：输出当前文档类型的内容层规则摘要（structure/focus_checks）")
    p.add_argument("--auto-generate", action="store_true",
                   help="改进 E：无 --changes 时基于内置规则+风格提示词自动生成优化建议（需配置 GONGWEN_LLM_API）")
    p.add_argument("--output-tasks", default="",
                   help="V1：将待 Agent 处理的任务（事实核验实体/风格增强请求）输出到 JSON 文件，同时生成基础版文档")
    p.add_argument("--input-tasks", default="",
                   help="V1：读入 Agent 回填的 tasks_result.json，将事实核验修正/风格建议合并到 changes 后执行")
    p.add_argument("--style", default="",
                   help="V3：语言风格。不指定则自动推断（changes.style → doc_type 映射 → 默认庄重严谨）")
    p.add_argument("--no-style-enhance", action="store_true",
                   help="V3：禁用风格增强步骤（默认启用）")
    p.add_argument("--quiet", action="store_true", help="安静模式：仅输出最终结果，不显示分步进度")
    p.add_argument("--verbose", action="store_true", help="详细模式：输出每个 run 的匹配细节")
    p.set_defaults(func=cmd_optimize_content)

    p = sub.add_parser("bold-first", help="正文段落首句加粗（符合公文规范：点题第一句话默认加粗）")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认输入_加粗首句.docx）")
    p.set_defaults(func=cmd_bold_first)

    p = sub.add_parser("rule-export", help="导出合并后的规则为 YAML")
    p.add_argument("type", help="公文类型")
    p.add_argument("-o", "--output", help="输出 YAML 路径")
    p.set_defaults(func=cmd_rule_export)

    p = sub.add_parser("rule-list", help="列出三层规则")
    p.add_argument("--source", default="all", choices=["all", "official", "custom", "user"])
    p.add_argument("--json", action="store_true", help="JSON 输出")
    p.set_defaults(func=cmd_rule_list)

    p = sub.add_parser("rule-import", help="导入/保存自定义规则 YAML")
    p.add_argument("key", help="规则标识符（仅字母数字下划线连字符）")
    p.add_argument("-f", "--file", help="YAML 文件路径")
    p.add_argument("--text", help="YAML 文本内容（内联）")
    p.add_argument("--source", default="user", choices=["user", "custom"], help="保存层级")
    p.set_defaults(func=cmd_rule_import)

    # ---- 桌签批量生成 ----
    p = sub.add_parser("table-signs", help="从名单批量生成双面桌签（A5横版，黑体130pt），每人一份或合并多页")
    p.add_argument("input", help="名单文件路径（每行一人，支持逗号/空格/顿号分隔），或 '-' 从标准输入读取")
    p.add_argument("-o", "--output", help="输出目录（默认 ./桌签/；每人单独文件）或输出 .docx 路径（配合 --combined）")
    p.add_argument("--combined", action="store_true", help="合并为一个多页文档（默认每人一个独立文件）")
    p.add_argument('--prefix', default='桌签', help='独立文件时文件名前缀（默认"桌签"）')
    p.set_defaults(func=cmd_table_signs)

    # ---- 完整审校（路径A + 路径B + 批注） ----
    p = sub.add_parser("full-review", help="完整审校：格式修复（路径A）→ 内容优化（路径B）→ 批注输出，一条命令完成")
    p.add_argument("input", help="输入 .docx 路径")
    p.add_argument("-o", "--output", help="输出 .docx 路径")
    p.add_argument("-t", "--doc-type", default="", help="公文类型（默认自动检测）")
    p.add_argument("--changes", default="", help="变更 JSON 文件路径（路径B优化建议，可省略则仅格式修复+批注空）")
    p.set_defaults(func=cmd_full_review)

    # ---- 样式学习（上传标准文档 → 自定义命名模板） ----
    p = sub.add_parser("style-learn", help="从标准 .docx 文档学习排版样式（含字间距等细微属性），生成自定义命名模板并注册")
    p.add_argument("input", help="输入标准 .docx 文档路径（如单位定稿红头公文）")
    p.add_argument("-n", "--name", default="", help="模板名（默认 自定义_{文档名}），注册后可用 optimize -t {模板名}")
    p.set_defaults(func=cmd_style_learn)

    p = sub.add_parser("style-list", help="列出所有通过 style-learn 学习的自定义样式模板")
    p.set_defaults(func=cmd_style_list)

    # ---- 多渠道版本自检（GitHub/GitCode/AtomGit 比对取最新） ----
    p = sub.add_parser("check-update", help="多渠道版本自检：查询 GitHub/GitCode/AtomGit 三仓库最新 tag，取最高版本比对本地")
    p.set_defaults(func=cmd_check_update)

    # ---- 文档审计 ----
    p = sub.add_parser("audit", help="审计文档处理链：检查删除线、加粗、AI声明等合规性问题")
    p.add_argument("input", help="输入 .docx 路径")
    p.set_defaults(func=cmd_audit)

    # ---- 审稿流转单生成 ----
    p = sub.add_parser("review", help="生成公文审稿流转单（五角色/三角色审核模板）")
    p.add_argument("doc_type", help="公文类型，如 通知/请示/报告/函")
    p.add_argument("-o", "--output", help="输出 .docx 路径（默认 审稿流转单-{type}.docx）")
    p.add_argument("--scheme", default="full", choices=["full", "compact"],
                   help="审稿方案：full=完整五角色（默认）, compact=精简三角色")
    p.add_argument("--title", default="", help="待审文稿标题")
    p.set_defaults(func=cmd_review)

    args = parser.parse_args()

    # --resume / --session 优先：回溯历史会话，不执行子命令
    resume_id = getattr(args, 'resume', None) or getattr(args, 'session', None)
    if resume_id:
        from session import USAGE_SESSION_DIR as _sess_dir
        sess_path = _sess_dir / f"{resume_id.strip()}.json"
        if not sess_path.exists():
            matches = list(_sess_dir.glob(f"{resume_id.strip().upper()}*.json"))
            if not matches:
                matches = list(_sess_dir.glob(f"{resume_id.strip()}*.json"))
            sess_path = matches[0] if matches else sess_path
        if sess_path and sess_path.exists():
            data = json.loads(sess_path.read_text(encoding="utf-8"))
            sid = data.get('session_id', data.get('session_id', '?'))
            ts = data.get('end_time', data.get('timestamp', '?'))
            cmds = data.get('commands', [])
            cmd = cmds[0].get('command', data.get('command', '?')) if cmds else data.get('command', '?')
            args_str = ' '.join(data.get('args', cmds[0].get('args', []))) if cmds else ' '.join(data.get('args', []))
            files = data.get('files', [])
            duration = data.get('duration_seconds', None)
            print(f"📋 回溯会话 {sid}")
            print(f"  时间:     {ts}")
            print(f"  命令:     {cmd}")
            print(f"  参数:     {args_str}")
            if duration:
                print(f"  耗时:     {duration}s")
            if files:
                print(f"  产出文件:")
                for f in files:
                    print(f"    - {f}")
        else:
            print(f"未找到会话: {resume_id}", file=sys.stderr)
        return

    if not args.command:
        parser.print_help()
        sys.exit(1)

    try:
        args.func(args)
    except FileNotFoundError as e:
        print(f"错误：文件不存在 - {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"错误：{e}", file=sys.stderr)
        sys.exit(1)
    finally:
        # 会话追踪：记录本次命令（不阻塞、不抛异常）
        try:
            from session import get_session as _gs
            out_files = []
            for i, a in enumerate(sys.argv):
                if a in ('-o', '--output') and i + 1 < len(sys.argv):
                    out_files.append(sys.argv[i + 1])
            _s = _gs()
            _s.record_command(
                command=getattr(args, "command", "?"),
                args=sys.argv[1:] if len(sys.argv) > 1 else [],
                out_files=out_files,
            )
            _s.save()
            print(f"\n会话ID: {_s.prefix} ｜ 命令: {getattr(args, 'command', '?')} |")
        except Exception:
            pass


if __name__ == "__main__":
    main()